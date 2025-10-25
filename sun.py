import streamlit as st
st.set_page_config(
    page_title="Dimensionnement Solaire Sénégal",
    page_icon="☀️",
    layout="wide"
)

import requests
import json
import re
from urllib.parse import quote
import pandas as pd
import io
from docx import Document
from docx.shared import Pt
import math
import os
from firebase_config import (
    login_user, logout_user, is_user_authenticated, is_admin_user,
    save_quote_to_firebase, get_all_quotes, save_equipment_prices, get_equipment_prices,
    save_client_request, get_all_client_requests, update_client_request_status, initialize_equipment_prices_in_firebase,
    delete_quote, delete_client_request,
    is_admin_email, save_labor_percentages, get_labor_percentages, initialize_labor_percentages_in_firebase,
    clear_labor_percentages_cache, save_accessories_rate, get_accessories_rate, clear_accessories_rate_cache,
    initialize_accessories_rate_in_firebase, get_change_history,
    # Fonctions de gestion de stock
    save_product_to_firebase, get_all_products_from_firebase, update_product_in_firebase, delete_product_from_firebase,
    save_client_to_firebase, get_all_clients_from_firebase,
    save_invoice_to_firebase, get_all_invoices_from_firebase,
    save_stock_movement_to_firebase, get_stock_movements_from_firebase,
    clear_stock_cache, sync_sqlite_to_firebase
)

# Import des nouveaux modules
from sync_products import get_stock_for_dimensioning_product, check_stock_availability, update_stock_after_quote
from invoice_editor import show_invoice_editor
from stock_ui_improvements import create_modern_metric_card, create_stock_alert_card, create_advanced_stock_chart, create_financial_overview, create_interactive_product_table, show_stock_alerts_sidebar

# Fonction pour synchroniser les données locales vers Firebase
def sync_local_to_firebase():
    """Synchronise les données locales PRIX_EQUIPEMENTS vers Firebase"""
    try:
        from firebase_config import save_equipment_prices
        success = save_equipment_prices(PRIX_EQUIPEMENTS)
        if success:
            st.success("✅ Données locales synchronisées vers Firebase avec succès!")
            # Vider le cache pour forcer le rechargement
            get_current_prices.clear()
            return True
        else:
            st.error("❌ Erreur lors de la synchronisation vers Firebase")
            return False
    except Exception as e:
        st.error(f"❌ Erreur de synchronisation: {e}")
        return False

# Fonction pour obtenir les prix actuels (Firebase uniquement)
@st.cache_data(ttl=3600)  # Cache pendant 1 heure
def get_current_prices():
    """Obtient les prix actuels depuis Firebase, avec fallback vers PRIX_EQUIPEMENTS"""
    firebase_prices = get_equipment_prices()
    if firebase_prices:
        # Vérifier si Firebase contient des onduleurs hybrides
        if "onduleurs" in firebase_prices:
            has_hybrid = any(specs.get("type") == "Hybride" for specs in firebase_prices["onduleurs"].values())
            if not has_hybrid:
                st.warning("⚠️ Aucun onduleur hybride trouvé dans Firebase. Synchronisation en cours...")
                if sync_local_to_firebase():
                    # Recharger les données après synchronisation
                    firebase_prices = get_equipment_prices()
        return firebase_prices
    else:
        # Firebase vide, synchroniser les données locales
        st.warning("⚠️ Firebase vide. Synchronisation des données locales en cours...")
        if sync_local_to_firebase():
            # Recharger les données après synchronisation
            firebase_prices = get_equipment_prices()
            if firebase_prices:
                return firebase_prices
        # Utilise les prix par défaut si Firebase n'a pas de données
        return PRIX_EQUIPEMENTS

def clear_prices_cache():
    """Vide le cache des prix pour forcer le rechargement"""
    get_current_prices.clear()

# --- Synchronisation croisée Stock ↔ Dimensionnement ---
STOCK_TO_DIM_CATEGORY = {
    "Panneaux Solaires": "panneaux",
    "Batteries": "batteries",
    "Onduleurs": "onduleurs",
    "Régulateurs": "regulateurs",
    "Regulateurs": "regulateurs",
}

def delete_dimensionnement_article_if_exists(article_name: str, stock_category: str | None = None) -> bool:
    """
    Supprime l'article du module 'dimensionnement' s'il existe, 
    en se basant sur le nom et, si fourni, la catégorie stock.
    Retourne True si une suppression a été effectuée.
    """
    try:
        if not article_name:
            return False
        prices = get_current_prices()
        if not isinstance(prices, dict):
            return False

        # Déterminer les catégories à parcourir
        if stock_category and stock_category in STOCK_TO_DIM_CATEGORY:
            search_categories = [STOCK_TO_DIM_CATEGORY[stock_category]]
        else:
            search_categories = ["panneaux", "batteries", "onduleurs", "regulateurs"]

        updated = prices.copy()
        removed = False
        for cat in search_categories:
            cat_items = updated.get(cat, {})
            if isinstance(cat_items, dict) and article_name in cat_items:
                del updated[cat][article_name]
                removed = True
                break

        if removed:
            if save_equipment_prices(updated):
                clear_prices_cache()
                st.cache_data.clear()
                return True
        return False
    except Exception:
        return False

def delete_stock_product_by_name_if_exists(product_name: str) -> bool:
    """
    Supprime le produit correspondant dans le stock (Firebase) en se basant sur le nom exact.
    Retourne True si une suppression a été effectuée.
    """
    try:
        if not product_name:
            return False
        products = get_all_products_from_firebase()
        if not isinstance(products, dict):
            return False
        for prod_id, prod in products.items():
            if prod.get('nom') == product_name:
                if delete_product_from_firebase(prod_id):
                    clear_stock_cache()
                    return True
                return False
        return False
    except Exception:
        return False

# Configuration de la page définie au début du script (bloc déplacé) 

# Valeurs par défaut pour éviter les erreurs si l’utilisateur n’a pas encore configuré tab1
if 'consommation' not in st.session_state:
    st.session_state.consommation = 10.0  # kWh/jour par défaut
# Valeur par défaut pour le widget de consommation journalière,
# gérée exclusivement via Session State pour éviter les conflits avec `value=`.
if 'conso_journaliere_input' not in st.session_state:
    st.session_state['conso_journaliere_input'] = 10.0
if 'choix' not in st.session_state:
    st.session_state.choix = {
        'type_batterie': 'Lithium',
        'type_onduleur': 'Hybride',
        'voltage': 48
    }

# Base de données complète des prix (en FCFA) - basée sur energiesolairesenegal.com
PRIX_EQUIPEMENTS = {
    "panneaux": {
        "50W Polycristallin": {"prix": 45000, "puissance": 50, "type": "Polycristallin"},
        "100W Polycristallin": {"prix": 75000, "puissance": 100, "type": "Polycristallin"},
        "150W Polycristallin": {"prix": 95000, "puissance": 150, "type": "Polycristallin"},
        "200W Polycristallin": {"prix": 115000, "puissance": 200, "type": "Polycristallin"},
        "250W Polycristallin": {"prix": 140000, "puissance": 250, "type": "Polycristallin"},
        "260W Polycristallin": {"prix": 145000, "puissance": 260, "type": "Polycristallin"},
        "270W Polycristallin": {"prix": 150000, "puissance": 270, "type": "Polycristallin"},
        "280W Polycristallin": {"prix": 155000, "puissance": 280, "type": "Polycristallin"},
        "320W Polycristallin": {"prix": 180000, "puissance": 320, "type": "Polycristallin"},
        "335W Polycristallin": {"prix": 195000, "puissance": 335, "type": "Polycristallin"},
        
        # Ajouts alignés sur energiesolairesenegal.com (prix promo)
        "375W Monocristallin": {"prix": 49174, "puissance": 375, "type": "Monocristallin"},
        "450W Monocristallin": {"prix": 56199, "puissance": 450, "type": "Monocristallin"},
        "550W Monocristallin": {"prix": 65233, "puissance": 550, "type": "Monocristallin"},
        "700W Monocristallin": {"prix": 78000, "puissance": 700, "type": "Monocristallin"},
    },
    "batteries": {
        # Batteries Plomb-Acide (traditionnelles) — prix promo alignés
        "Plomb 100Ah 12V": {"prix": 110395, "capacite": 100, "voltage": 12, "type": "Plomb", "cycles": 500, "decharge_max": 50},
        "Plomb 150Ah 12V": {"prix": 160574, "capacite": 150, "voltage": 12, "type": "Plomb", "cycles": 500, "decharge_max": 50},
        "Plomb 200Ah 12V": {"prix": 210759, "capacite": 200, "voltage": 12, "type": "Plomb", "cycles": 500, "decharge_max": 50},
        
        # Batteries AGM (Absorbed Glass Mat) — prix promo alignés
        "AGM 100Ah 12V": {"prix": 110395, "capacite": 100, "voltage": 12, "type": "AGM", "cycles": 800, "decharge_max": 70},
        "AGM 150Ah 12V": {"prix": 160574, "capacite": 150, "voltage": 12, "type": "AGM", "cycles": 800, "decharge_max": 70},
        "AGM 200Ah 12V": {"prix": 210759, "capacite": 200, "voltage": 12, "type": "AGM", "cycles": 800, "decharge_max": 70},
        "AGM 250Ah 12V": {"prix": 350000, "capacite": 250, "voltage": 12, "type": "AGM", "cycles": 800, "decharge_max": 70},
        
        # Batteries GEL — ajustement 200Ah
        "GEL 100Ah 12V": {"prix": 180000, "capacite": 100, "voltage": 12, "type": "GEL", "cycles": 1200, "decharge_max": 80},
        "GEL 150Ah 12V": {"prix": 270000, "capacite": 150, "voltage": 12, "type": "GEL", "cycles": 1200, "decharge_max": 80},
        "GEL 200Ah 12V": {"prix": 210759, "capacite": 200, "voltage": 12, "type": "GEL", "cycles": 1200, "decharge_max": 80},
        "GEL 250Ah 12V": {"prix": 450000, "capacite": 250, "voltage": 12, "type": "GEL", "cycles": 1200, "decharge_max": 80},
        
        # Batteries Lithium LiFePO4 Normales (12V-24V) — prix promo alignés
        "Lithium 100Ah 12V": {"prix": 450000, "capacite": 100, "voltage": 12, "type": "Lithium", "cycles": 3000, "decharge_max": 90},
        "Lithium 150Ah 12V": {"prix": 650000, "capacite": 150, "voltage": 12, "type": "Lithium", "cycles": 3000, "decharge_max": 90},
        "Lithium 200Ah 12V": {"prix": 850000, "capacite": 200, "voltage": 12, "type": "Lithium", "cycles": 3000, "decharge_max": 90},
        "Lithium 100Ah 24V": {"prix": 750000, "capacite": 100, "voltage": 24, "type": "Lithium", "cycles": 3000, "decharge_max": 90},
        "Lithium 150Ah 24V": {"prix": 950000, "capacite": 150, "voltage": 24, "type": "Lithium", "cycles": 3000, "decharge_max": 90},
        
        # Batteries Lithium Haute Tension (48V et plus) — prix promo alignés
        "Lithium HV 4.8kWh 48V": {"prix": 950000, "capacite": 100, "voltage": 48, "type": "Lithium HV", "cycles": 6000, "decharge_max": 95, "kwh": 4.8},
        "Lithium HV 7.2kWh 48V": {"prix": 1345883, "capacite": 150, "voltage": 48, "type": "Lithium HV", "cycles": 6000, "decharge_max": 95, "kwh": 7.2},
        "Lithium HV 9.6kWh 48V": {"prix": 1103959, "capacite": 200, "voltage": 48, "type": "Lithium HV", "cycles": 6000, "decharge_max": 95, "kwh": 9.6},
        "Lithium HV 12kWh 48V": {"prix": 1650000, "capacite": 250, "voltage": 48, "type": "Lithium HV", "cycles": 6000, "decharge_max": 95, "kwh": 12.0},
        "Lithium HV 14.4kWh 48V": {"prix": 1950000, "capacite": 300, "voltage": 48, "type": "Lithium HV", "cycles": 6000, "decharge_max": 95, "kwh": 14.4},
        
        # Batteries Lithium Très Haute Tension (96V et plus) pour installations industrielles
        "Lithium HV 9.6kWh 96V": {"prix": 1800000, "capacite": 100, "voltage": 96, "type": "Lithium HV", "cycles": 8000, "decharge_max": 98, "kwh": 9.6},
        "Lithium HV 14.4kWh 96V": {"prix": 2500000, "capacite": 150, "voltage": 96, "type": "Lithium HV", "cycles": 8000, "decharge_max": 98, "kwh": 14.4},
    },
    "onduleurs": {
        # Onduleurs Standard (Off-Grid) - Monophasés
        "1000W 12V Pur Sinus": {"prix": 150000, "puissance": 1000, "voltage": 12, "type": "Off-Grid", "phase": "monophase"},
        "1500W 24V Pur Sinus": {"prix": 240000, "puissance": 1500, "voltage": 24, "type": "Off-Grid", "phase": "monophase"},
        "2000W 24V Pur Sinus": {"prix": 350000, "puissance": 2000, "voltage": 24, "type": "Off-Grid", "phase": "monophase"},
        
        # Onduleurs Hybrides (avec MPPT intégré) - Monophasés — prix promo
        "Hybride 1KVA 12V MPPT": {"prix": 151002, "puissance": 1000, "voltage": 12, "type": "Hybride", "mppt": "30A", "phase": "monophase"},
        "Hybride 3KVA 24V MPPT": {"prix": 400482, "puissance": 3000, "voltage": 24, "type": "Hybride", "mppt": "60A", "phase": "monophase"},
        "Hybride 3KVA 48V MPPT": {"prix": 538000, "puissance": 3000, "voltage": 48, "type": "Hybride", "mppt": "80A", "phase": "monophase"},
        "Hybride 5KVA 48V MPPT": {"prix": 750000, "puissance": 5000, "voltage": 48, "type": "Hybride", "mppt": "100A", "phase": "monophase"},
        "Hybride 6KVA 48V MPPT": {"prix": 900000, "puissance": 6000, "voltage": 48, "type": "Hybride", "mppt": "120A", "phase": "monophase"},
        
        # Onduleurs Online (haute qualité) - Monophasés — prix promo
        "Online 2KVA": {"prix": 263137, "puissance": 2000, "voltage": 24, "type": "Online", "phase": "monophase"},
        "Online 3KVA": {"prix": 558049, "puissance": 3000, "voltage": 48, "type": "Online", "phase": "monophase"},
        "Online 6KVA": {"prix": 1220487, "puissance": 6000, "voltage": 48, "type": "Online", "phase": "monophase"},
        "Online 10KVA Mono": {"prix": 1750962, "puissance": 10000, "voltage": 48, "type": "Online", "phase": "monophase"},
        
        # Onduleurs Online Triphasés (haute qualité) — prix promo
        "Online 10KVA 3/3 HF": {"prix": 3157902, "puissance": 10000, "voltage": 48, "type": "Online Tri", "phase": "triphase"},
        "Online 20KVA 3/3 HF": {"prix": 4565499, "puissance": 20000, "voltage": 48, "type": "Online Tri", "phase": "triphase"},
        "Online 30KVA 3/3 HF": {"prix": 5974410, "puissance": 30000, "voltage": 48, "type": "Online Tri", "phase": "triphase"},
        
        # Onduleurs Haute Tension (>180V) pour installations industrielles
        "HV Hybride 10KVA 200V": {"prix": 2500000, "puissance": 10000, "voltage": 200, "type": "Hybride", "mppt": "150A", "phase": "monophase"},
        "HV Hybride 15KVA 300V": {"prix": 3500000, "puissance": 15000, "voltage": 300, "type": "Hybride", "mppt": "200A", "phase": "monophase"},
        "HV Online 20KVA 400V": {"prix": 4500000, "puissance": 20000, "voltage": 400, "type": "Online", "phase": "monophase"},
        "HV Online 30KVA 400V": {"prix": 6500000, "puissance": 30000, "voltage": 400, "type": "Online", "phase": "monophase"},
        "HV Online 50KVA 600V Tri": {"prix": 8500000, "puissance": 50000, "voltage": 600, "type": "Online Tri", "phase": "triphase"},
        "HV Online 100KVA 800V Tri": {"prix": 15000000, "puissance": 100000, "voltage": 800, "type": "Online Tri", "phase": "triphase"},
    },
    "regulateurs": {
        # Régulateurs PWM
        "PWM 10A 12/24V": {"prix": 15000, "amperage": 10, "type": "PWM", "voltage_max": 50},
        "PWM 20A 12/24V": {"prix": 25000, "amperage": 20, "type": "PWM", "voltage_max": 50},
        "PWM 30A 12/24V": {"prix": 35000, "amperage": 30, "type": "PWM", "voltage_max": 50},
        "PWM 40A 12/24V": {"prix": 45000, "amperage": 40, "type": "PWM", "voltage_max": 50},
        
        # Régulateurs MPPT (30% plus efficaces)
        "MPPT 20A 12/24V": {"prix": 45000, "amperage": 20, "type": "MPPT", "voltage_max": 100},
        "MPPT 30A 12/24/48V": {"prix": 65000, "amperage": 30, "type": "MPPT", "voltage_max": 100},
        "MPPT 40A 12/24/48V": {"prix": 85000, "amperage": 40, "type": "MPPT", "voltage_max": 150},
        "MPPT 60A 12/24/48V": {"prix": 120000, "amperage": 60, "type": "MPPT", "voltage_max": 150},
        "MPPT 80A 12/24/48V": {"prix": 160000, "amperage": 80, "type": "MPPT", "voltage_max": 150},
        "MPPT 100A 12/24/48V": {"prix": 200000, "amperage": 100, "type": "MPPT", "voltage_max": 150},
    }
}



# Informations sur les types de batteries
INFO_BATTERIES = {
    "Plomb": {
        "avantages": "✓ Prix bas\n✓ Technologie éprouvée\n✓ Facilement disponible",
        "inconvenients": "✗ Nécessite entretien (ajout d'eau)\n✗ Durée de vie courte (2-3 ans)\n✗ Décharge limitée à 50%",
        "usage": "Petit budget, usage occasionnel"
    },
    "AGM": {
        "avantages": "✓ Sans entretien\n✓ Bonne résistance aux chocs\n✓ Supporte bien la chaleur\n✓ Charge rapide",
        "inconvenients": "✗ Plus cher que le plomb\n✗ Décharge limitée à 70%",
        "usage": "Bon compromis prix/performance, idéal pour le Sénégal"
    },
    "GEL": {
        "avantages": "✓ Sans entretien\n✓ Excellente durée de vie (5-7 ans)\n✓ Décharge profonde possible (80%)\n✓ Supporte bien les températures élevées",
        "inconvenients": "✗ Plus cher que AGM\n✗ Charge plus lente",
        "usage": "Usage intensif, meilleur rapport qualité/durée"
    },
    "Lithium": {
        "avantages": "✓ Durée de vie exceptionnelle (10-12 ans)\n✓ Décharge profonde 90%\n✓ Très léger et compact\n✓ Sans entretien\n✓ Charge ultra-rapide",
        "inconvenients": "✗ Prix élevé (3-4x plus cher)\n✗ Nécessite BMS pour sécurité",
        "usage": "Meilleur investissement long terme, installations modernes"
    },
    "Lithium HV": {
        "avantages": "✓ Durée de vie exceptionnelle (15-20 ans)\n✓ Décharge profonde 95-98%\n✓ Très haute densité énergétique\n✓ BMS avancé intégré\n✓ Charge ultra-rapide\n✓ Idéal pour systèmes 48V+",
        "inconvenients": "✗ Prix très élevé (5-6x plus cher)\n✗ Nécessite onduleurs compatibles HV\n✗ Installation par professionnel requis",
        "usage": "Installations industrielles, systèmes haute puissance, autonomie maximale"
    }
}

# Régions du Sénégal pour la sélection de main d'œuvre
REGIONS_SENEGAL = [
    "Dakar",
    "Thiès", 
    "Saint-Louis",
    "Diourbel",
    "Louga",
    "Fatick",
    "Kaolack",
    "Kaffrine",
    "Tambacounda",
    "Kédougou",
    "Kolda",
    "Ziguinchor",
    "Sédhiou",
    "Matam"
]

# Pourcentages de main d'œuvre par défaut par région (en % du coût total des équipements)
POURCENTAGES_MAIN_OEUVRE_DEFAUT = {
    "Dakar": 15.0,
    "Thiès": 18.0,
    "Saint-Louis": 20.0,
    "Diourbel": 22.0,
    "Louga": 25.0,
    "Fatick": 25.0,
    "Kaolack": 20.0,
    "Kaffrine": 25.0,
    "Tambacounda": 30.0,
    "Kédougou": 35.0,
    "Kolda": 30.0,
    "Ziguinchor": 25.0,
    "Sédhiou": 30.0,
    "Matam": 30.0
}

# Taux accessoires par défaut (en %)
TAUX_ACCESSOIRES_DEFAUT = 15.0

# Estimation de surface des panneaux (approximation)
# Hypothèse réaliste: ~5 m² par kWc installé (modules 375–550W)
SURFACE_PAR_KWC_M2 = 5.0
# Marge d'implantation (espacement, orientation, accès)
MARGE_IMPLANTATION_SURFACE_PCT = 10.0

# Catalogue d'appareils par familles (puissances typiques)
APPAREILS_FAMILLES = {
    "Éclairage": [
        {"nom": "LED 7W", "puissance": 7},
        {"nom": "LED 10W", "puissance": 10},
        {"nom": "Lampe tube 18W", "puissance": 18},
        {"nom": "Néon 36W", "puissance": 36},
        {"nom": "Néon 58W", "puissance": 58},
        {"nom": "Fluocompacte 15W", "puissance": 15},
        {"nom": "Plafonnier LED 24W", "puissance": 24}
    ],
    "Ventilation": [
        {"nom": "Ventilateur 50W", "puissance": 50},
        {"nom": "Ventilateur 75W", "puissance": 75},
        {"nom": "Ventilateur 100W", "puissance": 100}
    ],
    "Électroménager": [
        {"nom": "TV 100W", "puissance": 100},
        {"nom": "Réfrigérateur 150W", "puissance": 150},
        {"nom": "Congélateur 200W", "puissance": 200},
        {"nom": "Machine à laver 500W", "puissance": 500},
        {"nom": "Micro-ondes 1000W", "puissance": 1000}
    ],
    "Informatique": [
        {"nom": "Ordinateur 200W", "puissance": 200},
        {"nom": "Laptop 60W", "puissance": 60},
        {"nom": "Routeur 10W", "puissance": 10},
        {"nom": "Chargeur téléphone 10W", "puissance": 10}
    ],
    "Cuisine": [
        {"nom": "Bouilloire 2000W", "puissance": 2000},
        {"nom": "Plaque électrique 1500W", "puissance": 1500},
        {"nom": "Mixeur 500W", "puissance": 500}
    ],
    "Pompage": [
        {"nom": "Pompe 500W", "puissance": 500},
        {"nom": "Pompe 1000W", "puissance": 1000}
    ],
    "Atelier": [
        {"nom": "Perceuse 600W", "puissance": 600}
    ],
    "Climatisation": [
        {"nom": "Climatiseur 1CV 900W", "puissance": 900},
        {"nom": "Climatiseur 1.5CV 1100W", "puissance": 1100},
        {"nom": "Climatiseur 2CV 1300W", "puissance": 1300},
        {"nom": "Climatiseur 2.5CV 1500W", "puissance": 1500},
        {"nom": "Climatiseur 3CV 1700W", "puissance": 1700}
    ],
    "Eau chaude": [
        {"nom": "Chauffe-eau instantané 3000W", "puissance": 3000},
        {"nom": "Chauffe-eau instantané 5000W", "puissance": 5000},
        {"nom": "Cumulus 50L 1500W", "puissance": 1500},
        {"nom": "Cumulus 100L 2000W", "puissance": 2000}
    ]
}

# Mots-clés simples pour suggestions IA → appareils
APPAREILS_KEYWORDS = {
    "tv": ("Électroménager", "TV 100W"),
    "télé": ("Électroménager", "TV 100W"),
    "television": ("Électroménager", "TV 100W"),
    "frigo": ("Électroménager", "Réfrigérateur 150W"),
    "réfrigérateur": ("Électroménager", "Réfrigérateur 150W"),
    "congelateur": ("Électroménager", "Congélateur 200W"),
    "ventilateur": ("Ventilation", "Ventilateur 75W"),
    "ordi": ("Informatique", "Ordinateur 200W"),
    "ordinateur": ("Informatique", "Ordinateur 200W"),
    "pc": ("Informatique", "Ordinateur 200W"),
    "laptop": ("Informatique", "Laptop 60W"),
    "routeur": ("Informatique", "Routeur 10W"),
    "wifi": ("Informatique", "Routeur 10W"),
    "chargeur": ("Informatique", "Chargeur téléphone 10W"),
    "pompe": ("Pompage", "Pompe 500W"),
    "machine": ("Électroménager", "Machine à laver 500W"),
    "micro": ("Électroménager", "Micro-ondes 1000W"),
    "bouilloire": ("Cuisine", "Bouilloire 2000W"),
    "neon": ("Éclairage", "Néon 36W"),
    "néon": ("Éclairage", "Néon 36W"),
    "fluorescent": ("Éclairage", "Néon 36W"),
    "tube": ("Éclairage", "Lampe tube 18W"),
    "chauffe": ("Eau chaude", "Chauffe-eau instantané 3000W"),
    "chauffeau": ("Eau chaude", "Chauffe-eau instantané 3000W"),
    "douche": ("Eau chaude", "Chauffe-eau instantané 3000W"),
    "cumulus": ("Eau chaude", "Cumulus 50L 1500W")
}

# Fonction pour appeler l'API DeepSeek
def appeler_assistant_ia(question, contexte=""):
    # Priorité aux secrets, fallback à la session
    api_key = None
    try:
        api_key = st.secrets["DEEPSEEK_API_KEY"]
    except Exception:
        api_key = st.session_state.get('api_key', '')

    if not api_key:
        return "⚠️ Veuillez entrer votre clé API DeepSeek dans la barre latérale ou dans secrets.toml."
    
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        prompt = f"""Tu es Pape, un expert en énergie solaire d'Energie Solaire Sénégal. Tu connais bien le climat local (chaleur, humidité) et les meilleures pratiques d'installation.
Contexte: {contexte}
Question: {question}
Réponds de manière claire et pratique en français, avec des conseils adaptés au Sénégal."""
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "Tu es Pape, un expert en installations solaires au Sénégal représentant Energie Solaire Sénégal. Tu aides les utilisateurs à comprendre leurs besoins en énergie solaire et à choisir les bons équipements. Pour les commandes et installations, dirige vers Energie Solaire Sénégal (energiesolairesenegal@gmail.com, +221 77 631 42 25, +221 78 177 39 26, +221 77 250 47 90)."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7
        }
        
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"❌ Erreur API: {response.status_code}"
    except Exception as e:
        return f"❌ Erreur: {str(e)}"

# Variante streaming pour réponses progressives (avec options de concision)
def appeler_assistant_ia_stream(question, contexte="", max_tokens=None, limite_caracteres=None, concis=False):
    api_key = None
    try:
        api_key = st.secrets["DEEPSEEK_API_KEY"]
    except Exception:
        api_key = st.session_state.get('api_key', '')
    if not api_key:
        yield "⚠️ Veuillez entrer votre clé API DeepSeek dans la barre latérale ou dans secrets.toml."
        return
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        consigne_concise = ("Réponds de façon concise en 5–7 phrases max, "
                             "avec des points clés si utile, sans longs détails.") if concis else ""
        prompt = f"""Tu es Pape, un expert en énergie solaire d'Energie Solaire Sénégal. Tu connais bien le climat local (chaleur, humidité) et les meilleures pratiques d'installation.

IMPORTANT: Tu représentes et recommandes Energie Solaire Sénégal, notre partenaire officiel pour tous les équipements et installations solaires au Sénégal.

Informations du partenaire à mentionner quand approprié:
- Nom: Energie Solaire Sénégal
- Adresse: Castor 221 Dakar, Sénégal (en face du terrain de Football)
- Adresse 2: Zac Mbao (pres du rond point AXA)
- Email: energiesolairesenegal@gmail.com
- Téléphones: +221 77 631 42 25 ou +221 78 177 39 26 ou +221 77 250 47 90
- Site web: energiesolairesenegal.com

Pour les commandes, devis personnalisés ou installations, dirige toujours vers Energie Solaire Sénégal.

Contexte: {contexte}
Question: {question}
{consigne_concise}
Réponds de manière claire et pratique en français, avec des conseils adaptés au Sénégal."""
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "Tu es Pape, un expert en installations solaires au Sénégal représentant Energie Solaire Sénégal. Tu aides les utilisateurs à comprendre leurs besoins en énergie solaire et à choisir les bons équipements. Pour les commandes et installations, dirige vers Energie Solaire Sénégal (energiesolairesenegal@gmail.com, +221 77 631 42 25, +221 78 177 39 26, +221 77 250 47 90)."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "stream": True
        }
        if max_tokens:
            data["max_tokens"] = max_tokens
        resp = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers,
            json=data,
            stream=True,
            timeout=60
        )
        if resp.status_code != 200:
            yield f"❌ Erreur API: {resp.status_code}"
            return
        chars_out = 0
        for raw in resp.iter_lines():
            if not raw:
                continue
            try:
                line = raw.decode("utf-8") if isinstance(raw, (bytes, bytearray)) else str(raw)
            except Exception:
                line = str(raw)
            if line.startswith("data: "):
                payload = line[6:].strip()
                if payload == "[DONE]":
                    break
                try:
                    obj = json.loads(payload)
                    choice = (obj.get("choices") or [{}])[0]
                    delta = choice.get("delta", {})
                    content = delta.get("content") or (choice.get("message", {}) or {}).get("content")
                    if content:
                        if concis and (limite_caracteres is not None):
                            restant = max(0, limite_caracteres - chars_out)
                            if restant <= 0:
                                break
                            morceau = content[:restant]
                            chars_out += len(morceau)
                            if morceau:
                                yield morceau
                            if chars_out >= (limite_caracteres or 0):
                                break
                        else:
                            yield content
                except Exception:
                    continue
    except Exception as e:
        yield f"❌ Erreur: {str(e)}"
        return

# --- Prix en ligne (energiesolairesenegal.com) ---
@st.cache_data(ttl=3600)
def _rechercher_url_produit_solairesenegal(nom: str):
    try:
        url = f"https://energiesolairesenegal.com/?s={quote(nom)}"
        r = requests.get(url, timeout=20)
        html = r.text
        m = re.search(r'href="(https?://[^"]*solairesenegal\\.com/(?:produit|product)/[^"]+)"', html, re.IGNORECASE)
        if m:
            return m.group(1)
    except Exception:
        return None
    return None


def _extraire_prix_fcfa(html: str):
    candidates = re.findall(r'([0-9]{3,9}(?:[\s.,][0-9]{3})*)\s*(?:FCFA|F\s*CFA|CFA)', html, flags=re.IGNORECASE)
    if not candidates:
        return None
    def _to_int(s):
        digits = re.sub(r'[^0-9]', '', s)
        try:
            return int(digits)
        except Exception:
            return None
    vals = [v for v in (_to_int(c) for c in candidates) if v is not None]
    plausible = [v for v in vals if 1000 <= v <= 20000000]
    return max(plausible) if plausible else (max(vals) if vals else None)

@st.cache_data(ttl=3600)
def obtenir_prix_depuis_site(nom_item: str):
    url_prod = _rechercher_url_produit_solairesenegal(nom_item)
    if not url_prod:
        return None, None
    try:
        r = requests.get(url_prod, timeout=20)
        prix = _extraire_prix_fcfa(r.text)
        return prix, url_prod
    except Exception:
        return None, url_prod

@st.cache_data(ttl=86400)
def get_pvgis_monthly_psh(lat, lon, optimalangles=True, angle=None, aspect=None, raddatabase="PVGIS-SARAH3"):
    """Récupère les PSH (E_d kWh/kWp/jour) mensuelles via PVGIS PVcalc.
    Retourne un dict: {"psh_by_month": {1: val, ..., 12: val}, "meta": {...}}
    """
    try:
        params = {
            "lat": float(lat),
            "lon": float(lon),
            "peakpower": 1,
            "loss": 14,
            "outputformat": "json",
            "raddatabase": raddatabase,
            "pvtechchoice": "crystSi"
        }
        if optimalangles:
            params["optimalangles"] = 1
        else:
            params["optimalangles"] = 0
            if angle is not None:
                params["angle"] = float(angle)
            if aspect is not None:
                params["aspect"] = float(aspect)
        resp = requests.get("https://re.jrc.ec.europa.eu/api/PVcalc", params=params, timeout=30)
        data = resp.json()
        monthly = data.get("outputs", {}).get("monthly", [])
        psh_by_month = {int(m.get("month")): float(m.get("E_d")) for m in monthly if m.get("E_d") is not None}
        meta = {
            "angle": data.get("inputs", {}).get("angle"),
            "aspect": data.get("inputs", {}).get("aspect"),
            "raddatabase": data.get("inputs", {}).get("raddatabase"),
            "location": data.get("inputs", {}).get("location")
        }
        return {"psh_by_month": psh_by_month, "meta": meta}
    except Exception as e:
        return {"psh_by_month": {}, "meta": {"error": str(e)}}

# Fonction de dimensionnement améliorée
def calculer_dimensionnement(consommation_journaliere, autonomie_jours=1, voltage=12, type_batterie="AGM", part_nuit=0.5):
    # Conversion du voltage pour les calculs
    if voltage == "High Voltage":
        # Pour High Voltage, utiliser une valeur représentative (400V par exemple)
        voltage_numeric = 400
    else:
        voltage_numeric = int(voltage)
    # Coefficients issus des secrets (avec valeurs par défaut en repli)
    try:
        panel_loss_factor = float(st.secrets["formulas"]["panel_loss_factor"])
    except Exception:
        panel_loss_factor = 1.25
    try:
        solar_hours = float(st.secrets["formulas"]["solar_hours"])
    except Exception:
        solar_hours = 5.0
    # Override des heures solaires si PVGIS a été utilisé
    try:
        override_sh = float(st.session_state.get("solar_hours_override", 0) or 0)
        if override_sh > 0:
            solar_hours = override_sh
    except Exception:
        pass
    try:
        inverter_peak_fraction = float(st.secrets["formulas"]["inverter_peak_fraction"])
    except Exception:
        inverter_peak_fraction = 1.0 / 3.0

    # Paramètres selon le type de batterie (priorité aux secrets)
    decharge_max = {}
    for k, default in [("Plomb", 0.5), ("AGM", 0.7), ("GEL", 0.8), ("Lithium", 0.9), ("Lithium HV", 0.95)]:
        try:
            decharge_max[k] = float(st.secrets["formulas"]["decharge_max"][k])
        except Exception:
            decharge_max[k] = default

    # Efficacité de cycle batterie (charge/décharge) selon la chimie
    efficacite_batterie_map = {}
    for k, default in [("Plomb", 0.85), ("AGM", 0.85), ("GEL", 0.85), ("Lithium", 0.93), ("Lithium HV", 0.96)]:
        try:
            efficacite_batterie_map[k] = float(st.secrets["formulas"]["battery_efficiency"][k])
        except Exception:
            efficacite_batterie_map[k] = default

    # Calcul de la puissance panneau nécessaire
    # Sortie en Watts-crête (Wc)
    puissance_panneaux = ((consommation_journaliere * panel_loss_factor) / max(solar_hours, 0.1)) * 1000

    # Hypothèse réaliste: charge le jour, décharge la nuit
    # On dimensionne la batterie sur la fraction nocturne de la consommation
    profondeur_decharge = decharge_max.get(type_batterie, 0.7)
    efficacite_batterie = efficacite_batterie_map.get(type_batterie, 0.85 if type_batterie in ("Plomb", "AGM", "GEL") else 0.93)
    consommation_nocturne = consommation_journaliere * max(0.1, min(part_nuit/100.0, 1.0))
    capacite_batterie = (consommation_nocturne * autonomie_jours * 1000) / (voltage_numeric * max(profondeur_decharge, 0.01) * max(efficacite_batterie, 0.01))

    # Puissance onduleur (fraction de la conso journalière)
    puissance_onduleur = consommation_journaliere * inverter_peak_fraction * 1000  # en W

    return {
        "puissance_panneaux": puissance_panneaux,
        "capacite_batterie": capacite_batterie,
        "puissance_onduleur": puissance_onduleur,
        "type_batterie": type_batterie,
        "profondeur_decharge": profondeur_decharge * 100,
        "efficacite_cycle": efficacite_batterie * 100
    }

# Fonction pour sélectionner les équipements
def selectionner_equipements(dimensionnement, choix_utilisateur):
    # Obtenir les prix actuels (Firebase ou par défaut)
    prix_equipements = get_current_prices()
    
    # Affichage du stock disponible pour les équipements
    if st.session_state.get('user_role') == 'admin':
        with st.expander("📦 Stock disponible pour les équipements", expanded=False):
            col_stock1, col_stock2, col_stock3 = st.columns(3)
            
            with col_stock1:
                st.markdown("**🌞 Panneaux Solaires**")
                for nom, specs in prix_equipements["panneaux"].items():
                    stock_info = get_stock_for_dimensioning_product("panneau", nom)
                    if stock_info:
                        stock_qty = stock_info.get('quantite', 0)
                        if stock_qty > 0:
                            st.success(f"✅ {nom}: {stock_qty} unités")
                        elif stock_qty == 0:
                            st.warning(f"⚠️ {nom}: Rupture de stock")
                        else:
                            st.info(f"ℹ️ {nom}: Stock non renseigné")
                    else:
                        st.info(f"ℹ️ {nom}: Non synchronisé")
            
            with col_stock2:
                st.markdown("**🔋 Batteries**")
                for nom, specs in prix_equipements["batteries"].items():
                    stock_info = get_stock_for_dimensioning_product("batterie", nom)
                    if stock_info:
                        stock_qty = stock_info.get('quantite', 0)
                        if stock_qty > 0:
                            st.success(f"✅ {nom}: {stock_qty} unités")
                        elif stock_qty == 0:
                            st.warning(f"⚠️ {nom}: Rupture de stock")
                        else:
                            st.info(f"ℹ️ {nom}: Stock non renseigné")
                    else:
                        st.info(f"ℹ️ {nom}: Non synchronisé")
            
            with col_stock3:
                st.markdown("**⚡ Onduleurs**")
                for nom, specs in prix_equipements["onduleurs"].items():
                    stock_info = get_stock_for_dimensioning_product("onduleur", nom)
                    if stock_info:
                        stock_qty = stock_info.get('quantite', 0)
                        if stock_qty > 0:
                            st.success(f"✅ {nom}: {stock_qty} unités")
                        elif stock_qty == 0:
                            st.warning(f"⚠️ {nom}: Rupture de stock")
                        else:
                            st.info(f"ℹ️ {nom}: Stock non renseigné")
                    else:
                        st.info(f"ℹ️ {nom}: Non synchronisé")
    
    type_batterie = choix_utilisateur["type_batterie"]
    type_onduleur = choix_utilisateur["type_onduleur"]
    # Supporte l'absence de type_regulateur (ex: onduleur Hybride)
    type_regulateur = choix_utilisateur.get("type_regulateur", "MPPT")
    voltage_systeme = choix_utilisateur["voltage"]
    # Conversion du voltage pour les calculs numériques
    if voltage_systeme == "High Voltage":
        voltage_systeme_numeric = 400  # Valeur représentative pour High Voltage
    else:
        voltage_systeme_numeric = int(voltage_systeme)
    phase_type = choix_utilisateur.get("phase_type", "monophase")
    
    # Sélection panneaux — choisir le module qui minimise le nombre de panneaux
    puissance_panneau_select = None
    nb_panneaux = 0
    puissance_min = dimensionnement["puissance_panneaux"]

    candidats = []
    for nom, specs in prix_equipements["panneaux"].items():
        p = specs["puissance"]
        if p <= 0:
            continue
        nb = math.ceil(puissance_min / p)
        prix_unitaire = specs.get("prix", 0)
        prix_par_watt = prix_unitaire / p if p else float("inf")
        candidats.append((nom, nb, prix_par_watt, -p))  # tie-break par prix/W puis par puissance plus élevée

    if candidats:
        # Trier: moins de panneaux, meilleur prix/W, puissance plus élevée
        candidats.sort(key=lambda x: (x[1], x[2], x[3]))
        puissance_panneau_select, nb_panneaux, _, _ = candidats[0]

    
    # Sélection batterie selon le type choisi
    batterie_select = None
    nb_batteries = 0
    
    # Filtrage des batteries selon le voltage système
    if voltage_systeme == "High Voltage":
        # Pour High Voltage, prendre les batteries Lithium HV avec voltage > 48V
        batteries_filtrees = {k: v for k, v in prix_equipements["batteries"].items() 
                             if v["type"] == "Lithium HV" and v["voltage"] > 48}
        
        # Pour les batteries HV, utiliser les kWh pour la comparaison
        if batteries_filtrees:
            # Convertir la capacité requise en kWh pour les batteries HV
            # voltage_systeme_numeric est défini plus haut dans la fonction
            capacite_requise_kwh = (dimensionnement["capacite_batterie"] * voltage_systeme_numeric) / 1000.0
            
            for nom, specs in sorted(batteries_filtrees.items(), key=lambda x: x[1].get("kwh", 0)):
                if specs.get("kwh", 0) >= capacite_requise_kwh:
                    batterie_select = nom
                    nb_batteries = 1
                    break
            
            # Si aucune batterie assez grande, prendre plusieurs petites
            if not batterie_select:
                nom_batterie = max(batteries_filtrees.keys(), key=lambda x: batteries_filtrees[x].get("kwh", 0))
                specs = batteries_filtrees[nom_batterie]
                kwh_unitaire = specs.get("kwh", 1)
                nb_batteries = int(capacite_requise_kwh / kwh_unitaire) + 1
                batterie_select = nom_batterie
    else:
        # Pour les voltages standards, filtrage exact avec logique Ah
        batteries_filtrees = {k: v for k, v in prix_equipements["batteries"].items() 
                             if v["type"] == type_batterie and v["voltage"] == voltage_systeme}
        
        if batteries_filtrees:
            for nom, specs in sorted(batteries_filtrees.items(), key=lambda x: x[1]["capacite"]):
                if specs["capacite"] >= dimensionnement["capacite_batterie"]:
                    batterie_select = nom
                    nb_batteries = 1
                    break
            
            # Si aucune batterie assez grande, prendre plusieurs petites
            if not batterie_select:
                nom_batterie = max(batteries_filtrees.keys(), key=lambda x: batteries_filtrees[x]["capacite"])
                specs = batteries_filtrees[nom_batterie]
                nb_batteries = int(dimensionnement["capacite_batterie"] / specs["capacite"]) + 1
                batterie_select = nom_batterie
    
    # Sélection onduleur selon le type choisi avec couplage si nécessaire
    onduleur_select = None
    nb_onduleurs = 1
    
    # Filtrage des onduleurs selon le voltage système
    if voltage_systeme == "High Voltage":
        # Pour High Voltage, prendre les onduleurs avec voltage >= 180V
        onduleurs_filtres = {k: v for k, v in prix_equipements["onduleurs"].items() 
                            if type_onduleur == v["type"] and v["voltage"] >= 180 and v.get("phase", "monophase") == phase_type}
    else:
        # Pour les voltages standards, filtrage exact
        onduleurs_filtres = {k: v for k, v in prix_equipements["onduleurs"].items() 
                            if type_onduleur == v["type"] and v["voltage"] == voltage_systeme and v.get("phase", "monophase") == phase_type}
    
    if onduleurs_filtres:
        # Essayer d'abord un seul onduleur du type choisi
        for nom, specs in sorted(onduleurs_filtres.items(), key=lambda x: x[1]["puissance"]):
            if specs["puissance"] >= dimensionnement["puissance_onduleur"]:
                onduleur_select = nom
                break
        
        # Si aucun onduleur unique du type choisi ne suffit, chercher dans d'autres types compatibles
        if not onduleur_select:
            # Définir les types compatibles selon le type choisi
            types_compatibles = []
            if type_onduleur == "Hybride":
                types_compatibles = ["Online", "Online Tri"]  # Hybride peut être remplacé par Online
            elif type_onduleur == "Off-Grid":
                types_compatibles = ["Hybride", "Online", "Online Tri"]  # Off-Grid peut être remplacé par tout
            elif type_onduleur == "Online":
                types_compatibles = ["Online Tri"]  # Online peut être remplacé par Online Tri
            elif type_onduleur == "Online Tri":
                types_compatibles = []  # Online Tri est le plus haut niveau
            
            # Chercher dans les types compatibles
            for type_compatible in types_compatibles:
                if voltage_systeme == "High Voltage":
                    onduleurs_compatibles = {k: v for k, v in prix_equipements["onduleurs"].items() 
                                           if type_compatible == v["type"] and v["voltage"] >= 180 and v.get("phase", "monophase") == phase_type}
                else:
                    onduleurs_compatibles = {k: v for k, v in prix_equipements["onduleurs"].items() 
                                           if type_compatible == v["type"] and v["voltage"] == voltage_systeme and v.get("phase", "monophase") == phase_type}
                
                for nom, specs in sorted(onduleurs_compatibles.items(), key=lambda x: x[1]["puissance"]):
                    if specs["puissance"] >= dimensionnement["puissance_onduleur"]:
                        onduleur_select = nom
                        break
                
                if onduleur_select:
                    break
        
        # Si toujours aucun onduleur unique ne suffit, essayer le couplage avec le type choisi
        if not onduleur_select:
            # Prendre l'onduleur le plus puissant disponible du type choisi (avec puissance > 0)
            onduleurs_valides = {k: v for k, v in onduleurs_filtres.items() if v["puissance"] > 0}
            
            if onduleurs_valides:
                onduleur_max = max(onduleurs_valides.items(), key=lambda x: x[1]["puissance"])
                nom_max, specs_max = onduleur_max
                
                # Calculer le nombre d'onduleurs nécessaires
                nb_onduleurs = int(dimensionnement["puissance_onduleur"] / specs_max["puissance"]) + 1
                
                # Limiter à 4 onduleurs maximum pour des raisons pratiques
                if nb_onduleurs <= 4:
                    onduleur_select = nom_max
    
    # Sélection régulateur (seulement si onduleur pas hybride)
    regulateur_select = None
    if type_onduleur != "Hybride" and puissance_panneau_select and batterie_select:
        puissance_panneaux_totale = nb_panneaux * prix_equipements["panneaux"][puissance_panneau_select]["puissance"]
        amperage_requis = (puissance_panneaux_totale / voltage_systeme_numeric) * 1.25
        
        regulateurs_filtres = {k: v for k, v in prix_equipements["regulateurs"].items() 
                              if v["type"] == type_regulateur}
        
        for nom, specs in sorted(regulateurs_filtres.items(), key=lambda x: x[1]["amperage"]):
            if specs["amperage"] >= amperage_requis:
                regulateur_select = nom
                break
    
    return {
        "panneau": (puissance_panneau_select, nb_panneaux),
        "batterie": (batterie_select, nb_batteries),
        "onduleur": (onduleur_select, nb_onduleurs),
        "regulateur": regulateur_select,
    }

# Estimation kWh mensuels à partir d'une facture Senelec
# Note: approximation des paliers, hors frais fixes/abonnement/taxes.
def estimer_kwh_depuis_facture(montant_fcfa: float, type_compteur: str = "mensuel") -> float:
    try:
        m = float(montant_fcfa)
    except Exception:
        return 0.0
    if m <= 0:
        return 0.0

    # Tarifs et paliers issus des secrets (avec valeurs par défaut)
    try:
        t1 = float(st.secrets["formulas"]["facture"]["tarifs"]["tier1_price"])
    except Exception:
        t1 = 124.17
    try:
        t2 = float(st.secrets["formulas"]["facture"]["tarifs"]["tier2_price"])
    except Exception:
        t2 = 136.49
    try:
        t3 = float(st.secrets["formulas"]["facture"]["tarifs"]["tier3_price"])
    except Exception:
        t3 = 159.36

    if type_compteur.lower().startswith("bimes"):
        try:
            p1_kwh = float(st.secrets["formulas"]["facture"]["bimestriel"]["tier1_kwh"])
        except Exception:
            p1_kwh = 300.0
        try:
            p2_kwh = float(st.secrets["formulas"]["facture"]["bimestriel"]["tier2_kwh"])
        except Exception:
            p2_kwh = 200.0
    else:
        try:
            p1_kwh = float(st.secrets["formulas"]["facture"]["mensuel"]["tier1_kwh"])
        except Exception:
            p1_kwh = 150.0
        try:
            p2_kwh = float(st.secrets["formulas"]["facture"]["mensuel"]["tier2_kwh"])
        except Exception:
            p2_kwh = 100.0

    cout_p1 = p1_kwh * t1
    cout_p2 = p2_kwh * t2

    if m <= cout_p1:
        return m / t1
    elif m <= cout_p1 + cout_p2:
        return p1_kwh + (m - cout_p1) / t2
    else:
        return p1_kwh + p2_kwh + (m - cout_p1 - cout_p2) / t3

# Fonction pour calculer le devis
def calculer_devis(equipements, use_online=False, accessoires_rate=0.15, region_selectionnee=None):
    # Obtenir les prix actuels (Firebase ou par défaut)
    prix_equipements = get_current_prices()
    
    total = 0
    details = []
    
    # Panneaux
    panneau_nom, nb_panneaux = equipements["panneau"]
    if panneau_nom:
        prix_unitaire = prix_equipements["panneaux"][panneau_nom]["prix"]
        source_prix = "local"
        url_source = None
        if use_online:
            prix_site, url_site = obtenir_prix_depuis_site(panneau_nom)
            if prix_site:
                prix_unitaire = prix_site
                source_prix = "site"
                url_source = url_site
        sous_total = prix_unitaire * nb_panneaux
        total += sous_total
        details.append({
            "item": f"Panneau solaire {panneau_nom}",
            "quantite": nb_panneaux,
            "prix_unitaire": prix_unitaire,
            "sous_total": sous_total,
            "source_prix": source_prix,
            "url_source": url_source
        })
        
        # Supports de panneaux (forfait par panneau)
        if panneau_nom and nb_panneaux > 0:
            prix_support = 25000
            sous_total_supports = prix_support * nb_panneaux
            total += sous_total_supports
            details.append({
                "item": "Supports de panneaux",
                "quantite": nb_panneaux,
                "prix_unitaire": prix_support,
                "sous_total": sous_total_supports,
                "source_prix": "forfait 25 000/panneau",
                "url_source": None
            })
    
    # Batteries
    batterie_nom, nb_batteries = equipements["batterie"]
    if batterie_nom:
        prix_unitaire = prix_equipements["batteries"][batterie_nom]["prix"]
        source_prix = "local"
        url_source = None
        if use_online:
            prix_site, url_site = obtenir_prix_depuis_site(batterie_nom)
            if prix_site:
                prix_unitaire = prix_site
                source_prix = "site"
                url_source = url_site
        sous_total = prix_unitaire * nb_batteries
        total += sous_total
        details.append({
            "item": f"Batterie {batterie_nom}",
            "quantite": nb_batteries,
            "prix_unitaire": prix_unitaire,
            "sous_total": sous_total,
            "source_prix": source_prix,
            "url_source": url_source
        })
    
    # Onduleur(s)
    onduleur_data = equipements["onduleur"]
    if onduleur_data:
        # Gérer le nouveau format avec couplage
        if isinstance(onduleur_data, tuple):
            onduleur_nom, nb_onduleurs = onduleur_data
        else:
            # Compatibilité avec l'ancien format
            onduleur_nom = onduleur_data
            nb_onduleurs = 1
            
        if onduleur_nom:
            prix_unitaire = prix_equipements["onduleurs"][onduleur_nom]["prix"]
            source_prix = "local"
            url_source = None
            if use_online:
                prix_site, url_site = obtenir_prix_depuis_site(onduleur_nom)
                if prix_site:
                    prix_unitaire = prix_site
                    source_prix = "site"
                    url_source = url_site
            sous_total = prix_unitaire * nb_onduleurs
            total += sous_total
            
            # Affichage adapté selon le nombre d'onduleurs
            if nb_onduleurs > 1:
                item_name = f"Onduleur {onduleur_nom} (couplage de {nb_onduleurs})"
            else:
                item_name = f"Onduleur {onduleur_nom}"
                
            details.append({
                "item": item_name,
                "quantite": nb_onduleurs,
                "prix_unitaire": prix_unitaire,
                "sous_total": sous_total,
                "source_prix": source_prix,
                "url_source": url_source
            })
    
    # Régulateur (si nécessaire)
    regulateur_nom = equipements["regulateur"]
    if regulateur_nom:
        prix_unitaire = prix_equipements["regulateurs"][regulateur_nom]["prix"]
        source_prix = "local"
        url_source = None
        if use_online:
            prix_site, url_site = obtenir_prix_depuis_site(regulateur_nom)
            if prix_site:
                prix_unitaire = prix_site
                source_prix = "site"
                url_source = url_site
        total += prix_unitaire
        details.append({
            "item": f"Régulateur {regulateur_nom}",
            "quantite": 1,
            "prix_unitaire": prix_unitaire,
            "sous_total": prix_unitaire,
            "source_prix": source_prix,
            "url_source": url_source
        })
    
    # Accessoires (câbles, connecteurs, protections)
    accessoires = int(total * accessoires_rate)
    total += accessoires
    details.append({
        "item": "Accessoires (câbles, connecteurs, protections)",
        "quantite": 1,
        "prix_unitaire": accessoires,
        "sous_total": accessoires,
        "source_prix": f"taux {int(accessoires_rate*100)}%",
        "url_source": None
    })
    
    # Calcul de la puissance totale
    puissance_totale = 0
    if panneau_nom:
        puissance_totale = nb_panneaux * prix_equipements["panneaux"][panneau_nom]["puissance"] / 1000
    
    # Calcul de l'installation et mise en service selon la région
    if region_selectionnee:
        # Récupérer les pourcentages depuis Firebase ou utiliser les valeurs par défaut
        pourcentages_firebase = get_labor_percentages()
        if pourcentages_firebase and region_selectionnee in pourcentages_firebase:
            pourcentage_main_oeuvre = pourcentages_firebase[region_selectionnee]
        elif region_selectionnee in POURCENTAGES_MAIN_OEUVRE_DEFAUT:
            pourcentage_main_oeuvre = POURCENTAGES_MAIN_OEUVRE_DEFAUT[region_selectionnee]
        else:
            pourcentage_main_oeuvre = 20.0  # Valeur par défaut si région non trouvée
        
        # Calculer l'installation et mise en service en pourcentage du coût des équipements
        cout_installation = round(total * (pourcentage_main_oeuvre / 100.0))
        
        total += cout_installation
        details.append({
            "item": f"Installation et mise en service - {region_selectionnee}",
            "quantite": 1,
            "prix_unitaire": cout_installation,
            "sous_total": cout_installation,
            "source_prix": "pourcentage régional",
            "url_source": None
        })
    else:
        # Si aucune région n'est sélectionnée, utiliser un forfait par défaut
        installation_defaut = 200000
        total += installation_defaut
        details.append({
            "item": "Installation et mise en service",
            "quantite": 1,
            "prix_unitaire": installation_defaut,
            "sous_total": installation_defaut,
            "source_prix": "forfait",
            "url_source": None
        })
    
    return {"details": details, "total": total, "puissance_totale": puissance_totale}

# Interface principale
st.title("☀️ Dimensionnement d'Installation Solaire - Sénégal")
st.markdown("### Calculez votre installation solaire complète et obtenez un devis estimatif détaillé")

# Barre latérale
with st.sidebar:
    # Logo dans la sidebar - centré
    try:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image("logo-solaire.svg", width=350)
    except:
        st.markdown("### ☀️ Energie Solaire Sénégal")
    st.markdown("### ☀️ Parlez avec Pape votre Conseiller solaire")
    
    # Callback: déclenché à l'appui sur Entrée
    def _trigger_sidebar_chat():
        st.session_state.sidebar_chat_go = True
    
    q_sidebar = st.text_input(
        "Votre question à Pape",
        placeholder="Ex: Décrivez vos appareils ou votre besoin",
        key="sidebar_chat_q",
        on_change=_trigger_sidebar_chat
    )
    
    # Soumission automatique sur Entrée
    if st.session_state.get("sidebar_chat_go"):
        if q_sidebar and len(q_sidebar.strip()) > 5:
            # Construit un contexte synthétique depuis l'état courant
            contexte_sb = ""
            if 'dimensionnement' in st.session_state:
                dim = st.session_state.dimensionnement
                choix = st.session_state.choix
                equip_actifs_ctx = st.session_state.get('equip_choisi', st.session_state.get('equipements', None))
                prod_kwh_j_ctx = 0.0
                auto_reelle_ctx = (st.session_state.autonomie_reelle_pct if 'autonomie_reelle_pct' in st.session_state else None)
                if not auto_reelle_ctx and equip_actifs_ctx and equip_actifs_ctx.get('panneau'):
                    pn_ctx, nb_ctx = equip_actifs_ctx['panneau']
                    if pn_ctx and nb_ctx > 0:
                        p_unit_ctx = PRIX_EQUIPEMENTS['panneaux'][pn_ctx]['puissance']
                        p_tot_ctx = p_unit_ctx * nb_ctx
                        prod_kwh_j_ctx = (p_tot_ctx / 1000.0) * 5.0 * 0.75
                conso_totale_ctx = st.session_state.consommation if 'consommation' in st.session_state else 10.0
                conso_couverte_ctx = st.session_state.get('consommation_couverte', None)
                autonomie_voulue_ctx = choix.get('autonomie_h', None)
                pack_info = f"{choix.get('type_batterie','?')} / {choix.get('type_onduleur','?')} / {('MPPT' if choix.get('type_regulateur')=='MPPT' else 'PWM' if choix.get('type_regulateur')=='PWM' else 'auto')} @ {choix.get('voltage', 12)}V"
                contexte_sb = f"Conso quotidienne: {conso_totale_ctx} kWh/j ; Couverture cible: {conso_couverte_ctx or 'N/A'} kWh/j ; Prod estimée: {round(prod_kwh_j_ctx,2)} kWh/j ; Autonomie cible: {autonomie_voulue_ctx or 'N/A'} h ; Pack choisi: {pack_info}"
            with st.spinner("🤔 Pape répond en streaming (réponse courte)..."):
                st.write_stream(appeler_assistant_ia_stream(q_sidebar, contexte_sb, concis=True, max_tokens=220, limite_caracteres=700))
            st.session_state.sidebar_chat_go = False
            st.caption("Réponse abrégée de Pape. Pour plus de détails, utilisez l'onglet Conseiller solaire.")
        else:
            st.session_state.sidebar_chat_go = False
            st.warning("⚠️ Veuillez entrer une question (minimum 6 caractères)")
    
    st.markdown("---")
    st.markdown("### 🏢 Energie Solaire Sénégal")
    st.markdown("""
 
    
    🥇 **Premier outil de dimensionnement solaire en ligne au Sénégal**
    
    📍 **Adresse :** Castor 221 Dakar, Sénégal  
    (En face du terrain de Football)  
    Zac Mbao (pres du rond point AXA)

    📧 **Email :** energiesolairesenegal@gmail.com

    📞 **Téléphones :**  
    • +221 77 631 42 25  
    • +221 78 177 39 26  
    • +221 77 250 47 90
    
    🌐 **Site web :** [energiesolairesenegal.com](https://energiesolairesenegal.com)
    """)
    
    st.markdown("---")
    st.markdown("### À propos")
    st.info("Application développée en partenariat avec Energie Solaire Sénégal pour le dimensionnement complet de votre installation solaire.")

# Interface d'authentification admin dans la sidebar
with st.sidebar:
    st.markdown("---")
    
    if not is_user_authenticated():
        with st.expander("🔐 Connexion Admin", expanded=False):
            # Connexion
            st.subheader("🔐 Connexion")
            with st.form("admin_login"):
                email = st.text_input("Email", placeholder="admin@energiesolairesenegal.com")
                password = st.text_input("Mot de passe", type="password")
                login_btn = st.form_submit_button("Se connecter")
                
                if login_btn and email and password:
                    with st.spinner("Connexion en cours..."):
                        user = login_user(email, password)
                        if user:
                            st.session_state['user_token'] = user['idToken']
                            st.session_state['user_email'] = email
                            st.session_state['is_admin'] = is_admin_email(email)
                            st.success("✅ Connexion réussie!")
                            st.rerun()
                        else:
                            st.error("❌ Échec de la connexion. Vérifiez vos identifiants.")
            

    else:
        if is_admin_user():
            st.success(f"👋 **Admin connecté**")
        else:
            st.info(f"👋 **Utilisateur connecté**")
        st.write(f"📧 {st.session_state.get('user_email', '')}")
        if st.button("🚪 Se déconnecter", use_container_width=True):
            logout_user()
            st.rerun()

# (Supprimé) Mode développement et Debug Info retirés selon demande

# Onglets principaux avec admin si connecté
if is_user_authenticated() and is_admin_user():
    tab1, tab2, tab3, tab_contact, tab_admin = st.tabs(["📊 Dimensionnement", "💰 Devis", "☀️ Conseiller solaire", "📞 Contact", "⚙️ Admin"])
else:
    tab1, tab2, tab3, tab_contact = st.tabs(["📊 Dimensionnement", "💰 Devis", "☀️ Conseiller Technique", "📞 Contact"])

with tab1:
    st.header("Calculez vos besoins en énergie solaire")
    
    # Charger les prix des équipements
    prix_equipements = get_current_prices()
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("1️⃣ Consommation")
        mode_calcul = st.radio("Méthode de calcul", ["Simple", "Détaillée"], horizontal=True)
        
        if mode_calcul == "Simple":
            consommation_simple = st.number_input(
                "Consommation électrique journalière (kWh/jour)",
                min_value=0.5,
                max_value=1000.0,
                step=0.5,
                help="Estimez votre consommation quotidienne moyenne",
                key="conso_journaliere_input"
            )
            with st.expander("🧮 Estimer depuis facture Senelec", expanded=False):
                type_compteur = st.selectbox(
                    "Type de compteur",
                    ["Normal (bimestriel)", "Woyofal (mensuel)"],
                    index=1,
                    key="type_compteur_select"
                )
                def _update_conso_from_montant():
                    montant_val = st.session_state.get("montant_fcfa_input", 0)
                    periodicite = "bimestriel" if "bimestriel" in st.session_state.get("type_compteur_select", "").lower() else "mensuel"
                    kwh_mensuel_estime = estimer_kwh_depuis_facture(montant_val, periodicite)
                    jours_cycle = 60 if periodicite == "bimestriel" else 30
                    conso_jour_estimee = kwh_mensuel_estime / jours_cycle if kwh_mensuel_estime > 0 else 0.0
                    # On écrit directement dans le champ principal
                    st.session_state.use_estimation = False
                    st.session_state.consommation_estimee = None
                    st.session_state["conso_journaliere_input"] = round(conso_jour_estimee, 2)
                montant_fcfa = st.number_input(
                    "Montant facture/achat (FCFA)",
                    min_value=0,
                    max_value=10000000,
                    value=st.session_state.get("montant_fcfa_input", 0),
                    step=1000,
                    help="Estimation approximative hors frais fixes/abonnements",
                    key="montant_fcfa_input",
                    on_change=_update_conso_from_montant
                )
                kwh_mensuel_estime = estimer_kwh_depuis_facture(
                    montant_fcfa,
                    "bimestriel" if "bimestriel" in type_compteur.lower() else "mensuel"
                )
                jours_cycle = 60 if "bimestriel" in type_compteur.lower() else 30
                conso_jour_estimee = kwh_mensuel_estime / jours_cycle if kwh_mensuel_estime > 0 else 0.0
                st.info(f"Estimation: {kwh_mensuel_estime:.0f} kWh/mois • {conso_jour_estimee:.2f} kWh/jour")
                col_est1, col_est2 = st.columns(2)
                with col_est1:
                    use_est = st.button("Utiliser cette estimation", key="use_estimation_btn")
                with col_est2:
                    reset_est = st.button("Revenir à la saisie manuelle", key="reset_estimation_btn")
                if use_est:
                    st.session_state.use_estimation = True
                    st.session_state.consommation_estimee = conso_jour_estimee
                if reset_est:
                    st.session_state.use_estimation = False
                    st.session_state.consommation_estimee = None
            if st.session_state.get("use_estimation", False) and st.session_state.get("consommation_estimee"):
                consommation_finale = float(st.session_state.consommation_estimee)
                st.success(f"✅ Estimation utilisée: {consommation_finale:.2f} kWh/jour")
            else:
                consommation_finale = consommation_simple
        else:
            with st.expander("📱 Calculer par appareils", expanded=True):
                st.markdown("Sélectionnez une famille d’appareils, choisissez un appareil et ajoutez-le.")
                
                # Initialisation de la liste sélectionnée
                if 'appareils_selectionnes' not in st.session_state:
                    st.session_state.appareils_selectionnes = []
                consommation_rapide = 0.0
                
                tab_cat, tab_rapide = st.tabs(["Catalogue par familles", "Liste rapide"])
                
                with tab_cat:
                    col_sel1, col_sel2 = st.columns(2)
                    famille = col_sel1.selectbox("Famille d'appareils", sorted(APPAREILS_FAMILLES.keys()))
                    appareils_list = APPAREILS_FAMILLES.get(famille, [])
                    noms_list = [a['nom'] for a in appareils_list]
                    appareil_nom = col_sel2.selectbox("Appareil", noms_list, index=0)
                    app_sel = next((a for a in appareils_list if a['nom'] == appareil_nom), None)
                    p_typ = app_sel['puissance'] if app_sel else 10
                    col_in1, col_in2, col_in3 = st.columns([1, 1, 1])
                    quant = col_in1.number_input("Nombre", min_value=0, max_value=50, value=1, step=1)
                    heures = col_in2.number_input("Heures/jour", min_value=0.0, max_value=24.0, value=6.0, step=0.5)
                    puissance_w = col_in3.number_input("Puissance (W)", min_value=1, max_value=5000, value=int(p_typ), step=10)
                    add_btn = st.button("➕ Ajouter cet appareil", use_container_width=True)
                    if add_btn:
                        if quant > 0 and heures > 0 and puissance_w > 0:
                            st.session_state.appareils_selectionnes.append({
                                "famille": famille,
                                "nom": appareil_nom,
                                "puissance": puissance_w,
                                "quantite": int(quant),
                                "heures": float(heures),
                                "kwh_j": (int(quant) * puissance_w * float(heures)) / 1000.0
                            })
                            st.success(f"Ajouté: {appareil_nom} • {quant} × {puissance_w}W • {heures} h/j")
                
                    # Ajouter via Conseiller solaire (sans expander)
                    show_ai = st.checkbox("Ajouter via Pape (mots-clés simples)", value=False, key="ai_show_checkbox")
                    if show_ai:
                        phrase = st.text_input("Décrivez vos appareils (ex: 2 tv, 1 frigo, routeur wifi)", key="ai_phrase_input")
                        if st.button("Proposer et ajouter", key="ai_proposer_btn"):
                            tokens = re.findall(r"\w+", phrase.lower())
                            last_num = 1
                            for tk in tokens:
                                if tk.isdigit():
                                    try:
                                        last_num = int(tk)
                                    except Exception:
                                        last_num = 1
                                    continue
                                if tk in APPAREILS_KEYWORDS:
                                    fam, nom = APPAREILS_KEYWORDS[tk]
                                    appareils_fam = APPAREILS_FAMILLES.get(fam, [])
                                    app_def = next((a for a in appareils_fam if a['nom'] == nom), None)
                                    p = app_def['puissance'] if app_def else 10
                                    st.session_state.appareils_selectionnes.append({
                                        "famille": fam,
                                        "nom": nom,
                                        "puissance": p,
                                        "quantite": last_num,
                                        "heures": 6.0,
                                        "kwh_j": (last_num * p * 6.0) / 1000.0
                                    })
                                    last_num = 1
                            st.success("Suggestions ajoutées. Ajustez les heures ou puissances au besoin.")
                        st.caption("Mots-clés: tv, frigo, ventilateur, ordi, laptop, routeur, pompe, micro, bouilloire")
                
                with tab_rapide:
                    st.markdown("*Liste rapide (optionnelle) pour quelques appareils courants*")
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_a:
                        st.markdown("**Appareil**")
                    with col_b:
                        st.markdown("**Nombre**")
                    with col_c:
                        st.markdown("**Heures/jour**")
                    
                    appareils_conso = {}
                    
                    st.markdown("🔦 **Éclairage**")
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_b:
                        nb_led = st.number_input("", 0, 50, 0, key="led", label_visibility="collapsed")
                    with col_c:
                        h_led = st.number_input("", 1, 24, 6, key="h_led", label_visibility="collapsed")
                    appareils_conso["LED 10W"] = nb_led * 10 * h_led / 1000
                    
                    st.markdown("💨 **Ventilation**")
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_b:
                        nb_vent = st.number_input("", 0, 20, 0, key="vent", label_visibility="collapsed")
                    with col_c:
                        h_vent = st.number_input("", 1, 24, 10, key="h_vent", label_visibility="collapsed")
                    appareils_conso["Ventilateur 75W"] = nb_vent * 75 * h_vent / 1000
                    
                    st.markdown("📺 **Électroménager**")
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_b:
                        nb_tv = st.number_input("TV", 0, 10, 0, key="tv", label_visibility="collapsed")
                    with col_c:
                        h_tv = st.number_input("", 1, 24, 6, key="h_tv", label_visibility="collapsed")
                    appareils_conso["TV 100W"] = nb_tv * 100 * h_tv / 1000
                    
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_b:
                        nb_frigo = st.number_input("Frigo", 0, 5, 0, key="frigo", label_visibility="collapsed")
                    with col_c:
                        h_frigo = st.number_input("", 8, 24, 12, key="h_frigo", label_visibility="collapsed")
                    appareils_conso["Réfrigérateur 150W"] = nb_frigo * 150 * h_frigo / 1000
                    
                    st.markdown("💻 **Informatique**")
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    with col_b:
                        nb_pc = st.number_input("Ordi", 0, 10, 0, key="pc", label_visibility="collapsed")
                    with col_c:
                        h_pc = st.number_input("", 1, 24, 8, key="h_pc", label_visibility="collapsed")
                    appareils_conso["Ordinateur 200W"] = nb_pc * 200 * h_pc / 1000
                    
                    consommation_rapide = sum(appareils_conso.values())
                
                # Tableau et total
                if st.session_state.appareils_selectionnes:
                    st.markdown("### Appareils sélectionnés")
                    for i, it in enumerate(st.session_state.appareils_selectionnes):
                        col_a, col_b, col_c, col_d, col_e, col_f = st.columns([3, 1, 1, 1, 1, 1])
                        col_a.write(f"{it['famille']} • {it['nom']}")
                        col_b.write(f"{it['quantite']}")
                        col_c.write(f"{it['puissance']} W")
                        col_d.write(f"{it['heures']} h/j")
                        col_e.write(f"{it['kwh_j']:.2f} kWh/j")
                        if col_f.button("🗑️", key=f"rm_{i}"):
                            st.session_state.appareils_selectionnes.pop(i)
                            st.rerun()
                    conso_familles = sum([x['kwh_j'] for x in st.session_state.appareils_selectionnes])
                else:
                    conso_familles = 0.0
                
                consommation_finale = conso_familles + consommation_rapide
                st.success(f"**Consommation totale: {consommation_finale:.2f} kWh/jour**")
        
        # Répartition jour/nuit et week-end
        with st.expander("🌙 Répartition jour/nuit et week-end", expanded=False):
            part_jour = st.slider("Part jour (%)", 0, 100, 40, step=1, help="Pour ménages: nuit souvent plus élevée")
            part_nuit = 100 - part_jour
            # Stocker dans session_state pour accès global
            st.session_state.part_jour = part_jour
            st.session_state.part_nuit = part_nuit
            facteur_weekend = st.slider("Facteur week-end (%)", 80, 150, 110, step=5, help="Ex: 110% = +10% de conso le week-end")
            # Moyenne jour/nuit en tenant compte du week-end (5 jours semaine + 2 jours weekend)
            conso_jour_semaine = consommation_finale * (part_jour/100) * 5
            conso_nuit_semaine = consommation_finale * (part_nuit/100) * 5
            conso_jour_weekend = consommation_finale * (part_jour/100) * (facteur_weekend/100) * 2
            conso_nuit_weekend = consommation_finale * (part_nuit/100) * (facteur_weekend/100) * 2
            conso_jour_moy = (conso_jour_semaine + conso_jour_weekend) / 7
            conso_nuit_moy = (conso_nuit_semaine + conso_nuit_weekend) / 7
            col_rep1, col_rep2, col_rep3 = st.columns(3)
            with col_rep1:
                st.metric("Jour moyen", f"{conso_jour_moy:.2f} kWh/jour")
            with col_rep2:
                st.metric("Nuit moyenne", f"{conso_nuit_moy:.2f} kWh/jour")
            with col_rep3:
                st.metric("Total", f"{(conso_jour_moy+conso_nuit_moy):.2f} kWh/jour")
            # Sauvegarde profil dans session
            st.session_state.usage_profile = {
                "part_jour": part_jour,
                "part_nuit": part_nuit,
                "facteur_weekend": facteur_weekend,
                "jour_moy": conso_jour_moy,
                "nuit_moy": conso_nuit_moy
            }
    
    with col2:
        st.subheader("2️⃣ Configuration du Système")
        
        # Type de batterie
        type_batterie = st.selectbox(
            "🔋 Type de batterie",
            ["Plomb", "AGM", "GEL", "Lithium", "Lithium HV"],
            index=3,
            help="AGM recommandé pour le climat sénégalais, Lithium HV pour installations haute puissance"
        )
        
        # Affichage des caractéristiques de la batterie choisie
        with st.expander(f"ℹ️ Pourquoi {type_batterie} ?"):
            info = INFO_BATTERIES[type_batterie]
            st.markdown(info["avantages"])
            st.markdown(info["inconvenients"])
            st.info(f"💡 **Recommandé pour:** {info['usage']}")
        
        # Type d'onduleur
        type_onduleur = st.selectbox(
            "⚡ Type d'onduleur",
            ["Off-Grid", "Hybride", "Online"],
            index=1,
            help="Hybride = avec régulateur MPPT intégré"
        )
        
        # Choix monophasé/triphasé
        phase_type = st.selectbox(
            "🔌 Type de phase",
            ["monophase", "triphase"],
            index=0,
            help="Monophasé pour usage domestique, Triphasé pour usage industriel"
        )
        
        # Type de régulateur (si nécessaire)
        type_regulateur = "MPPT"
        if type_onduleur != "Hybride":
            type_regulateur = st.selectbox(
                "🎛️ Type de régulateur",
                ["PWM", "MPPT"],
                index=1,
                help="MPPT 30% plus efficace que PWM"
            )
            if type_regulateur == "MPPT":
                st.success("✅ MPPT = 30% de rendement en plus")
            else:
                st.warning("⚠️ PWM = moins cher mais moins efficace")
        
        # Voltage du système
        voltage = st.selectbox(
            "⚡ Voltage du système",
            [12, 24, 48, "High Voltage"],
            index=2,
            help="48V recommandé pour usage domestique, High Voltage pour installations industrielles (>180V)"
        )
        
        # Niveau d'autonomie (pourcentage de besoins couverts)
        autonomie_pct = st.slider(
        "🔄 Niveau d’autonomie (%)",
        min_value=0,
        max_value=100,
        value=100,
        step=5,
        help="Objectif de couverture souhaitée par le solaire (cible)"
    )

    # Localisation et PSH PVGIS
    with st.expander("📍 Localisation et PSH PVGIS", expanded=False):
        villes = {
            "Dakar": (14.6937, -17.4441),
            "Thiès": (14.7900, -16.9240),
            "Saint-Louis": (16.0170, -16.4890),
            "Ziguinchor": (12.5560, -16.2720),
            "Tambacounda": (13.7700, -13.6670),
            "Kaolack": (14.1470, -16.0740),
            "Kolda": (12.8850, -14.9550),
            "Louga": (15.6180, -16.2240),
            "Matam": (15.6600, -13.3430),
            "Sédhiou": (12.7100, -15.5540),
            "Kaffrine": (14.1050, -15.5480),
            "Kédougou": (12.5540, -12.1740),
            "Autre (coordonnées)": None
        }
        ville = st.selectbox("Ville", list(villes.keys()), index=list(villes.keys()).index("Dakar"), key="pvgis_city")
        if ville == "Autre (coordonnées)":
            lat = st.number_input("Latitude", -90.0, 90.0, value=st.session_state.get("location_lat", 14.6937), step=0.001, key="location_lat")
            lon = st.number_input("Longitude", -180.0, 180.0, value=st.session_state.get("location_lon", -17.4441), step=0.001, key="location_lon")
        else:
            lat, lon = villes[ville]
            st.session_state["location_lat"] = lat
            st.session_state["location_lon"] = lon
            st.caption(f"Coordonnées: lat {lat:.4f}, lon {lon:.4f}")
        use_optimal = st.checkbox("Utiliser l’angle optimal PVGIS", value=True, key="pvgis_optimal")
        if use_optimal:
            st.caption("Optimalangles activé (inclinaison/azimut calculés par PVGIS)")
        else:
            angle = st.slider("Inclinaison (°)", 0, 60, value=20, key="pvgis_angle")
            aspect = st.slider("Azimut (°)", -90, 90, value=0, help="0=Sud, -90=Est, 90=Ouest", key="pvgis_aspect")
        mode_mois = st.selectbox("Saison / Mois dimensionnant", ["Mois le plus défavorable (PSH min)", "Choisir un mois"], index=0, key="pvgis_month_mode")
        if mode_mois == "Choisir un mois":
            mois_label = ["Janvier","Février","Mars","Avril","Mai","Juin","Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
            mois = st.selectbox("Mois", mois_label, index=11, key="pvgis_selected_month")
    
    # Bouton de calcul
    st.markdown("---")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        calculer_btn = st.button("🔍 CALCULER LE DIMENSIONNEMENT", type="primary", use_container_width=True)
    
    if calculer_btn:
        if consommation_finale > 0:
            with st.spinner("⚙️ Calcul en cours..."):
                # Calcul du dimensionnement avec niveau d’autonomie appliqué
                consommation_couverte = consommation_finale * autonomie_pct / 100.0

                # PVGIS: récupérer PSH mensuelles et choisir la saison
                lat = st.session_state.get("location_lat", 14.6937)
                lon = st.session_state.get("location_lon", -17.4441)
                optimal = st.session_state.get("pvgis_optimal", True)
                angle = st.session_state.get("pvgis_angle", None)
                aspect = st.session_state.get("pvgis_aspect", None)
                pvgis = get_pvgis_monthly_psh(lat, lon, optimalangles=optimal, angle=angle, aspect=aspect)
                monthly_psh = pvgis.get("psh_by_month", {})
                pvgis_mode = st.session_state.get("pvgis_month_mode", "Mois le plus défavorable (PSH min)")
                if monthly_psh:
                    if pvgis_mode.startswith("Choisir"):
                        mois_label = ["Janvier","Février","Mars","Avril","Mai","Juin","Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
                        selected = st.session_state.get("pvgis_selected_month", "Décembre")
                        mois_num = mois_label.index(selected) + 1 if selected in mois_label else 12
                        solar_hours_use = monthly_psh.get(mois_num, min(monthly_psh.values()))
                    else:
                        solar_hours_use = min(monthly_psh.values())
                    st.session_state["solar_hours_override"] = float(solar_hours_use)
                    st.session_state["pvgis_monthly_psh"] = monthly_psh
                    st.session_state["pvgis_meta"] = pvgis.get("meta", {})
                else:
                    st.session_state["solar_hours_override"] = None

                dim = calculer_dimensionnement(consommation_couverte, voltage=voltage, type_batterie=type_batterie, part_nuit=part_nuit)
                
                # Choix utilisateur
                choix_utilisateur = {
                    "type_batterie": type_batterie,
                    "type_onduleur": type_onduleur,
                    "type_regulateur": type_regulateur,
                    "voltage": voltage,
                    "phase_type": phase_type
                }
                
                # Sélection des équipements
                equip = selectionner_equipements(dim, choix_utilisateur)
                
                # Sauvegarde dans session
                st.session_state.dimensionnement = dim
                st.session_state.equipements = equip
                st.session_state.consommation = consommation_finale  # totale
                st.session_state.consommation_couverte = consommation_couverte
                st.session_state.autonomie_pct = autonomie_pct
                st.session_state.choix = choix_utilisateur
                
                st.success("✅ Dimensionnement effectué avec succès !")


                
            # Affichage des résultats
            st.markdown("---")
            
            # En-tête amélioré avec design attractif
            st.markdown("""
            <div style="background: linear-gradient(135deg, #4CAF50, #45a049); padding: 20px; border-radius: 15px; margin: 20px 0; box-shadow: 0 4px 15px rgba(76, 175, 80, 0.3);">
                <h2 style="color: white; margin: 0; text-align: center; font-size: 28px;">
                    🎯 Résultats du Dimensionnement
                </h2>
                <p style="color: #E8F5E8; text-align: center; margin: 10px 0 0 0; font-size: 16px;">
                    Votre installation solaire personnalisée pour le Sénégal
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # Résumé rapide du système
            autonomie_reelle = st.session_state.get('autonomie_reelle_pct', st.session_state.get('autonomie_pct', 100))
            production_kwh = st.session_state.get('production_solaire_kwh_j', st.session_state.get('consommation_couverte', st.session_state.consommation))
            
            col_info1, col_info2, col_info3 = st.columns(3)
            with col_info1:
                st.info(f"🏠 **Consommation:** {st.session_state.consommation:.1f} kWh/jour")
            with col_info2:
                st.info(f"⚡ **Production estimée:** {production_kwh:.1f} kWh/jour")
            with col_info3:
                autonomie_color = "🟢" if autonomie_reelle >= 90 else "🟡" if autonomie_reelle >= 70 else "🔴"
                st.info(f"{autonomie_color} **Autonomie:** {autonomie_reelle:.0f}%")
            
            # Métriques principales avec design amélioré
            st.markdown("### 🔧 Équipements Dimensionnés")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Panneaux solaires avec indicateurs visuels
                puissance_kw = dim['puissance_panneaux'] / 1000.0
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #FF9800, #F57C00); padding: 15px; border-radius: 10px; margin: 10px 0; color: white; text-align: center;">
                    <h3 style="margin: 0; font-size: 18px;">🌞 Panneaux Solaires</h3>
                    <h2 style="margin: 5px 0; font-size: 24px;">{dim['puissance_panneaux']:.0f} Wc</h2>
                    <p style="margin: 0; font-size: 14px; opacity: 0.9;">({puissance_kw:.1f} kWc)</p>
                </div>
                """, unsafe_allow_html=True)
                
                panneau_nom, nb = equip["panneau"]
                if panneau_nom:
                    st.success(f"✅ **{nb} x {panneau_nom}**")
                    surface_dim_m2 = puissance_kw * SURFACE_PAR_KWC_M2 * (1 + MARGE_IMPLANTATION_SURFACE_PCT / 100.0)
                    st.caption(f"📐 Surface nécessaire: ~{surface_dim_m2:.1f} m²")
                    
                    # Indicateur de qualité du panneau
                    if "Monocristallin" in panneau_nom:
                        st.caption("🏆 Technologie Monocristallin - Haut rendement")
                    else:
                        st.caption("💎 Technologie Polycristallin - Bon rapport qualité/prix")
            
            with col2:
                # Batteries avec indicateurs de performance
                # Convertir voltage en valeur numérique pour les calculs
                voltage_numeric = 400 if voltage == "High Voltage" else int(voltage)
                capacite_kwh = (dim['capacite_batterie'] * voltage_numeric) / 1000.0
                autonomie_jours = capacite_kwh / st.session_state.consommation
                
                # Affichage adapté selon le type de batterie
                if voltage == "High Voltage":
                    # Pour les batteries HV, afficher en kWh en priorité
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #2196F3, #1976D2); padding: 15px; border-radius: 10px; margin: 10px 0; color: white; text-align: center;">
                        <h3 style="margin: 0; font-size: 18px;">🔋 Batteries HV</h3>
                        <h2 style="margin: 5px 0; font-size: 24px;">{capacite_kwh:.1f} kWh</h2>
                        <p style="margin: 0; font-size: 14px; opacity: 0.9;">(≈{dim['capacite_batterie']:.0f} Ah équivalent à High Voltage)</p>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    # Pour les batteries standards, afficher en Ah en priorité
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #2196F3, #1976D2); padding: 15px; border-radius: 10px; margin: 10px 0; color: white; text-align: center;">
                        <h3 style="margin: 0; font-size: 18px;">🔋 Batteries</h3>
                        <h2 style="margin: 5px 0; font-size: 24px;">{dim['capacite_batterie']:.0f} Ah</h2>
                        <p style="margin: 0; font-size: 14px; opacity: 0.9;">({capacite_kwh:.1f} kWh à {voltage}V)</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                batterie_nom, nb = equip["batterie"]
                if batterie_nom:
                    st.success(f"✅ **{nb} x {batterie_nom}**")
                    st.caption(f"⏱️ Autonomie théorique: ~{autonomie_jours:.1f} jours")
                    st.caption(f"🔄 Décharge max: {dim['profondeur_decharge']:.0f}%")
                    
                    # Indicateur de qualité de la batterie
                    type_batterie = st.session_state.choix['type_batterie']
                    if type_batterie == "Lithium HV":
                        st.caption("🔥 Lithium HV - Durée de vie 15-20 ans, haute performance")
                    elif type_batterie == "Lithium":
                        st.caption("🚀 Lithium - Durée de vie 10-12 ans")
                    elif type_batterie == "GEL":
                        st.caption("⭐ GEL - Durée de vie 5-7 ans")
                    elif type_batterie == "AGM":
                        st.caption("👍 AGM - Bon compromis pour le Sénégal")
                    else:
                        st.caption("⚠️ Plomb - Entretien requis")
            
            with col3:
                # Onduleur avec indicateurs de capacité
                puissance_kw_ond = dim['puissance_onduleur'] / 1000.0
                marge_puissance = (dim['puissance_onduleur'] / (st.session_state.consommation * 1000 / 24)) * 100
                
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #9C27B0, #7B1FA2); padding: 15px; border-radius: 10px; margin: 10px 0; color: white; text-align: center;">
                    <h3 style="margin: 0; font-size: 18px;">⚡ Onduleur</h3>
                    <h2 style="margin: 5px 0; font-size: 24px;">{dim['puissance_onduleur']:.0f} W</h2>
                    <p style="margin: 0; font-size: 14px; opacity: 0.9;">({puissance_kw_ond:.1f} kW)</p>
                </div>
                """, unsafe_allow_html=True)
                
                onduleur_data = equip["onduleur"]
                if onduleur_data:
                    # Gérer le nouveau format avec couplage
                    if isinstance(onduleur_data, tuple):
                        onduleur_nom, nb_onduleurs = onduleur_data
                        if onduleur_nom:  # Vérifier que onduleur_nom n'est pas None
                            if nb_onduleurs > 1:
                                st.success(f"✅ **{nb_onduleurs} x {onduleur_nom}** (couplage)")
                                st.caption(f"🔗 Puissance totale: {nb_onduleurs * prix_equipements['onduleurs'][onduleur_nom]['puissance']}W")
                            else:
                                st.success(f"✅ **{onduleur_nom}**")
                        else:
                            st.warning("⚠️ Aucun onduleur approprié trouvé pour cette configuration")
                    else:
                        # Compatibilité avec l'ancien format
                        if onduleur_data:  # Vérifier que onduleur_data n'est pas None
                            st.success(f"✅ **{onduleur_data}**")
                        else:
                            st.warning("⚠️ Aucun onduleur approprié trouvé pour cette configuration")
                    
                    # Indicateur de marge de puissance
                    if marge_puissance >= 150:
                        st.caption("🟢 Excellente marge de puissance")
                    elif marge_puissance >= 120:
                        st.caption("🟡 Bonne marge de puissance")
                    else:
                        st.caption("🔴 Marge de puissance juste")
                    
                    # Type d'onduleur et phase
                    type_onduleur = st.session_state.choix['type_onduleur']
                    phase_type = st.session_state.choix.get('phase_type', 'monophase')
                    phase_display = "Monophasé" if phase_type == "monophase" else "Triphasé"
                    
                    if type_onduleur == "Hybride":
                        st.caption(f"🔄 Hybride - MPPT intégré - {phase_display}")
                    elif type_onduleur == "Online":
                        st.caption(f"🏆 Online - Qualité premium - {phase_display}")
                    else:
                        st.caption(f"⚡ Off-Grid - Solution basique - {phase_display}")
                
            # 📊 Indicateurs de performance du système
            st.markdown("---")
            st.markdown("""
            <div style="background: linear-gradient(135deg, #E3F2FD, #BBDEFB); padding: 20px; border-radius: 15px; margin: 20px 0; border-left: 5px solid #2196F3;">
                <h3 style="color: #1976D2; margin: 0 0 15px 0; display: flex; align-items: center;">
                    📊 Indicateurs de Performance & Efficacité
                </h3>
                <p style="color: #424242; margin: 0; font-size: 14px;">
                    Analyse détaillée des performances de votre installation solaire
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # Calculs des indicateurs
            production_annuelle_kwh = (st.session_state.production_solaire_kwh_j if 'production_solaire_kwh_j' in st.session_state else st.session_state.consommation) * 365
            consommation_annuelle_kwh = st.session_state.consommation * 365
            taux_autosuffisance = min(100, (production_annuelle_kwh / consommation_annuelle_kwh) * 100)
            
            # Efficacité énergétique
            kWc = dim['puissance_panneaux'] / 1000.0
            rendement_specifique = production_annuelle_kwh / kWc if kWc > 0 else 0
            
            # Facteur de charge
            facteur_charge = (production_annuelle_kwh / (kWc * 8760)) * 100 if kWc > 0 else 0
            
            # Impact environnemental
            co2_evite_kg = production_annuelle_kwh * 0.82  # 0.82 kg CO2/kWh évité au Sénégal
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Taux d'autosuffisance
                color_auto = "#4CAF50" if taux_autosuffisance >= 80 else "#FF9800" if taux_autosuffisance >= 60 else "#F44336"
                st.markdown(f"""
                <div style="background: white; padding: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid {color_auto};">
                    <div style="font-size: 32px; font-weight: bold; color: {color_auto}; margin-bottom: 8px;">
                        {taux_autosuffisance:.1f}%
                    </div>
                    <div style="font-size: 16px; font-weight: 600; color: #333; margin-bottom: 5px;">
                        Taux d'Autosuffisance
                    </div>
                    <div style="font-size: 12px; color: #666;">
                        {"🟢 Excellent" if taux_autosuffisance >= 80 else "🟡 Bon" if taux_autosuffisance >= 60 else "🔴 À améliorer"}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col2:
                # Rendement spécifique
                color_rend = "#4CAF50" if rendement_specifique >= 1200 else "#FF9800" if rendement_specifique >= 1000 else "#F44336"
                st.markdown(f"""
                <div style="background: white; padding: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid {color_rend};">
                    <div style="font-size: 32px; font-weight: bold; color: {color_rend}; margin-bottom: 8px;">
                        {rendement_specifique:.0f}
                    </div>
                    <div style="font-size: 16px; font-weight: 600; color: #333; margin-bottom: 5px;">
                        Rendement Spécifique
                    </div>
                    <div style="font-size: 12px; color: #666;">
                        kWh/kWc/an - {"🟢 Optimal" if rendement_specifique >= 1200 else "🟡 Correct" if rendement_specifique >= 1000 else "🔴 Faible"}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col3:
                # Facteur de charge
                color_fc = "#4CAF50" if facteur_charge >= 15 else "#FF9800" if facteur_charge >= 12 else "#F44336"
                st.markdown(f"""
                <div style="background: white; padding: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid {color_fc};">
                    <div style="font-size: 32px; font-weight: bold; color: {color_fc}; margin-bottom: 8px;">
                        {facteur_charge:.1f}%
                    </div>
                    <div style="font-size: 16px; font-weight: 600; color: #333; margin-bottom: 5px;">
                        Facteur de Charge
                    </div>
                    <div style="font-size: 12px; color: #666;">
                        {"🟢 Très bon" if facteur_charge >= 15 else "🟡 Acceptable" if facteur_charge >= 12 else "🔴 Sous-optimal"}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Indicateurs supplémentaires en ligne
            st.markdown("<br>", unsafe_allow_html=True)
            col4, col5 = st.columns(2)
            
            with col4:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #E8F5E8, #C8E6C9); padding: 15px; border-radius: 10px; border-left: 4px solid #4CAF50;">
                    <div style="display: flex; align-items: center; margin-bottom: 8px;">
                        <span style="font-size: 20px; margin-right: 10px;">🌱</span>
                        <span style="font-weight: 600; color: #2E7D32;">Impact Environnemental</span>
                    </div>
                    <div style="color: #388E3C; font-size: 14px;">
                        <strong>{co2_evite_kg:.0f} kg CO₂</strong> évités par an<br>
                        Équivalent à <strong>{co2_evite_kg/22:.1f} arbres</strong> plantés
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col5:
                duree_vie_systeme = 20  # années
                production_totale_vie = production_annuelle_kwh * duree_vie_systeme
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #FFF3E0, #FFE0B2); padding: 15px; border-radius: 10px; border-left: 4px solid #FF9800;">
                    <div style="display: flex; align-items: center; margin-bottom: 8px;">
                        <span style="font-size: 20px; margin-right: 10px;">⚡</span>
                        <span style="font-weight: 600; color: #F57C00;">Production sur 20 ans</span>
                    </div>
                    <div style="color: #EF6C00; font-size: 14px;">
                        <strong>{production_totale_vie:,.0f} kWh</strong> au total<br>
                        Soit <strong>{production_totale_vie/1000:.1f} MWh</strong> d'énergie propre
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            # 📅 Simulateur de production mensuelle (Sénégal)
            st.markdown("---")
            st.markdown("""
            <div style="background: linear-gradient(135deg, #FFF8E1, #FFECB3); padding: 20px; border-radius: 15px; margin: 20px 0; border-left: 5px solid #FFC107;">
                <h3 style="color: #F57C00; margin: 0 0 15px 0; display: flex; align-items: center;">
                    📅 Simulateur de Production Mensuelle
                </h3>
                <p style="color: #424242; margin: 0; font-size: 14px;">
                    Prévisions détaillées basées sur l'ensoleillement sénégalais et les conditions climatiques locales
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            kWc = dim['puissance_panneaux'] / 1000.0
            
            # Données d'ensoleillement détaillées pour le Sénégal
            heures_par_jour = {
                'Jan': 6.2, 'Fév': 6.5, 'Mar': 6.7, 'Avr': 6.6, 'Mai': 6.5, 'Juin': 6.0,
                'Juil': 5.5, 'Août': 5.4, 'Sep': 5.8, 'Oct': 6.0, 'Nov': 6.2, 'Déc': 6.1
            }
            jours_mois = {'Jan':31,'Fév':28,'Mar':31,'Avr':30,'Mai':31,'Juin':30,'Juil':31,'Août':31,'Sep':30,'Oct':31,'Nov':30,'Déc':31}
            
            # Facteurs de performance selon les conditions sénégalaises
            PR = 0.80  # Performance Ratio moyen
            facteurs_saisonniers = {
                'Jan': 0.95, 'Fév': 0.95, 'Mar': 0.90, 'Avr': 0.85, 'Mai': 0.80, 'Juin': 0.75,  # Saison sèche à chaude
                'Juil': 0.70, 'Août': 0.70, 'Sep': 0.75, 'Oct': 0.85, 'Nov': 0.90, 'Déc': 0.95   # Saison des pluies et retour
            }

            data = []
            production_totale = 0
            for m in heures_par_jour:
                # Production de base
                prod_base = kWc * heures_par_jour[m] * PR * jours_mois[m]
                # Application du facteur saisonnier (température, humidité, poussière)
                prod_ajustee = prod_base * facteurs_saisonniers[m]
                production_totale += prod_ajustee
                
                # Calcul du taux de couverture mensuel
                conso_mensuelle = st.session_state.consommation * jours_mois[m]
                taux_couverture = min(100, (prod_ajustee / conso_mensuelle) * 100)
                
                data.append({
                    'Mois': m, 
                    'Production (kWh)': round(prod_ajustee, 1),
                    'Consommation (kWh)': round(conso_mensuelle, 1),
                    'Taux de couverture (%)': round(taux_couverture, 1),
                    'Ensoleillement (h/j)': heures_par_jour[m],
                    'Facteur saisonnier': facteurs_saisonniers[m]
                })

            df_prod = pd.DataFrame(data)
            
            # Affichage du graphique principal avec Plotly pour plus d'interactivité
            try:
                import plotly.graph_objects as go
                from plotly.subplots import make_subplots
                
                # Création du graphique combiné
                fig = make_subplots(
                    rows=2, cols=1,
                    subplot_titles=('Production vs Consommation Mensuelle', 'Taux de Couverture Mensuel'),
                    vertical_spacing=0.15,
                    row_heights=[0.7, 0.3]
                )
                
                # Graphique 1: Production vs Consommation
                fig.add_trace(
                    go.Bar(
                        x=df_prod['Mois'],
                        y=df_prod['Production (kWh)'],
                        name='Production solaire',
                        marker_color='#FFC107',
                        hovertemplate='<b>%{x}</b><br>Production: %{y:.1f} kWh<br>Ensoleillement: %{customdata:.1f}h/j<extra></extra>',
                        customdata=df_prod['Ensoleillement (h/j)']
                    ),
                    row=1, col=1
                )
                
                fig.add_trace(
                    go.Scatter(
                        x=df_prod['Mois'],
                        y=df_prod['Consommation (kWh)'],
                        mode='lines+markers',
                        name='Consommation',
                        line=dict(color='#F44336', width=3),
                        marker=dict(size=8),
                        hovertemplate='<b>%{x}</b><br>Consommation: %{y:.1f} kWh<extra></extra>'
                    ),
                    row=1, col=1
                )
                
                # Graphique 2: Taux de couverture
                colors = ['#4CAF50' if x >= 100 else '#FF9800' if x >= 80 else '#F44336' for x in df_prod['Taux de couverture (%)']]
                fig.add_trace(
                    go.Bar(
                        x=df_prod['Mois'],
                        y=df_prod['Taux de couverture (%)'],
                        name='Taux de couverture',
                        marker_color=colors,
                        hovertemplate='<b>%{x}</b><br>Couverture: %{y:.1f}%<br>Facteur saisonnier: %{customdata:.2f}<extra></extra>',
                        customdata=df_prod['Facteur saisonnier'],
                        showlegend=False
                    ),
                    row=2, col=1
                )
                
                # Ligne de référence à 100%
                fig.add_hline(y=100, line_dash="dash", line_color="green", opacity=0.7, row=2, col=1)
                
                # Mise en forme
                fig.update_layout(
                    height=600,
                    title_text="📊 Analyse Détaillée de la Production Solaire Mensuelle",
                    title_x=0.5,
                    showlegend=True,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    hovermode='x unified'
                )
                
                fig.update_xaxes(title_text="Mois", row=2, col=1)
                fig.update_yaxes(title_text="Énergie (kWh)", row=1, col=1)
                fig.update_yaxes(title_text="Couverture (%)", row=2, col=1)
                
                st.plotly_chart(fig, use_container_width=True)
                
            except ImportError:
                # Fallback vers le graphique Streamlit standard si Plotly n'est pas disponible
                st.bar_chart(df_prod.set_index('Mois')[['Production (kWh)', 'Consommation (kWh)']])
            
            # Tableau détaillé
            with st.expander("📋 Détails mensuels complets", expanded=False):
                st.dataframe(
                    df_prod.style.format({
                        'Production (kWh)': '{:.1f}',
                        'Consommation (kWh)': '{:.1f}',
                        'Taux de couverture (%)': '{:.1f}%',
                        'Ensoleillement (h/j)': '{:.1f}',
                        'Facteur saisonnier': '{:.2f}'
                    }),
                    use_container_width=True
                )
            
            # Résumé des performances annuelles
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown(f"""
                <div style="background: white; padding: 15px; border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid #4CAF50;">
                    <div style="font-size: 24px; font-weight: bold; color: #4CAF50; margin-bottom: 5px;">
                        {production_totale:,.0f} kWh
                    </div>
                    <div style="font-size: 14px; color: #666;">
                        Production annuelle totale
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col2:
                mois_optimal = df_prod.loc[df_prod['Production (kWh)'].idxmax(), 'Mois']
                prod_max = df_prod['Production (kWh)'].max()
                st.markdown(f"""
                <div style="background: white; padding: 15px; border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid #FF9800;">
                    <div style="font-size: 24px; font-weight: bold; color: #FF9800; margin-bottom: 5px;">
                        {mois_optimal}
                    </div>
                    <div style="font-size: 14px; color: #666;">
                        Meilleur mois ({prod_max:.0f} kWh)
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col3:
                mois_faible = df_prod.loc[df_prod['Production (kWh)'].idxmin(), 'Mois']
                prod_min = df_prod['Production (kWh)'].min()
                st.markdown(f"""
                <div style="background: white; padding: 15px; border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; border-left: 4px solid #F44336;">
                    <div style="font-size: 24px; font-weight: bold; color: #F44336; margin-bottom: 5px;">
                        {mois_faible}
                    </div>
                    <div style="font-size: 14px; color: #666;">
                        Mois le plus faible ({prod_min:.0f} kWh)
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Informations contextuelles
            st.markdown("""
            <div style="background: #F5F5F5; padding: 15px; border-radius: 10px; margin: 15px 0;">
                <h4 style="color: #333; margin: 0 0 10px 0;">🌍 Facteurs climatiques pris en compte :</h4>
                <ul style="color: #666; margin: 0; padding-left: 20px;">
                    <li><strong>Saison sèche (Nov-Mai)</strong> : Conditions optimales, facteur 0.80-0.95</li>
                    <li><strong>Saison des pluies (Juin-Oct)</strong> : Réduction due à l'humidité et aux nuages, facteur 0.70-0.85</li>
                    <li><strong>Température</strong> : Impact de la chaleur sur le rendement des panneaux</li>
                    <li><strong>Poussière harmattan</strong> : Réduction temporaire en décembre-février</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Régulateur si nécessaire
            if equip["regulateur"]:
                st.markdown("### 🎛️ Régulateur de charge")
                st.info(f"**{equip['regulateur']}**")
            
            # Avertissements et recommandations enrichies
            st.markdown("---")
            st.markdown("""
            <div style="background: linear-gradient(135deg, #E8F5E8, #C8E6C9); padding: 20px; border-radius: 15px; margin: 20px 0; border-left: 5px solid #4CAF50;">
                <h3 style="color: #2E7D32; margin: 0 0 15px 0; display: flex; align-items: center;">
                    💡 Recommandations Techniques & Pratiques
                </h3>
                <p style="color: #424242; margin: 0; font-size: 14px;">
                    Conseils d'experts pour optimiser votre installation solaire au Sénégal
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # Analyse des choix techniques
            col_rec1, col_rec2 = st.columns(2)
            
            with col_rec1:
                st.markdown("#### 🔋 Analyse de votre choix de batterie")
                if type_batterie == "Lithium HV":
                    st.markdown("""
                    <div style="background: #E3F2FD; padding: 15px; border-radius: 10px; border-left: 4px solid #2196F3;">
                        <h5 style="color: #1565C0; margin: 0 0 10px 0;">🔥 Choix Premium - Haute Performance !</h5>
                        <ul style="color: #1976D2; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Durée de vie :</strong> 15-20 ans (6x plus que plomb)</li>
                            <li><strong>Décharge :</strong> 95-98% utilisable</li>
                            <li><strong>Densité énergétique :</strong> Maximale</li>
                            <li><strong>BMS avancé :</strong> Protection intelligente</li>
                            <li><strong>Idéal pour :</strong> Systèmes 48V+ haute puissance</li>
                            <li><strong>ROI :</strong> Excellent sur très long terme</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                elif type_batterie == "Lithium":
                    st.markdown("""
                    <div style="background: #E8F5E8; padding: 15px; border-radius: 10px; border-left: 4px solid #4CAF50;">
                        <h5 style="color: #2E7D32; margin: 0 0 10px 0;">✅ Excellent choix !</h5>
                        <ul style="color: #388E3C; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Durée de vie :</strong> 10-12 ans (3x plus que plomb)</li>
                            <li><strong>Décharge :</strong> 90% utilisable</li>
                            <li><strong>Maintenance :</strong> Aucune</li>
                            <li><strong>Poids :</strong> 3x plus léger</li>
                            <li><strong>ROI :</strong> Rentable sur le long terme</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                elif type_batterie == "GEL":
                    st.markdown("""
                    <div style="background: #E8F5E8; padding: 15px; border-radius: 10px; border-left: 4px solid #4CAF50;">
                        <h5 style="color: #2E7D32; margin: 0 0 10px 0;">✅ Très bon choix pour le Sénégal</h5>
                        <ul style="color: #388E3C; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Durée de vie :</strong> 5-7 ans</li>
                            <li><strong>Décharge :</strong> 80% utilisable</li>
                            <li><strong>Résistance :</strong> Excellente à la chaleur</li>
                            <li><strong>Maintenance :</strong> Aucune</li>
                            <li><strong>Idéal pour :</strong> Climat tropical</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                elif type_batterie == "AGM":
                    st.markdown("""
                    <div style="background: #FFF3E0; padding: 15px; border-radius: 10px; border-left: 4px solid #FF9800;">
                        <h5 style="color: #F57C00; margin: 0 0 10px 0;">👍 Bon compromis qualité/prix</h5>
                        <ul style="color: #EF6C00; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Durée de vie :</strong> 3-5 ans</li>
                            <li><strong>Décharge :</strong> 70% utilisable</li>
                            <li><strong>Avantage :</strong> Charge rapide</li>
                            <li><strong>Maintenance :</strong> Aucune</li>
                            <li><strong>Recommandé pour :</strong> Budget moyen</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background: #FFEBEE; padding: 15px; border-radius: 10px; border-left: 4px solid #F44336;">
                        <h5 style="color: #C62828; margin: 0 0 10px 0;">⚠️ Attention : Maintenance requise</h5>
                        <ul style="color: #D32F2F; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Durée de vie :</strong> 2-3 ans seulement</li>
                            <li><strong>Décharge :</strong> 50% max (risque de dégâts)</li>
                            <li><strong>Entretien :</strong> Eau distillée tous les 3 mois</li>
                            <li><strong>Surveillance :</strong> Niveau d'électrolyte</li>
                            <li><strong>Coût total :</strong> Plus élevé à long terme</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
            
            with col_rec2:
                st.markdown("#### ⚡ Analyse de votre système de charge")
                if type_regulateur == "MPPT" or type_onduleur == "Hybride":
                    st.markdown("""
                    <div style="background: #E8F5E8; padding: 15px; border-radius: 10px; border-left: 4px solid #4CAF50;">
                        <h5 style="color: #2E7D32; margin: 0 0 10px 0;">✅ MPPT : Choix optimal</h5>
                        <ul style="color: #388E3C; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Rendement :</strong> +25-30% vs PWM</li>
                            <li><strong>Suivi :</strong> Point de puissance max</li>
                            <li><strong>Température :</strong> Compensation automatique</li>
                            <li><strong>Monitoring :</strong> Données en temps réel</li>
                            <li><strong>ROI :</strong> Amortissement rapide</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background: #FFF3E0; padding: 15px; border-radius: 10px; border-left: 4px solid #FF9800;">
                        <h5 style="color: #F57C00; margin: 0 0 10px 0;">💡 Conseil : Upgrade vers MPPT</h5>
                        <ul style="color: #EF6C00; margin: 0; padding-left: 20px; font-size: 14px;">
                            <li><strong>Gain potentiel :</strong> +30% de production</li>
                            <li><strong>Surtout efficace :</strong> Temps nuageux</li>
                            <li><strong>Température :</strong> Meilleure gestion chaleur</li>
                            <li><strong>Investissement :</strong> Rentable rapidement</li>
                            <li><strong>Évolutivité :</strong> Facilite extensions</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Recommandations spécifiques au Sénégal
            st.markdown("#### 🌍 Recommandations spécifiques au climat sénégalais")
            
            col_climat1, col_climat2 = st.columns(2)
            
            with col_climat1:
                st.markdown("""
                <div style="background: #E3F2FD; padding: 15px; border-radius: 10px; border-left: 4px solid #2196F3;">
                    <h5 style="color: #1976D2; margin: 0 0 10px 0;">🌡️ Gestion de la température</h5>
                    <ul style="color: #1565C0; margin: 0; padding-left: 20px; font-size: 14px;">
                        <li><strong>Ventilation :</strong> Espace 15cm sous panneaux</li>
                        <li><strong>Orientation :</strong> Éviter exposition directe batteries</li>
                        <li><strong>Ombrage :</strong> Local technique ventilé</li>
                        <li><strong>Câblage :</strong> Section adaptée (pertes thermiques)</li>
                        <li><strong>Monitoring :</strong> Surveillance température batteries</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                
            with col_climat2:
                st.markdown("""
                <div style="background: #F3E5F5; padding: 15px; border-radius: 10px; border-left: 4px solid #9C27B0;">
                    <h5 style="color: #7B1FA2; margin: 0 0 10px 0;">💧 Protection contre l'humidité</h5>
                    <ul style="color: #6A1B9A; margin: 0; padding-left: 20px; font-size: 14px;">
                        <li><strong>Étanchéité :</strong> IP65 minimum pour équipements</li>
                        <li><strong>Drainage :</strong> Évacuation eau de pluie</li>
                        <li><strong>Corrosion :</strong> Fixations inox ou galvanisées</li>
                        <li><strong>Câbles :</strong> Gaines étanches et UV-résistantes</li>
                        <li><strong>Maintenance :</strong> Inspection post-hivernage</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            
            # Conseils de maintenance et optimisation
            st.markdown("#### 🔧 Plan de maintenance recommandé")
            
            maintenance_tabs = st.tabs(["📅 Mensuel", "🔄 Trimestriel", "📋 Annuel"])
            
            with maintenance_tabs[0]:
                st.markdown("""
                **🗓️ Tâches mensuelles (15 min)**
                - ✅ Nettoyage panneaux (eau + brosse douce)
                - ✅ Vérification niveau batteries (si plomb)
                - ✅ Contrôle visuel câblage
                - ✅ Relevé production (monitoring)
                - ✅ Test fonctionnement onduleur
                """)
                
            with maintenance_tabs[1]:
                st.markdown("""
                **🔄 Tâches trimestrielles (30 min)**
                - 🔧 Serrage connexions électriques
                - 🔧 Nettoyage bornes batteries
                - 🔧 Vérification fixations panneaux
                - 🔧 Test alarmes et protections
                - 🔧 Calibrage régulateur (si nécessaire)
                """)
                
            with maintenance_tabs[2]:
                st.markdown("""
                **📋 Tâches annuelles (2h - Technicien recommandé)**
                - 🔬 Test capacité batteries
                - 🔬 Mesure isolement installation
                - 🔬 Vérification mise à la terre
                - 🔬 Contrôle performances vs prévisions
                - 🔬 Mise à jour firmware équipements
                """)
            
            # Conseils d'optimisation
            st.markdown("#### 🚀 Conseils d'optimisation énergétique")
            
            optimisation_data = [
                {"Conseil": "Utiliser appareils énergivores en journée", "Économie": "15-20%", "Difficulté": "Facile"},
                {"Conseil": "Installer minuteries sur éclairage", "Économie": "10-15%", "Difficulté": "Facile"},
                {"Conseil": "Remplacer ampoules par LED", "Économie": "60-80%", "Difficulté": "Facile"},
                {"Conseil": "Optimiser température frigo (4-6°C)", "Économie": "10-15%", "Difficulté": "Facile"},
                {"Conseil": "Ajouter délestage automatique", "Économie": "5-10%", "Difficulté": "Moyen"},
                {"Conseil": "Installer système de monitoring", "Économie": "5-15%", "Difficulté": "Moyen"}
            ]
            
            df_optim = pd.DataFrame(optimisation_data)
            st.dataframe(
                df_optim,
                use_container_width=True,
                hide_index=True
            )
            
            # Section PSH PVGIS séparée
            if st.session_state.get("pvgis_monthly_psh"):
                with st.expander("🌍 Données PVGIS utilisées", expanded=False):
                    psh_used = st.session_state.get("solar_hours_override")
                    pvgis_mode = st.session_state.get("pvgis_month_mode")
                    label_mode = "mois choisi" if (pvgis_mode or "").startswith("Choisir") else "PSH minimale (saison creuse)"
                    
                    col_psh1, col_psh2 = st.columns(2)
                    with col_psh1:
                        st.metric("PSH utilisé pour le calcul", f"{psh_used:.2f} h/jour", help="Heures de soleil équivalent utilisées dans le dimensionnement")
                        st.info(f"**Mode sélectionné :** {label_mode}")
                    
                    with col_psh2:
                        st.subheader("📊 PSH mensuel PVGIS")
                        df_psh = pd.DataFrame({
                            "Mois": ["Jan","Fév","Mar","Avr","Mai","Jun","Juil","Aoû","Sep","Oct","Nov","Déc"],
                            "PSH (h)": [st.session_state["pvgis_monthly_psh"].get(i+1, None) for i in range(12)]
                        })
                        st.bar_chart(df_psh.set_index("Mois"))
                    
                    st.caption("📡 Source: PVGIS (Photovoltaic Geographical Information System) - Commission Européenne")
        else:
            st.error("❌ Veuillez entrer une consommation supérieure à 0")

with tab2:
    st.header("💰 Devis Estimatif Détaillé")
    
    if 'equipements' not in st.session_state:
        st.warning("⚠️ Veuillez d'abord effectuer un dimensionnement dans l'onglet 'Dimensionnement'")
    else:
        st.markdown("### ⚙️ Options du devis")
        
        # Sélection de la région pour le calcul de la main d'œuvre
        region_selectionnee = st.selectbox(
            "🌍 Région d'installation",
            options=REGIONS_SENEGAL,
            index=0,
            help="Sélectionnez la région où sera installé le système solaire. Le pourcentage de main d'œuvre sera appliqué automatiquement."
        )
        
        # Nom du demandeur
        nom_demandeur = st.text_input(
            "👤 Nom du demandeur",
            placeholder="Entrez le nom du demandeur du devis",
            help="Le nom du demandeur apparaîtra sur le devis généré"
        )
        
        # Récupération du taux accessoires depuis les paramètres admin (extrait valeur numérique)
        taux_accessoires_admin_data = get_accessories_rate()
        if isinstance(taux_accessoires_admin_data, dict):
            taux_accessoires_admin = taux_accessoires_admin_data.get('rate')
        else:
            taux_accessoires_admin = taux_accessoires_admin_data
        if taux_accessoires_admin is None:
            initialize_accessories_rate_in_firebase({'rate': TAUX_ACCESSOIRES_DEFAUT})
            taux_accessoires_admin = TAUX_ACCESSOIRES_DEFAUT
        
        devis = calculer_devis(st.session_state.equipements, use_online=False, accessoires_rate=float(taux_accessoires_admin)/100.0, region_selectionnee=region_selectionnee)
        
        # Résumé du système
        st.markdown("### 📋 Résumé de votre installation")
        col_info1, col_info2, col_info3, col_info4 = st.columns(4)
        
        with col_info1:
            st.metric("Consommation", f"{st.session_state.consommation:.1f} kWh/jour")
        with col_info2:
            st.metric("Puissance totale", f"{devis['puissance_totale']:.2f} kWc")
        with col_info3:
            voltage_display = st.session_state.choix['voltage']
            voltage_text = voltage_display if voltage_display == "High Voltage" else f"{voltage_display}V"
            st.metric("Type système", f"{voltage_text} {st.session_state.choix['type_batterie']}")
        with col_info4:
            surface_m2_resume = devis['puissance_totale'] * SURFACE_PAR_KWC_M2 * (1 + MARGE_IMPLANTATION_SURFACE_PCT/100.0)
            st.metric("Surface panneaux approx.", f"{surface_m2_resume:.1f} m²")
        
        st.caption(f"🎯 Autonomie souhaitée: {(st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)}% • Estimée: {(st.session_state.autonomie_reelle_pct if 'autonomie_reelle_pct' in st.session_state else (st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)):.0f}%")
        
        st.markdown("---")
        st.markdown("### 📦 Détails du devis")
        
        # Style CSS pour le tableau Excel
        table_style = """
        <style>
        .excel-table {
            width: 100%;
            border-collapse: collapse;
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 10px 0;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }
        .excel-table th {
            background: linear-gradient(135deg, #4CAF50, #45a049);
            color: white;
            font-weight: bold;
            padding: 15px 10px;
            text-align: left;
            border: 1px solid #ddd;
            font-size: 16px;
        }
        .excel-table td {
            padding: 12px 10px;
            border: 1px solid #ddd;
            font-size: 15px;
        }
        .excel-table tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        .excel-table tr:nth-child(odd) {
            background-color: #ffffff;
        }
        .excel-table tr:hover {
            background-color: #f5f5f5;
        }
        .excel-table .price-cell {
            text-align: right;
            font-weight: 500;
        }
        .excel-table .total-cell {
            font-weight: bold;
            color: #2E7D32;
        }
        .excel-table .qty-cell {
            text-align: center;
            font-weight: 500;
        }
        .total-row {
            background: linear-gradient(135deg, #E8F5E8, #C8E6C9) !important;
            font-weight: bold;
            font-size: 18px;
        }
        .total-row td {
            border-top: 2px solid #4CAF50;
            padding: 18px 10px;
        }
        </style>
        """
        
        # Construire le tableau HTML
        table_html = table_style + """
        <div style="overflow-x:auto;">
<table class="excel-table">
            <thead>
                <tr>
                    <th style="width: 40%;">📦 Équipement</th>
                    <th style="width: 10%;">📊 Qté</th>
                    <th style="width: 25%;">💰 Prix unitaire (FCFA)</th>
                    <th style="width: 25%;">💵 Sous-total (FCFA)</th>
                </tr>
            </thead>
            <tbody>
        """
        
        # Ajouter les lignes du devis
        for item in devis["details"]:
            table_html += f"""
                <tr>
                    <td>{item['item']}</td>
                    <td class="qty-cell">x{item['quantite']}</td>
                    <td class="price-cell">{item['prix_unitaire']:,}</td>
                    <td class="price-cell total-cell">{item['sous_total']:,}</td>
                </tr>
            """
        
        # Ajouter la ligne de total
        table_html += f"""
                <tr class="total-row">
                    <td colspan="3"><strong>💰 TOTAL ESTIMATIF</strong></td>
                    <td class="price-cell"><strong>{devis['total']:,}</strong></td>
                </tr>
            </tbody>
        </table>
</div>
        """
        
        # Afficher le tableau avec st.components.v1.html pour un rendu garanti
        import streamlit.components.v1 as components
        components.html(table_html, height=400)
        
        # Estimation facture électricité (Senelec)
        st.markdown("---")
        st.markdown("### ⚡ Estimation facture électricité (Senelec)")
        kwh_mensuel_total = (st.session_state.consommation if 'consommation' in st.session_state else 10.0) * 30

        # Production solaire estimée à partir des équipements actifs (option choisie ou dimensionnement)
        equip_actifs = st.session_state.get('equip_choisi', st.session_state.get('equipements', None))
        prod_kwh_j = 0.0
        autonomie_reelle_pct = 0.0
        if equip_actifs and equip_actifs.get('panneau'):
            panneau_nom, nb = equip_actifs['panneau']
            if panneau_nom and nb > 0:
                current_prices = get_current_prices()
                if current_prices and 'panneaux' in current_prices and panneau_nom in current_prices['panneaux']:
                    puissance_unitaire = current_prices['panneaux'][panneau_nom]['puissance']
                    puissance_totale_w = puissance_unitaire * nb
                else:
                    puissance_totale_w = 0
                # 5h d'ensoleillement/jour avec pertes ~25%
                prod_kwh_j = (puissance_totale_w / 1000.0) * 5.0 * 0.75
                conso_totale = st.session_state.consommation if 'consommation' in st.session_state else 10.0
                autonomie_reelle_pct = min(100.0, (prod_kwh_j / conso_totale) * 100.0)

        kwh_mensuel_solaire = prod_kwh_j * 30.0
        kwh_mensuel_apres = max(kwh_mensuel_total - kwh_mensuel_solaire, 0.0)

        # Sauvegarde pour autres sections
        st.session_state.production_solaire_kwh_j = prod_kwh_j
        st.session_state.autonomie_reelle_pct = autonomie_reelle_pct

        # Calcul coût Senelec après solaire
        palier1 = min(150.0, kwh_mensuel_apres)
        palier2 = min(max(kwh_mensuel_apres - 150.0, 0.0), 100.0)
        palier3 = max(kwh_mensuel_apres - 250.0, 0.0)
        cout_mensuel_senelec = palier1 * 124.17 + palier2 * 136.49 + palier3 * 159.36

        # Affichage en montants (FCFA/mois)
        palier1_av = min(150.0, kwh_mensuel_total)
        palier2_av = min(max(kwh_mensuel_total - 150.0, 0.0), 100.0)
        palier3_av = max(kwh_mensuel_total - 250.0, 0.0)
        cout_mensuel_avant = palier1_av * 124.17 + palier2_av * 136.49 + palier3_av * 159.36
        economie_mensuelle = max(cout_mensuel_avant - cout_mensuel_senelec, 0.0)

        col_sen1, col_sen2, col_sen3 = st.columns(3)
        with col_sen1:
            st.metric("Avant solaire", f"{cout_mensuel_avant:,.0f} FCFA/mois")
        with col_sen2:
            st.metric("Après solaire estimé", f"{cout_mensuel_senelec:,.0f} FCFA/mois")
        with col_sen3:
            st.metric("Économie estimée", f"{economie_mensuelle:,.0f} FCFA/mois")
        st.caption(f"Couverture réelle estimée: {autonomie_reelle_pct:.0f}%")
        
        # (Section paiement supprimée; notes importantes déplacées en bas)
        
        # Économies sur 10 ans
        st.markdown("---")
        st.markdown("### 💡 Analyse financière")
        
        # Calcul des économies basées sur la couverture réelle
        cout_electricite_kwh = 100  # FCFA par kWh (Senelec)
        conso_couverte_reelle = st.session_state.production_solaire_kwh_j if 'production_solaire_kwh_j' in st.session_state else (st.session_state.consommation_couverte if 'consommation_couverte' in st.session_state else st.session_state.consommation)
        conso_totale = st.session_state.consommation if 'consommation' in st.session_state else conso_couverte_reelle
        conso_couverte_reelle = min(conso_couverte_reelle, conso_totale)
        economie_annuelle = conso_couverte_reelle * 365 * cout_electricite_kwh
        economie_10ans = economie_annuelle * 10
        retour_investissement = devis['total'] / economie_annuelle if economie_annuelle > 0 else float('inf')
        
        col_eco1, col_eco2, col_eco3 = st.columns(3)
        
        with col_eco1:
            st.metric("💰 Économie annuelle", f"{economie_annuelle:,.0f} FCFA")
        with col_eco2:
            st.metric("📈 Économie sur 10 ans", f"{economie_10ans:,.0f} FCFA")
        with col_eco3:
            st.metric("⏱️ Retour sur investissement", f"{retour_investissement:.1f} ans")
        
        if retour_investissement < 5:
            st.success(f"✅ Excellent investissement ! Rentabilisé en {retour_investissement:.1f} ans")
        elif retour_investissement < 8:
            st.info(f"👍 Bon investissement ! Rentabilisé en {retour_investissement:.1f} ans")
        else:
            st.warning(f"⚠️ Investissement long terme : {retour_investissement:.1f} ans")
        
        # Boutons de téléchargement
        st.markdown("---")
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            
            # Génération du devis texte
            devis_text = f"""
╔════════════════════════════════════════════════════════════════╗
║        DEVIS ESTIMATIF - INSTALLATION SOLAIRE SÉNÉGAL         ║
╚════════════════════════════════════════════════════════════════╝

👤 INFORMATIONS CLIENT
{'─' * 64}
Nom du demandeur        : {nom_demandeur if nom_demandeur else "Non renseigné"}
Région d'installation   : {region_selectionnee}

📊 RÉSUMÉ DU SYSTÈME
{'─' * 64}
Consommation totale     : {st.session_state.consommation:.1f} kWh/jour
Autonomie souhaitée     : {(st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)} %
Autonomie estimée       : {(st.session_state.autonomie_reelle_pct if 'autonomie_reelle_pct' in st.session_state else (st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)):.0f} %
Couverte estimée        : {(st.session_state.production_solaire_kwh_j if 'production_solaire_kwh_j' in st.session_state else (st.session_state.consommation_couverte if 'consommation_couverte' in st.session_state else st.session_state.consommation)):.1f} kWh/jour
Puissance installée     : {devis['puissance_totale']:.2f} kWc
Surface panneaux approx. : {devis['puissance_totale'] * SURFACE_PAR_KWC_M2 * (1 + MARGE_IMPLANTATION_SURFACE_PCT/100.0):.1f} m²
Type de batterie        : {st.session_state.choix['type_batterie']}
Voltage système         : {st.session_state.choix['voltage'] if st.session_state.choix['voltage'] == "High Voltage" else f"{st.session_state.choix['voltage']}V"}
Type onduleur           : {st.session_state.choix['type_onduleur']}

📦 DÉTAILS DES ÉQUIPEMENTS
{'─' * 64}
"""
            for item in devis["details"]:
                devis_text += f"""
{item['item']}
  Quantité        : {item['quantite']}
  Prix unitaire   : {item['prix_unitaire']:,} FCFA
  Sous-total      : {item['sous_total']:,} FCFA
"""
            
            devis_text += f"""
{'═' * 64}
💰 TOTAL ESTIMATIF : {devis['total']:,} FCFA
{'═' * 64}


💡 ANALYSE FINANCIÈRE
{'─' * 64}
Économie annuelle estimée      : {economie_annuelle:,.0f} FCFA
Économie sur 10 ans            : {economie_10ans:,.0f} FCFA
Retour sur investissement      : {retour_investissement:.1f} ans

📝 NOTES IMPORTANTES
{'─' * 64}
- Prix indicatifs
- Installation standard incluse
- Garantie selon fabricant (panneaux: 25 ans, batteries: variable)
- Maintenance recommandée tous les 6 mois

{'═' * 64}
Document généré automatiquement
Pour plus d'informations : energiesolairesenegal.com
{'═' * 64}
"""
            
            # Génération du devis Word (.docx) avec tableau professionnel
            doc = Document()
            
            # En-tête avec logo
            header_paragraph = doc.add_paragraph()
            header_paragraph.alignment = 1  # Centré
            
            # Ajouter le logo s'il existe
            try:
                if os.path.exists("logo-solaire.svg"):
                    # Convertir SVG en image temporaire pour Word (python-docx ne supporte pas SVG directement)
                    # Pour l'instant, on ajoute juste le texte avec emoji
                    run = header_paragraph.add_run("☀️ ENERGIE SOLAIRE SÉNÉGAL\n")
                    run.font.size = Pt(16)
                    run.bold = True
                else:
                    run = header_paragraph.add_run("☀️ ENERGIE SOLAIRE SÉNÉGAL\n")
                    run.font.size = Pt(16)
                    run.bold = True
            except:
                run = header_paragraph.add_run("☀️ ENERGIE SOLAIRE SÉNÉGAL\n")
                run.font.size = Pt(16)
                run.bold = True
            
            # Titre principal
            title = doc.add_heading('DEVIS ESTIMATIF - INSTALLATION SOLAIRE SÉNÉGAL', 0)
            title.alignment = 1  # Centré
            
            # Informations client
            doc.add_heading('👤 INFORMATIONS CLIENT', level=1)
            client_table = doc.add_table(rows=2, cols=2)
            client_table.style = 'Table Grid'
            client_table.cell(0, 0).text = 'Nom du demandeur'
            client_table.cell(0, 1).text = nom_demandeur if nom_demandeur else "Non renseigné"
            client_table.cell(1, 0).text = 'Région d\'installation'
            client_table.cell(1, 1).text = region_selectionnee
            
            # Résumé du système
            doc.add_heading('📊 RÉSUMÉ DU SYSTÈME', level=1)
            resume_table = doc.add_table(rows=6, cols=2)
            resume_table.style = 'Table Grid'
            resume_table.cell(0, 0).text = 'Consommation totale'
            resume_table.cell(0, 1).text = f"{st.session_state.consommation:.1f} kWh/jour"
            resume_table.cell(1, 0).text = 'Autonomie souhaitée'
            resume_table.cell(1, 1).text = f"{(st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)} %"
            resume_table.cell(2, 0).text = 'Puissance installée'
            resume_table.cell(2, 1).text = f"{devis['puissance_totale']:.2f} kWc"
            resume_table.cell(3, 0).text = 'Type de batterie'
            resume_table.cell(3, 1).text = st.session_state.choix['type_batterie']
            resume_table.cell(4, 0).text = 'Voltage système'
            voltage_display = st.session_state.choix['voltage']
            resume_table.cell(4, 1).text = voltage_display if voltage_display == "High Voltage" else f"{voltage_display}V"
            resume_table.cell(5, 0).text = 'Type onduleur'
            resume_table.cell(5, 1).text = st.session_state.choix['type_onduleur']
            
            # Tableau des équipements
            doc.add_heading('📦 DÉTAILS DES ÉQUIPEMENTS', level=1)
            equip_table = doc.add_table(rows=len(devis['details']) + 1, cols=4)
            equip_table.style = 'Table Grid'
            
            # En-têtes du tableau
            hdr_cells = equip_table.rows[0].cells
            hdr_cells[0].text = 'Équipement'
            hdr_cells[1].text = 'Quantité'
            hdr_cells[2].text = 'Prix unitaire (FCFA)'
            hdr_cells[3].text = 'Sous-total (FCFA)'
            
            # Données du tableau
            for i, item in enumerate(devis['details']):
                row_cells = equip_table.rows[i + 1].cells
                row_cells[0].text = item['item']
                row_cells[1].text = str(item['quantite'])
                row_cells[2].text = f"{item['prix_unitaire']:,}"
                row_cells[3].text = f"{item['sous_total']:,}"
            
            # Total
            doc.add_heading('💰 TOTAL ESTIMATIF', level=1)
            total_table = doc.add_table(rows=1, cols=2)
            total_table.style = 'Table Grid'
            total_table.cell(0, 0).text = 'TOTAL'
            total_table.cell(0, 1).text = f"{devis['total']:,} FCFA"
            
            # Analyse financière
            doc.add_heading('💡 ANALYSE FINANCIÈRE', level=1)
            analyse_table = doc.add_table(rows=3, cols=2)
            analyse_table.style = 'Table Grid'
            analyse_table.cell(0, 0).text = 'Économie annuelle estimée'
            analyse_table.cell(0, 1).text = f"{economie_annuelle:,.0f} FCFA"
            analyse_table.cell(1, 0).text = 'Économie sur 10 ans'
            analyse_table.cell(1, 1).text = f"{economie_10ans:,.0f} FCFA"
            analyse_table.cell(2, 0).text = 'Retour sur investissement'
            analyse_table.cell(2, 1).text = f"{retour_investissement:.1f} ans"
            
            # Notes importantes
            doc.add_heading('📝 NOTES IMPORTANTES', level=1)
            notes = [
                "• Prix indicatifs",
                "• Installation standard incluse",
                "• Garantie selon fabricant (panneaux: 25 ans, batteries: variable)",
                "• Maintenance recommandée tous les 6 mois"
            ]
            for note in notes:
                doc.add_paragraph(note)
            
            # Informations de contact Energie Solaire Sénégal
            doc.add_heading('📞 INFORMATIONS DE CONTACT', level=1)
            contact_table = doc.add_table(rows=5, cols=2)
            contact_table.style = 'Table Grid'
            contact_table.cell(0, 0).text = '🏢 Entreprise'
            contact_table.cell(0, 1).text = 'Energie Solaire Sénégal'
            contact_table.cell(1, 0).text = '📍 Adresse'
            contact_table.cell(1, 1).text = 'Castor 221 Dakar, Sénégal (En face du terrain de Football)'
            contact_table.cell(2, 0).text = '📧 Email'
            contact_table.cell(2, 1).text = 'energiesolairesenegal@gmail.com'
            contact_table.cell(3, 0).text = '📞 Téléphones'
            contact_table.cell(3, 1).text = '+221 77 631 42 25 / +221 78 177 39 26'
            contact_table.cell(4, 0).text = '🌐 Site web'
            contact_table.cell(4, 1).text = 'energiesolairesenegal.com'
            
            # Pied de page
            doc.add_paragraph()
            footer = doc.add_paragraph("Document généré automatiquement")
            footer.alignment = 1  # Centré
            footer_info = doc.add_paragraph("Votre partenaire de confiance pour l'énergie solaire au Sénégal")
            footer_info.alignment = 1  # Centré
            
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            st.download_button(
                "📥 Télécharger le devis (Word .docx)",
                docx_buffer.getvalue(),
                file_name=f"devis_solaire_{st.session_state.choix['voltage']}V.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col_dl2:
            # Génération Excel (.xlsx) avec mise en forme professionnelle
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
            from openpyxl.utils.dataframe import dataframe_to_rows
            from openpyxl import Workbook
            
            # Créer un nouveau classeur
            wb = Workbook()
            ws = wb.active
            ws.title = "Devis Solaire"
            
            # Styles
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="4CAF50", end_color="4CAF50", fill_type="solid")
            title_font = Font(bold=True, size=14, color="2E7D32")
            border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                          top=Side(style='thin'), bottom=Side(style='thin'))
            center_alignment = Alignment(horizontal='center', vertical='center')
            
            # En-tête du document
            ws.merge_cells('A1:E1')
            ws['A1'] = "DEVIS ESTIMATIF - INSTALLATION SOLAIRE SÉNÉGAL"
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            
            # Informations client
            row = 3
            ws[f'A{row}'] = "INFORMATIONS CLIENT"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            nom_demandeur = st.session_state.get('nom_demandeur', 'Non renseigné')
            region_selectionnee = st.session_state.get('region_selectionnee', 'Non spécifiée')
            
            ws[f'A{row}'] = f"Nom du demandeur: {nom_demandeur}"
            row += 1
            ws[f'A{row}'] = f"Région d'installation: {region_selectionnee}"
            row += 2
            
            # Résumé du système
            ws[f'A{row}'] = "RÉSUMÉ DU SYSTÈME"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            ws[f'A{row}'] = f"Consommation totale: {st.session_state.consommation:.1f} kWh/jour"
            row += 1
            ws[f'A{row}'] = f"Puissance installée: {devis['puissance_totale']:.2f} kWc"
            row += 1
            ws[f'A{row}'] = f"Type de batterie: {st.session_state.choix['type_batterie']}"
            row += 1
            voltage_display = st.session_state.choix['voltage']
            voltage_text = voltage_display if voltage_display == "High Voltage" else f"{voltage_display}V"
            ws[f'A{row}'] = f"Voltage système: {voltage_text}"
            row += 1
            ws[f'A{row}'] = f"Type onduleur: {st.session_state.choix['type_onduleur']}"
            row += 2
            
            # Tableau des équipements
            ws[f'A{row}'] = "DÉTAILS DES ÉQUIPEMENTS"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            # En-têtes du tableau
            headers = ["Équipement", "Quantité", "Prix unitaire (FCFA)", "Sous-total (FCFA)"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = border
                cell.alignment = center_alignment
            
            row += 1
            
            # Données des équipements
            for item in devis["details"]:
                ws.cell(row=row, column=1, value=item["item"]).border = border
                ws.cell(row=row, column=2, value=item["quantite"]).border = border
                ws.cell(row=row, column=3, value=f"{item['prix_unitaire']:,}").border = border
                ws.cell(row=row, column=4, value=f"{item['sous_total']:,}").border = border
                row += 1
            
            # Ligne de total
            ws.cell(row=row, column=1, value="TOTAL").font = Font(bold=True)
            ws.cell(row=row, column=1).border = border
            ws.cell(row=row, column=2, value="").border = border
            ws.cell(row=row, column=3, value="").border = border
            total_cell = ws.cell(row=row, column=4, value=f"{devis['total']:,}")
            total_cell.font = Font(bold=True)
            total_cell.border = border
            total_cell.fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
            
            # Ajuster la largeur des colonnes
            ws.column_dimensions['A'].width = 40
            ws.column_dimensions['B'].width = 12
            ws.column_dimensions['C'].width = 20
            ws.column_dimensions['D'].width = 20
            
            # Notes importantes
            row += 3
            ws[f'A{row}'] = "NOTES IMPORTANTES"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            notes = [
                "• Prix indicatifs",
                "• Installation standard incluse",
                "• Garantie selon fabricant (panneaux: 25 ans, batteries: variable)",
                "• Maintenance recommandée tous les 6 mois"
            ]
            
            for note in notes:
                ws[f'A{row}'] = note
                row += 1
            
            # Contact
            row += 2
            ws[f'A{row}'] = "CONTACT - ENERGIE SOLAIRE SÉNÉGAL"
            ws[f'A{row}'].font = Font(bold=True, size=12, color="4CAF50")
            row += 1
            ws[f'A{row}'] = "📍 Castor 221 Dakar, Sénégal (En face du terrain de Football)"
            row += 1
            ws[f'A{row}'] = "📧 energiesolairesenegal@gmail.com"
            row += 1
            ws[f'A{row}'] = "📞 +221 77 631 42 25 / +221 78 177 39 26"
            row += 1
            ws[f'A{row}'] = "🌐 energiesolairesenegal.com"
            
            # Sauvegarder dans un buffer
            xlsx_buffer = io.BytesIO()
            wb.save(xlsx_buffer)
            xlsx_buffer.seek(0)
            
            st.download_button(
                "📊 Télécharger (Excel .xlsx)",
                xlsx_buffer.getvalue(),
                file_name=f"devis_solaire_{st.session_state.choix['voltage']}V.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        # Partage de devis avec coordonnées (formulaire détaillé)
        st.markdown("---")
        st.markdown("### 📤 Partager mon devis au service technique")
        
        with st.expander("📋 Partager mon devis au service technique", expanded=False):
            st.info("✉️ Remplissez ce formulaire pour partager votre devis au service technique. Ces informations facilitent un suivi rapide.")
            
            with st.form("form_partage_devis"):
                col_contact1_dev, col_contact2_dev = st.columns(2)
                
                with col_contact1_dev:
                    nom_dev = st.text_input("👤 Nom complet *", placeholder="Ex: Amadou Diallo")
                    tel_dev = st.text_input("📱 Téléphone *", placeholder="Ex: +221 77 123 45 67")
                    email_dev = st.text_input("📧 Email *", placeholder="Ex: amadou@example.com")
                
                with col_contact2_dev:
                    ville_dev = st.text_input("🏙️ Ville *", placeholder="Ex: Dakar")
                    quartier_dev = st.text_input("📍 Quartier/Zone", placeholder="Ex: Plateau, Almadies...")
                    type_batiment_dev = st.selectbox("🏠 Type de bâtiment", 
                                                   ["Maison individuelle", "Appartement", "Commerce", "Bureau", "Industrie", "Autre"])
                
                # Informations sur le projet
                st.markdown("#### 🔧 Détails du projet")
                col_projet1_dev, col_projet2_dev = st.columns(2)
                
                with col_projet1_dev:
                    urgence_dev = st.selectbox("⏰ Urgence du projet", 
                                             ["Pas urgent (> 6 mois)", "Moyen terme (3-6 mois)", "Court terme (1-3 mois)", "Urgent (< 1 mois)"])
                    budget_estime_dev = st.selectbox("💰 Budget estimé", 
                                                   ["< 500 000 FCFA", "500 000 - 1 000 000 FCFA", "1 000 000 - 2 000 000 FCFA", 
                                                    "2 000 000 - 5 000 000 FCFA", "> 5 000 000 FCFA", "À définir"])
                
                with col_projet2_dev:
                    installation_existante_dev = st.radio("⚡ Installation électrique existante", 
                                                     ["Raccordé au réseau SENELEC", "Groupe électrogène", "Aucune installation", "Autre"])
                    visite_technique_dev = st.checkbox("🔍 Demander une visite technique sur site")
                
                # Zone de commentaires
                commentaires_dev = st.text_area("💬 Questions ou commentaires spécifiques", 
                                              placeholder="Décrivez vos besoins spécifiques, contraintes, questions...", 
                                              height=100)
                
                # Consentement
                consent_dev = st.checkbox("✅ J'accepte d'être contacté par l'équipe technique d'Energie Solaire Sénégal *")
                
                # Bouton de soumission
                if st.form_submit_button("📤 Envoyer mon devis", type="primary", use_container_width=True):
                    # Validation des champs obligatoires
                    if not nom_dev or not tel_dev or not ville_dev or not email_dev or not consent_dev:
                        st.error("❌ Veuillez remplir les champs obligatoires (*) dont l’email, et accepter d'être contacté.")
                    elif '@' not in email_dev or '.' not in email_dev.split('@')[-1]:
                        st.error("❌ Email invalide.")
                    else:
                        quote_data = {
                            'timestamp': pd.Timestamp.now().isoformat(),
                            'consommation_kwh_jour': st.session_state.consommation,
                            'voltage_systeme': st.session_state.choix['voltage'],
                            'type_batterie': st.session_state.choix['type_batterie'],
                            'type_onduleur': st.session_state.choix['type_onduleur'],
                            'puissance_totale_kwc': devis['puissance_totale'],
                            'autonomie_souhaitee_pct': st.session_state.get('autonomie_pct', 100),
                            'autonomie_reelle_pct': st.session_state.get('autonomie_reelle_pct', 100),
                            'prix_total_fcfa': devis['total'],
                            'details_equipements': devis['details'],
                            'economie_mensuelle_fcfa': economie_mensuelle,
                            'retour_investissement_ans': retour_investissement,
                            'contact_info': {
                                'name': nom_dev.strip(),
                                'phone': tel_dev.strip(),
                                'email': email_dev.strip(),
                                'ville': ville_dev.strip(),
                                'quartier': quartier_dev.strip(),
                                'type_batiment': type_batiment_dev,
                                'urgence': urgence_dev,
                                'budget_estime': budget_estime_dev,
                                'installation_existante': installation_existante_dev,
                                'visite_technique': bool(visite_technique_dev),
                                'commentaires': commentaires_dev.strip(),
                                'demande_contact': bool(consent_dev),
                                'source': 'Application Dimensionnement Solaire - Devis Client'
                            }
                        }
                        quote_id = save_quote_to_firebase(quote_data)
                        if quote_id:
                            st.success(f"✅ Devis envoyé au service technique ! Référence: {quote_id[:8]}")
                            st.balloons()
                        else:
                            st.error("❌ Erreur lors du partage")
        
        # (Section Demander un contact du support technique supprimée)
        
    # Notes importantes (placées en bas)
    st.markdown("---")
    st.markdown("### 📝 Notes importantes")
    st.warning("""

    **Le prix final peut varier selon :**
    - La complexité de l'installation
    - L'accessibilité du site
    - Les promotions en cours
    """)
    
    # Section de gestion du stock pour les administrateurs
    if st.session_state.get('user_role') == 'admin' and 'equipements' in st.session_state:
        st.markdown("---")
        st.markdown("### 📦 Gestion du Stock - Équipements Dimensionnés")
        
        with st.expander("🔄 Actions sur le stock", expanded=False):
            st.info("💡 Gérez le stock des équipements sélectionnés pour ce dimensionnement")
            
            # Récupération des équipements dimensionnés
            equip = st.session_state.equipements
            
            # Vérification de la disponibilité en stock
            st.markdown("#### 📊 Vérification de la disponibilité")
            
            equipements_necessaires = []
            if equip["panneau"][0]:  # Si un panneau est sélectionné
                equipements_necessaires.append({
                    "type": "panneau",
                    "nom": equip["panneau"][0],
                    "quantite": equip["panneau"][1]
                })
            
            if equip["batterie"][0]:  # Si une batterie est sélectionnée
                equipements_necessaires.append({
                    "type": "batterie", 
                    "nom": equip["batterie"][0],
                    "quantite": equip["batterie"][1]
                })
            
            if equip["onduleur"][0]:  # Si un onduleur est sélectionné
                onduleur_nom, nb_onduleurs = equip["onduleur"] if isinstance(equip["onduleur"], tuple) else (equip["onduleur"], 1)
                equipements_necessaires.append({
                    "type": "onduleur",
                    "nom": onduleur_nom,
                    "quantite": nb_onduleurs
                })
            
            if equip["regulateur"]:  # Si un régulateur est sélectionné
                equipements_necessaires.append({
                    "type": "regulateur",
                    "nom": equip["regulateur"],
                    "quantite": 1
                })
            
            # Vérification du stock
            stock_status = check_stock_availability(equipements_necessaires)
            
            if stock_status["disponible"]:
                st.success("✅ Tous les équipements sont disponibles en stock !")
            else:
                st.warning("⚠️ Certains équipements ne sont pas disponibles en quantité suffisante")
                
                if stock_status["manquants"]:
                    st.markdown("**Équipements manquants :**")
                    for item in stock_status["manquants"]:
                        st.error(f"❌ {item['nom']} - Besoin: {item['quantite_demandee']}, Stock: {item['stock_disponible']}")
                
                if stock_status["stock_faible"]:
                    st.markdown("**Stock faible :**")
                    for item in stock_status["stock_faible"]:
                        st.warning(f"⚠️ {item['nom']} - Besoin: {item['quantite_demandee']}, Stock: {item['stock_disponible']}")
            
            # Actions sur le stock
            st.markdown("#### ⚡ Actions rapides")
            
            col_action1, col_action2, col_action3 = st.columns(3)
            
            with col_action1:
                pass  # Synchronisation dimensionnement → stock retirée
            
            with col_action2:
                if st.button("📋 Créer devis/facture", use_container_width=True):
                    # Rediriger vers l'éditeur de factures avec les équipements pré-remplis
                    st.session_state['equipements_pour_facture'] = equipements_necessaires
                    st.info("💡 Rendez-vous dans l'onglet 'Gestion de Stock' > 'Factures' pour créer le document")
            
            with col_action3:
                if st.button("📦 Réserver le stock", use_container_width=True):
                    if stock_status["disponible"]:
                        # Simuler une réservation en créant un mouvement de stock
                        for item in equipements_necessaires:
                            # Ici on pourrait implémenter une vraie réservation
                            # Pour l'instant, on affiche juste un message
                            pass
                        st.success("✅ Stock réservé pour ce devis (fonctionnalité à implémenter)")
                    else:
                        st.error("❌ Impossible de réserver - stock insuffisant")
            
            # Affichage détaillé du stock pour chaque équipement
            st.markdown("#### 📋 Détail du stock par équipement")
            
            for item in equipements_necessaires:
                with st.container():
                    col_detail1, col_detail2, col_detail3 = st.columns([2, 1, 1])
                    
                    with col_detail1:
                        st.markdown(f"**{item['nom']}** ({item['type']})")
                    
                    with col_detail2:
                        st.markdown(f"Besoin: **{item['quantite']}**")
                    
                    with col_detail3:
                        stock_info = get_stock_for_dimensioning_product(item['type'], item['nom'])
                        if stock_info:
                            stock_qty = stock_info.get('quantite', 0)
                            if stock_qty >= item['quantite']:
                                st.success(f"Stock: {stock_qty}")
                            elif stock_qty > 0:
                                st.warning(f"Stock: {stock_qty}")
                            else:
                                st.error("Rupture")
                        else:
                            st.info("Non sync.")
        
with tab3:
    st.header("☀️ Conseiller solaire")
    
    api_ready = ('DEEPSEEK_API_KEY' in st.secrets) and bool(st.secrets.get('DEEPSEEK_API_KEY', ''))
    if not api_ready:
        st.warning("⚠️ Clé API DeepSeek manquante. Ajoutez-la au fichier '.streamlit/secrets.toml' sous 'DEEPSEEK_API_KEY'.")
        st.info("👉 La configuration se fait uniquement via le fichier de secrets.")
    else:
        # Contexte du dimensionnement
        contexte = ""
        if 'dimensionnement' in st.session_state:
            dim = st.session_state.dimensionnement
            choix = st.session_state.choix

            # Estimation de la couverture réelle à partir des équipements actifs
            equip_actifs_ctx = st.session_state.get('equip_choisi', st.session_state.get('equipements', None))
            prod_kwh_j_ctx = 0.0
            auto_reelle_ctx = (st.session_state.autonomie_reelle_pct if 'autonomie_reelle_pct' in st.session_state else None)
            if not auto_reelle_ctx and equip_actifs_ctx and equip_actifs_ctx.get('panneau'):
                pn_ctx, nb_ctx = equip_actifs_ctx['panneau']
                if pn_ctx and nb_ctx > 0:
                    current_prices_ctx = get_current_prices()
                    if current_prices_ctx and 'panneaux' in current_prices_ctx and pn_ctx in current_prices_ctx['panneaux']:
                        p_unit_ctx = current_prices_ctx['panneaux'][pn_ctx]['puissance']
                        p_tot_ctx = p_unit_ctx * nb_ctx
                        prod_kwh_j_ctx = (p_tot_ctx / 1000.0) * 5.0 * 0.75
                        conso_totale_ctx = st.session_state.consommation if 'consommation' in st.session_state else 10.0
                        auto_reelle_ctx = min(100.0, (prod_kwh_j_ctx / conso_totale_ctx) * 100.0)
                    else:
                        prod_kwh_j_ctx = 0.0
                        auto_reelle_ctx = 0.0
            prod_kwh_j_ctx = prod_kwh_j_ctx if prod_kwh_j_ctx > 0 else (st.session_state.production_solaire_kwh_j if 'production_solaire_kwh_j' in st.session_state else 0.0)
            auto_reelle_ctx = auto_reelle_ctx if auto_reelle_ctx is not None else (st.session_state.autonomie_reelle_pct if 'autonomie_reelle_pct' in st.session_state else (st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100))

            contexte = f"""
L'utilisateur a dimensionné une installation avec:
- Consommation totale: {st.session_state.consommation:.1f} kWh/jour
- Couverture souhaitée: {(st.session_state.autonomie_pct if 'autonomie_pct' in st.session_state else 100)}% ({(st.session_state.consommation_couverte if 'consommation_couverte' in st.session_state else st.session_state.consommation):.1f} kWh/j)
- Couverture estimée: {auto_reelle_ctx:.0f}% ({prod_kwh_j_ctx:.1f} kWh/j)
- Puissance panneaux: {dim['puissance_panneaux']:.0f} Wc
- Capacité batteries: {f"{(dim['capacite_batterie'] * (400 if choix['voltage'] == 'High Voltage' else choix['voltage'])) / 1000:.1f} kWh" if choix['voltage'] == "High Voltage" else f"{dim['capacite_batterie']:.0f} Ah"} ({choix['type_batterie']})
- Puissance onduleur: {dim['puissance_onduleur']:.0f} W ({choix['type_onduleur']})
- Voltage système: {choix['voltage'] if choix['voltage'] == "High Voltage" else f"{choix['voltage']}V"}
- Climat: Sénégal (chaleur, humidité, 5h ensoleillement moyen)
"""
        
        st.subheader("🎛️ Options d'équipements avec totaux")
        
        # Récupération du taux accessoires depuis les paramètres admin (extrait valeur numérique)
        options_accessoires_data = get_accessories_rate()
        if isinstance(options_accessoires_data, dict):
            options_accessoires_pct = options_accessoires_data.get('rate')
        else:
            options_accessoires_pct = options_accessoires_data
        if options_accessoires_pct is None:
            initialize_accessories_rate_in_firebase({'rate': TAUX_ACCESSOIRES_DEFAUT})
            options_accessoires_pct = TAUX_ACCESSOIRES_DEFAUT
        
        base_voltage = st.session_state.choix['voltage'] if 'choix' in st.session_state else 48

        options_spec = [
            {'nom':'Option Économique','type_batterie':'AGM','type_onduleur':'Off-Grid','type_regulateur':'PWM','voltage':12,'phase_type':'monophase'},
            {'nom':'Option Équilibrée','type_batterie':'GEL','type_onduleur':'Hybride','type_regulateur':None,'voltage':12,'phase_type':'monophase'},
            {'nom':'Option Premium','type_batterie':'Lithium','type_onduleur':'Online','type_regulateur':'MPPT','voltage':48,'phase_type':'monophase'},
            {'nom':'Option Ultra Premium','type_batterie':'Lithium HV','type_onduleur':'Hybride','type_regulateur':None,'voltage':48,'phase_type':'monophase'},
        ]

        for opt in options_spec:
            # Consommation couverte (si disponible), sinon consommation totale
            consommation_opt = (st.session_state.consommation_couverte if 'consommation_couverte' in st.session_state else (st.session_state.consommation if 'consommation' in st.session_state else 10.0))
            # Dimensionnement pour l’option
            dim_opt = calculer_dimensionnement(
                consommation_opt,
                voltage=opt.get('voltage', base_voltage),
                type_batterie=opt['type_batterie'],
                part_nuit=st.session_state.get('part_nuit', 55)  # Valeur par défaut si pas encore définie
            )
            choix_opt = {
                'type_batterie': opt['type_batterie'],
                'type_onduleur': opt['type_onduleur'],
                'voltage': opt.get('voltage', base_voltage),
                'phase_type': opt.get('phase_type', 'monophase')
            }
            if opt['type_onduleur'] != 'Hybride':
                choix_opt['type_regulateur'] = opt['type_regulateur']

            equip_opt = selectionner_equipements(dim_opt, choix_opt)
            # S'assurer que options_accessoires_pct n'est jamais None et convertir en float
            taux_accessoires_final = options_accessoires_pct if options_accessoires_pct is not None else TAUX_ACCESSOIRES_DEFAUT
            devis_opt = calculer_devis(equip_opt, use_online=False, accessoires_rate=float(taux_accessoires_final)/100.0)
            with st.expander(f"{opt['nom']} – Total: {devis_opt['total']:,} FCFA", expanded=False):
                st.markdown(f"• Batterie: {opt['type_batterie']}")
                
                # Affichage onduleur avec gestion du couplage
                onduleur_data = equip_opt['onduleur']
                if isinstance(onduleur_data, tuple):
                    onduleur_nom, nb_onduleurs = onduleur_data
                    if nb_onduleurs > 1:
                        st.markdown(f"• Onduleur: {nb_onduleurs} x {onduleur_nom} (couplage)")
                    else:
                        st.markdown(f"• Onduleur: {onduleur_nom}")
                else:
                    st.markdown(f"• Onduleur: {opt['type_onduleur']}")
                
                if equip_opt['regulateur']:
                    st.markdown(f"• Régulateur: {equip_opt['regulateur']}")
                st.markdown(f"• Panneaux: {equip_opt['panneau'][1]} x {equip_opt['panneau'][0]}")
                
                # Autonomie estimée pour cette option
                try:
                    pn = equip_opt['panneau'][0]
                    nbp = equip_opt['panneau'][1]
                    current_prices_opt = get_current_prices()
                    punit = 0
                    if current_prices_opt and 'panneaux' in current_prices_opt and pn in current_prices_opt['panneaux']:
                        punit = current_prices_opt['panneaux'][pn].get('puissance', 0)
                    prod_opt_kwh_j = (punit * nbp / 1000.0) * 5.0 * 0.75 if (pn and nbp > 0 and punit > 0) else 0.0
                    conso_tot = st.session_state.consommation if 'consommation' in st.session_state else 10.0
                    auto_opt_pct = min(100.0, (prod_opt_kwh_j / conso_tot) * 100.0) if conso_tot > 0 else 0.0
                    st.markdown(f"• Autonomie estimée: {auto_opt_pct:.0f}% ({prod_opt_kwh_j:.1f} kWh/j)")
                except Exception:
                    pass
                
                st.markdown("—")
                for item in devis_opt['details']:
                    tag = "site" if item['source_prix']=='site' else ("local" if item['source_prix']=='local' else "estimé")
                    line = f"{item['item']}: {item['quantite']} × {item['prix_unitaire']:,} FCFA ({tag})"
                    if item.get('url_source'):
                        st.markdown(f"{line}  • [Lien]({item['url_source']})")
                    else:
                        st.markdown(line)

                if st.button(f"Appliquer {opt['nom']}", key=f"apply_{opt['nom']}"):
                    st.session_state.option_choisie = opt['nom']
                    st.session_state.equip_choisi = equip_opt
                    st.session_state.devis_choisi = devis_opt
                    st.success("Option appliquée. Allez à l’onglet Devis pour exporter.")

        st.markdown("---")

        st.subheader("💬 Questions fréquentes")
        
        col_q1, col_q2, col_q3 = st.columns(3)
        
        with col_q1:
            if st.button("🔧 Entretien des panneaux", use_container_width=True):
                question = "Comment entretenir mes panneaux solaires au Sénégal avec la poussière et le sable ?"
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de l'expert (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        with col_q2:
            if st.button("⚡ Durée de vie", use_container_width=True):
                question = "Quelle est la durée de vie de mon installation et quand faut-il remplacer les équipements ?"
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        with col_q3:
            if st.button("🌧️ Saison des pluies", use_container_width=True):
                question = "Comment optimiser ma production pendant la saison des pluies au Sénégal ?"
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        st.markdown("---")
        
        col_q4, col_q5, col_q6 = st.columns(3)
        
        with col_q4:
            if st.button("🔋 Batterie Lithium vs AGM", use_container_width=True):
                question = "Pour le climat du Sénégal, quelle est la meilleure batterie : Lithium ou AGM ? Explique les avantages et inconvénients."
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        with col_q5:
            if st.button("🔌 Onduleur hybride", use_container_width=True):
                question = "Pourquoi choisir un onduleur hybride plutôt qu'un onduleur standard ?"
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        with col_q6:
            if st.button("💰 Rentabilité", use_container_width=True):
                question = "Mon installation est-elle rentable ? Comment calculer le retour sur investissement ?"
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("**Question:**")
                    st.info(question)
                    st.markdown("**Réponse de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question, contexte))
        
        st.markdown("---")
        st.subheader("✍️ Posez votre question personnalisée")
        
        # Question personnalisée
        question_utilisateur = st.text_area(
            "Votre question sur l'énergie solaire :",
            placeholder="Ex: Comment protéger mon installation contre la foudre pendant l'hivernage ?",
            height=100
        )
        
        col_send, col_clear = st.columns([3, 1])
        
        with col_send:
            envoyer_btn = st.button("📤 Envoyer la question", type="primary", use_container_width=True)
        
        with col_clear:
            if st.button("🗑️ Effacer", use_container_width=True):
                st.rerun()
        
        if envoyer_btn:
            if question_utilisateur and len(question_utilisateur.strip()) > 5:
                with st.spinner("🤔 Pape répond en streaming..."):
                    st.markdown("---")
                    st.markdown("**Votre question:**")
                    st.info(question_utilisateur)
                    st.markdown("**Réponse détaillée de Pape (streaming):**")
                    st.write_stream(appeler_assistant_ia_stream(question_utilisateur, contexte))
            else:
                st.warning("⚠️ Veuillez entrer une question (minimum 5 caractères)")

# Onglet Contact
with tab_contact:
    st.header("📞 Contact & Partenaire Officiel")
    
    # Section principale du partenaire
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        ## 🏢 **Energie Solaire Sénégal**
        ### Votre partenaire de confiance pour l'énergie solaire
        
        🥇 **Premier outil de dimensionnement solaire en ligne au Sénégal**
        
        Cette application a été développée en **partenariat officiel** avec **Energie Solaire Sénégal**, 
        leader dans la fourniture et l'installation d'équipements solaires au Sénégal.
        """)
        
        # Informations de contact
        st.markdown("### 📍 **Nos Coordonnées**")
        
        contact_col1, contact_col2 = st.columns(2)
        
        with contact_col1:
            st.markdown("""
            **🏠 Adresse :**  
            Castor 221 Dakar, Sénégal  
            *(En face du terrain de Football)*  
            Zac Mbao (pres du rond point AXA)
            
            **📧 Email :**  
            energiesolairesenegal@gmail.com
            """)
        
        with contact_col2:
            st.markdown("""
            **📞 Téléphones :**  
            • +221 77 631 42 25  
            • +221 78 177 39 26  
            • +221 77 250 47 90
            
            **🌐 Site web :**  
            [energiesolairesenegal.com](https://energiesolairesenegal.com)
            """)
    
    with col2:
        try:
            st.image("logo-solaire.svg", caption="Energie Solaire Sénégal", use_container_width=True)
        except:
            st.markdown("### ☀️ Energie Solaire Sénégal")
            st.markdown("*Logo de l'entreprise*")
    
    st.markdown("---")
    
    # Services proposés
    st.markdown("### 🔧 **Nos Services**")
    
    service_col1, service_col2, service_col3 = st.columns(3)
    
    with service_col1:
        st.markdown("""
        **⚡ Installation Solaire**
        - Dimensionnement personnalisé
        - Installation professionnelle
        - Mise en service complète
        - Formation utilisateur
        """)
    
    with service_col2:
        st.markdown("""
        **🛠️ Maintenance & SAV**
        - Maintenance préventive
        - Réparations et dépannage
        - Remplacement de pièces
        - Support technique 24/7
        """)
    
    with service_col3:
        st.markdown("""
        **📦 Fourniture d'Équipements**
        - Panneaux solaires de qualité
        - Batteries haute performance
        - Onduleurs et régulateurs
        - Accessoires d'installation
        """)
    
    st.markdown("---")
    
    # Section commande
    st.markdown("### 🛒 **Commander votre Installation**")
    
    st.info("""
    **💡 Vous avez dimensionné votre installation ? Passez à l'action !**
    
    Pour commander votre installation solaire ou obtenir un devis personnalisé :
    
    1. **📞 Appelez-nous** : +221 77 631 42 25, +221 78 177 39 26 ou +221 77 250 47 90
    2. **📧 Envoyez-nous un email** : energiesolairesenegal@gmail.com
    3. **🌐 Visitez notre site** : [energiesolairesenegal.com](https://energiesolairesenegal.com)
    4. **🏠 Rendez-vous sur place** : Castor 221 Dakar (en face du terrain de Football) ou Zac Mbao (pres du rond point AXA)
    """)
    
    # Formulaire de contact rapide
    st.markdown("### 📝 **Contact Rapide**")
    
    with st.form("contact_form"):
        contact_col1, contact_col2 = st.columns(2)
        
        with contact_col1:
            nom_contact = st.text_input("Nom complet *")
            telephone_contact = st.text_input("Téléphone *")
        
        with contact_col2:
            email_contact = st.text_input("Email")
            region_contact = st.selectbox("Région", REGIONS_SENEGAL)
        
        type_demande = st.selectbox("Type de demande", [
            "Devis personnalisé",
            "Information produit",
            "Installation",
            "Maintenance/SAV",
            "Autre"
        ])
        
        message_contact = st.text_area("Votre message", 
                                     placeholder="Décrivez votre projet ou votre besoin...")
        
        submitted = st.form_submit_button("📤 Envoyer la demande")
        
        if submitted:
            if nom_contact and telephone_contact and message_contact:
                # Ici on pourrait intégrer avec Firebase pour sauvegarder la demande
                st.success("""
                ✅ **Demande envoyée avec succès !**
                
                Nous vous contacterons dans les plus brefs délais.
                
                **En attendant, vous pouvez nous joindre directement :**
                - 📞 +221 77 631 42 25
                - 📞 +221 78 177 39 26
                - 📞 +221 77 250 47 90
                - 📧 energiesolairesenegal@gmail.com
                """)
            else:
                st.error("⚠️ Veuillez remplir au minimum : Nom, Téléphone et Message")
    
    st.markdown("---")
    
    # Avantages du partenariat
    st.markdown("### 🤝 **Pourquoi choisir Energie Solaire Sénégal ?**")
    
    avantage_col1, avantage_col2 = st.columns(2)
    
    with avantage_col1:
        st.markdown("""
        **✅ Expertise Locale**
        - Connaissance du climat sénégalais
        - Adaptation aux conditions locales
        - Équipe technique qualifiée
        
        **✅ Qualité Garantie**
        - Équipements certifiés
        - Installation selon normes
        - Garantie fabricant respectée
        """)
    
    with avantage_col2:
        st.markdown("""
        **✅ Service Complet**
        - De l'étude à la mise en service
        - Formation et accompagnement
        - Maintenance et SAV
        
        **✅ Prix Compétitifs**
        - Tarifs transparents
        - Pas de frais cachés
        - Facilités de paiement
        """)

# Onglet Admin (seulement si connecté en tant qu'admin)
if is_user_authenticated() and is_admin_user():
    with tab_admin:
        st.header("⚙️ Panneau d'Administration")
        
        # Sous-onglets admin
        admin_tab1, admin_tab2, admin_tab3, admin_tab4, admin_tab5, admin_tab6, admin_tab7 = st.tabs(["💰 Gestion des Prix", "🔧 Main d'œuvre", "📋 Devis Clients", "📞 Demandes Clients", "🕘 Historique", "🧮 Calculateur", "📦 Gestion de Stock"])
        
        # Historique des modifications
        with admin_tab5:
            st.subheader("🕘 Historique des modifications")
            st.caption("Consultez les actions d’administration enregistrées.")

            # Libellés FR pour clarifier les types et objets
            EVENT_TYPE_LABELS = {
                "equipment_prices.update": "Prix des équipements — Mise à jour",
                "equipment_prices.init": "Prix des équipements — Initialisation",
                "labor_percentages.update": "Main d’œuvre — Mise à jour",
                "labor_percentages.init": "Main d’œuvre — Initialisation",
                "accessories_rate.update": "Taux accessoires — Mise à jour",
                "accessories_rate.init": "Taux accessoires — Initialisation",
                "quote.create": "Devis — Création",
                "quote.delete": "Devis — Suppression",
                "client_request.create": "Demande client — Création",
                "client_request.update_status": "Demande client — Mise à jour du statut",
                "client_request.delete": "Demande client — Suppression",
            }
            ITEM_ID_LABELS = {
                "equipment_prices": "Prix des équipements",
                "labor_percentages": "Pourcentages de main d’œuvre",
                "accessories_rate": "Taux accessoires",
                "quote": "Devis",
                "client_request": "Demande client",
                "global": "Global",
            }

            options_labels = ["Tous"] + list(EVENT_TYPE_LABELS.values())
            inverse_event_map = {v: k for k, v in EVENT_TYPE_LABELS.items()}

            colf1, colf2, colf3 = st.columns([2,2,1])
            with colf1:
                selected_label = st.selectbox(
                    "Type d’évènement",
                    options_labels,
                    index=0,
                    help="Filtre sur le type de changement"
                )
            with colf2:
                email_filter = st.text_input("Filtrer par email (optionnel)", value="")
            with colf3:
                limit = st.number_input("Limite", min_value=5, max_value=200, value=50, step=5)
            refresh = st.button("🔄 Recharger l'historique")

            # Récupération
            if refresh or True:
                et = None if selected_label == "Tous" else inverse_event_map.get(selected_label)
                email = email_filter.strip() or None
                try:
                    history = get_change_history(limit=int(limit), event_type=et, user_email=email)
                except Exception as e:
                    st.error(f"Erreur chargement historique: {e}")
                    history = []

                st.markdown(f"### {len(history)} évènement(s)")
                if not history:
                    st.info("Aucun évènement trouvé avec ces filtres.")
                else:
                    rows = []
                    for h in history:
                        ts = h.get('timestamp')
                        try:
                            ts_str = ts.strftime('%Y-%m-%d %H:%M:%S') if hasattr(ts, 'strftime') else str(ts)
                        except Exception:
                            ts_str = str(ts)
                        ev = h.get('event_type','')
                        obj = h.get('item_id','')
                        rows.append({
                            'Date': ts_str,
                            'Type': EVENT_TYPE_LABELS.get(ev, ev),
                            'Objet': ITEM_ID_LABELS.get(obj, obj),
                            'Utilisateur': h.get('user_email',''),
                            'Description': h.get('description',''),
                        })
                    try:
                        st.dataframe(pd.DataFrame(rows), use_container_width=True)
                    except Exception:
                        for r in rows:
                            st.write(r)

                    st.markdown("---")

                    # Helper pour afficher JSON joliment
                    def _show_json_block(titre, raw):
                        st.write(titre + ":")
                        if raw in (None, ""):
                            st.caption("—")
                            return
                        try:
                            import json
                            obj = json.loads(raw)
                            st.json(obj)
                        except Exception:
                            st.code(raw, language='json')

                    for i, h in enumerate(history, start=1):
                        ev = h.get('event_type','')
                        obj = h.get('item_id','')
                        with st.expander(f"Détail {i} • {EVENT_TYPE_LABELS.get(ev, ev)} • {ITEM_ID_LABELS.get(obj, obj)}"):
                            st.write(f"Date: {rows[i-1]['Date']}")
                            st.write(f"Utilisateur: {h.get('user_email','')}")
                            st.write(f"Description: {h.get('description','')}")
                            _show_json_block("Avant", h.get('before', ''))
                            _show_json_block("Après", h.get('after', ''))
        
        with admin_tab1:
            st.subheader("💰 Gestion des Prix des Équipements")
            
            # Boutons de gestion
            col_refresh, col_info = st.columns([1, 3])
            with col_refresh:
                if st.button("🔄 Recharger les prix (vider le cache)"):
                    st.cache_data.clear()
                    st.success("Cache vidé. Les prix seront rechargés.")
                    st.rerun()
            with col_info:
                st.caption("Rechargez le cache pour recharger les prix depuis Firebase.")
            
            # Charger les prix actuels depuis Firebase (sans fusion avec les valeurs par défaut)
            current_prices = get_current_prices()
            if current_prices and any(current_prices.get(cat) for cat in ["panneaux", "batteries", "onduleurs", "regulateurs"]):
                st.success("✅ Prix chargés depuis Firebase")
            else:
                st.warning("ℹ️ Aucun prix disponible dans Firebase.")
                col_reset, col_clear = st.columns([1, 1])
                with col_reset:
                    if st.button("Importer les prix par défaut"):
                        if save_equipment_prices(PRIX_EQUIPEMENTS):
                            st.success("✅ Prix par défaut importés dans Firebase.")
                            clear_prices_cache()
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors de l'import des prix par défaut")
                with col_clear:
                    if st.button("Vider tous les prix"):
                        if save_equipment_prices({"panneaux": {}, "batteries": {}, "onduleurs": {}, "regulateurs": {}}):
                            st.success("✅ Tous les prix ont été vidés.")
                            clear_prices_cache()
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors du vidage des prix")
            
            # Interface de modification des articles
            st.markdown("### 🔧 Modifier les articles existants")
            
            # Sélection de catégorie (liste fixe et filtrée)
            categories = ["panneaux", "batteries", "onduleurs", "regulateurs"]
            selected_category = st.selectbox("Choisir une catégorie", categories)
            
            if selected_category:
                st.markdown(f"#### {selected_category.title()}")
                
                # Afficher les équipements de la catégorie (Firebase uniquement)
                equipements = current_prices.get(selected_category, {})
                if not isinstance(equipements, dict):
                    equipements = {}
                
                if equipements:
                    # Sélection de l'article à modifier
                    article_names = list(equipements.keys())
                    selected_article = st.selectbox("Choisir un article à modifier", article_names)
                    
                    if selected_article:
                        st.markdown(f"**Modification de : {selected_article}**")
                        article_details = equipements[selected_article]
                        
                        # Formulaire de modification selon la catégorie
                        with st.form(f"modify_{selected_category}_{selected_article}"):
                            modified_item = {}
                            
                            if selected_category == "panneaux":
                                col1, col2 = st.columns(2)
                                with col1:
                                    new_puissance = st.number_input("Puissance (W)", min_value=0, step=10, value=article_details.get('puissance', 0))
                                    new_type = st.selectbox("Type", ["Monocristallin", "Polycristallin"], index=0 if article_details.get('type') == "Monocristallin" else 1)
                                with col2:
                                    new_voltage = st.number_input("Voltage (V)", min_value=0, step=12, value=article_details.get('voltage', 12))
                                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, value=article_details.get('prix', 0))
                                
                                modified_item = {
                                    "puissance": int(new_puissance),
                                    "voltage": int(new_voltage),
                                    "type": new_type,
                                    "prix": int(new_price)
                                }
                                
                            elif selected_category == "batteries":
                                col1, col2 = st.columns(2)
                                with col1:
                                    new_voltage = st.number_input("Voltage (V)", min_value=0, step=12, value=article_details.get('voltage', 12))
                                    battery_types = ["Plomb", "AGM", "GEL", "Lithium", "Lithium HV"]
                                    current_type = article_details.get('type', 'Plomb')
                                    type_index = battery_types.index(current_type) if current_type in battery_types else 0
                                    new_type = st.selectbox("Type", battery_types, index=type_index, 
                                                          help="Choisissez 'Lithium' pour voltage standard (12V-24V) ou 'Lithium HV' pour haute tension (48V+)",
                                                          key="battery_type_modify")
                                    
                                    # Champs de capacité conditionnels selon le type de batterie
                                    new_capacite = 0
                                    new_kwh = None
                                    
                                    # Utilisation d'un container pour forcer le rafraîchissement
                                    capacity_container = st.container()
                                    
                                    with capacity_container:
                                        if new_type == "Lithium HV":
                                            current_kwh = article_details.get('kwh', 0.0)
                                            new_kwh = st.number_input("Capacité (kWh)", min_value=0.0, step=0.1, value=current_kwh,
                                                                    help="Pour les batteries haute tension, spécifiez la capacité en kWh",
                                                                    key=f"battery_kwh_modify_{new_type}")
                                            st.info("💡 Les batteries Lithium HV sont spécifiées en kWh plutôt qu'en Ah")
                                            # Calcul automatique des Ah équivalents pour la compatibilité
                                            if new_kwh > 0 and new_voltage > 0:
                                                new_capacite = int((new_kwh * 1000) / new_voltage)
                                                st.caption(f"Équivalent: ~{new_capacite} Ah à {new_voltage}V")
                                            else:
                                                new_capacite = article_details.get('capacite', 0)
                                        else:
                                            new_capacite = st.number_input("Capacité (Ah)", min_value=0, step=10, value=article_details.get('capacite', 0),
                                                                         key=f"battery_ah_modify_{new_type}")
                                
                                with col2:
                                    new_cycles = st.number_input("Cycles", min_value=0, step=100, value=article_details.get('cycles', 0))
                                    new_decharge = st.number_input("Décharge max (%)", min_value=0, max_value=100, step=5, value=article_details.get('decharge_max', 50))
                                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, value=article_details.get('prix', 0))
                                
                                modified_item = {
                                    "capacite": int(new_capacite),
                                    "voltage": int(new_voltage),
                                    "type": new_type,
                                    "cycles": int(new_cycles),
                                    "decharge_max": int(new_decharge),
                                    "prix": int(new_price)
                                }
                                
                                # Ajouter le champ kWh si c'est une batterie Lithium HV
                                if new_type == "Lithium HV" and new_kwh is not None:
                                    modified_item["kwh"] = float(new_kwh)
                                
                            elif selected_category == "onduleurs":
                                col1, col2 = st.columns(2)
                                with col1:
                                    new_puissance = st.number_input("Puissance (W)", min_value=0, step=100, value=article_details.get('puissance', 0))
                                    new_voltage = st.number_input("Voltage (V)", min_value=0, step=12, value=article_details.get('voltage', 12))
                                    new_type = st.selectbox("Type", ["Off-Grid", "Hybride", "Online", "Online Tri"], index=["Off-Grid", "Hybride", "Online", "Online Tri"].index(article_details.get('type', 'Off-Grid')))
                                with col2:
                                    new_phase = st.selectbox("Phase", ["monophase", "triphase"], index=0 if article_details.get('phase', 'monophase') == 'monophase' else 1, help="Monophasé pour usage domestique, Triphasé pour usage industriel")
                                    new_mppt = st.text_input("MPPT (optionnel)", value=article_details.get('mppt', ''))
                                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, value=article_details.get('prix', 0))
                                
                                modified_item = {
                                    "puissance": int(new_puissance),
                                    "voltage": int(new_voltage),
                                    "type": new_type,
                                    "phase": new_phase,
                                    "mppt": new_mppt,
                                    "prix": int(new_price)
                                }
                                
                            elif selected_category == "regulateurs":
                                col1, col2 = st.columns(2)
                                with col1:
                                    new_amperage = st.number_input("Ampérage (A)", min_value=0, step=5, value=article_details.get('amperage', 0))
                                    new_type = st.selectbox("Type", ["PWM", "MPPT"], index=0 if article_details.get('type') == "PWM" else 1)
                                with col2:
                                    new_voltage_max = st.number_input("Voltage max (V)", min_value=0, step=12, value=article_details.get('voltage_max', 12))
                                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, value=article_details.get('prix', 0))
                                
                                modified_item = {
                                    "amperage": int(new_amperage),
                                    "type": new_type,
                                    "voltage_max": int(new_voltage_max),
                                    "prix": int(new_price)
                                }
                            
                            col_save, col_delete = st.columns(2)
                            with col_save:
                                save_button = st.form_submit_button("💾 Sauvegarder", type="primary")
                            with col_delete:
                                delete_button = st.form_submit_button("🗑️ Supprimer", type="secondary")
                            
                            if save_button:
                                # Validation spécifique pour les onduleurs
                                if selected_category == "onduleurs" and modified_item.get("puissance", 0) <= 0:
                                    st.warning("⚠️ La puissance de l'onduleur doit être supérieure à 0")
                                else:
                                    # Mettre à jour l'article dans la structure complète
                                    updated_prices = current_prices.copy()
                                    updated_prices[selected_category][selected_article] = modified_item
                                    
                                    # Sauvegarder dans Firebase
                                    if save_equipment_prices(updated_prices):
                                        st.success(f"✅ Article '{selected_article}' modifié avec succès!")
                                        # Vider le cache spécifique des prix
                                        clear_prices_cache()
                                        st.cache_data.clear()
                                        st.rerun()
                                    else:
                                        st.error("❌ Erreur lors de la sauvegarde")
                            
                            if delete_button:
                                updated_prices = current_prices.copy()
                                del updated_prices[selected_category][selected_article]
                                if save_equipment_prices(updated_prices):
                                    # Suppression croisée: retirer aussi le produit du stock si présent
                                    delete_stock_product_by_name_if_exists(selected_article)
                                    st.success(f"✅ Article '{selected_article}' supprimé avec succès!")
                                    clear_prices_cache()
                                    st.cache_data.clear()
                                    st.rerun()
                                else:
                                    st.error("❌ Erreur lors de la suppression")
                else:
                    st.info("Aucun article trouvé dans cette catégorie.")
                
                # Section séparée pour la suppression d'équipements
                st.markdown("---")
                st.markdown("#### 🗑️ Supprimer un équipement")
                
                if equipements:
                    # Sélection de l'équipement à supprimer
                    equipement_a_supprimer = st.selectbox(
                        "Choisir l'équipement à supprimer :",
                        options=[""] + list(equipements.keys()),
                        key=f"select_delete_{selected_category}"
                    )
                    
                    if equipement_a_supprimer:
                        col1, col2 = st.columns([1, 1])
                        
                        with col1:
                            if st.button(f"🗑️ Supprimer '{equipement_a_supprimer}'", key=f"delete_btn_{selected_category}"):
                                updated_prices = current_prices.copy()
                                if selected_category in updated_prices and equipement_a_supprimer in updated_prices[selected_category]:
                                    del updated_prices[selected_category][equipement_a_supprimer]
                                    if save_equipment_prices(updated_prices):
                                        # Suppression croisée: retirer aussi le produit du stock si présent
                                        delete_stock_product_by_name_if_exists(equipement_a_supprimer)
                                        st.success(f"✅ Article '{equipement_a_supprimer}' supprimé avec succès!")
                                        clear_prices_cache()
                                        st.cache_data.clear()
                                        st.rerun()
                                    else:
                                        st.error("❌ Erreur lors de la suppression")
                        
                        with col2:
                            st.info(f"Prix actuel : {equipements[equipement_a_supprimer]['prix']:,} FCFA")
                else:
                    st.info("Aucun équipement disponible dans cette catégorie.")
            
            # Ajout d'un nouvel article
            st.markdown("---")
            st.markdown("### ➕ Ajouter un nouvel article")
            
            # Affichage d'informations contextuelles selon la catégorie
            if selected_category == "batteries":
                st.info("""
                💡 **Guide des types de batteries :**
                - **Plomb/AGM/GEL** : Capacité en Ah (Ampères-heures)
                - **Lithium** : Capacité en Ah pour 12V-24V
                - **Lithium HV** : Capacité en kWh (kilowatts-heures) pour 48V+
                """)
            elif selected_category == "panneaux":
                st.info("💡 **Panneaux solaires** : Spécifiez la puissance en Watts et le type de technologie")
            elif selected_category == "onduleurs":
                st.info("💡 **Onduleurs** : Choisissez le type selon l'usage (Off-Grid, Hybride, Online)")
            elif selected_category == "regulateurs":
                st.info("💡 **Régulateurs** : MPPT sont plus efficaces que PWM (+30% de rendement)")
            
            with st.form(f"add_item_{selected_category}"):
                new_name = st.text_input("Nom de l'article", help="Nom descriptif de l'équipement (ex: 'Lithium HV 5.2kWh 48V')")
                
                if selected_category == "panneaux":
                    col1, col2 = st.columns(2)
                    with col1:
                        new_puissance = st.number_input("Puissance (W)", min_value=0, step=10, help="Puissance crête en Watts")
                        new_type = st.selectbox("Type", ["Polycristallin", "Monocristallin"], 
                                               help="Monocristallin = meilleur rendement, Polycristallin = plus économique") 
                    with col2:
                        new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, help="Prix de vente en FCFA")
                        # Calcul automatique du prix par Watt
                        if new_puissance > 0 and new_price > 0:
                            prix_par_watt = new_price / new_puissance
                            st.caption(f"💰 Prix par Watt: {prix_par_watt:.0f} FCFA/W")
                    
                    new_item = {
                        "puissance": int(new_puissance),
                        "type": new_type,
                        "prix": int(new_price)
                    }
                    
                elif selected_category == "batteries":
                    col1, col2 = st.columns(2)
                    with col1:
                        new_voltage = st.number_input("Voltage (V)", min_value=0, step=12, value=12, 
                                                    help="Tension nominale de la batterie")
                        battery_types = ["Plomb", "AGM", "GEL", "Lithium", "Lithium HV"]
                        new_type = st.selectbox("Type", battery_types, 
                                              help="Choisissez selon l'usage : Plomb (économique), AGM/GEL (sans entretien), Lithium (haute performance), Lithium HV (très haute performance)",
                                              key="battery_type_add") 
                    
                    with col2:
                        new_cycles = st.number_input("Cycles de vie", min_value=0, step=100, value=500,
                                                   help="Nombre de cycles charge/décharge")
                        new_decharge = st.number_input("Décharge max (%)", min_value=0, max_value=100, step=5, value=50,
                                                     help="Pourcentage de décharge maximale recommandée")
                    
                    # Section capacité avec logique conditionnelle améliorée
                    st.markdown("**Capacité de la batterie :**")
                    new_capacite = 0
                    new_kwh = None
                    
                    if new_type == "Lithium HV":
                        # Pour Lithium HV : priorité aux kWh
                        col_kwh, col_info = st.columns([2, 1])
                        with col_kwh:
                            new_kwh = st.number_input("Capacité (kWh)", min_value=0.0, step=0.1, value=4.8,
                                                    help="Capacité énergétique en kilowatts-heures",
                                                    key=f"battery_kwh_add_{new_type}")
                        with col_info:
                            st.markdown("""
                            <div style="background: #E3F2FD; padding: 10px; border-radius: 5px; margin-top: 25px;">
                                <small><strong>💡 Lithium HV</strong><br>
                                Spécification en kWh<br>
                                (plus précise que Ah)</small>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Calcul automatique des Ah équivalents
                        if new_kwh > 0 and new_voltage > 0:
                            new_capacite = int((new_kwh * 1000) / new_voltage)
                            st.success(f"✅ Équivalent calculé: ~{new_capacite} Ah à {new_voltage}V")
                        else:
                            st.warning("⚠️ Veuillez renseigner la capacité en kWh et le voltage")
                    else:
                        # Pour autres types : priorité aux Ah
                        col_ah, col_info = st.columns([2, 1])
                        with col_ah:
                            new_capacite = st.number_input("Capacité (Ah)", min_value=0, step=10, value=100,
                                                         help="Capacité en Ampères-heures",
                                                         key=f"battery_ah_add_{new_type}")
                        with col_info:
                            st.markdown(f"""
                            <div style="background: #E8F5E8; padding: 10px; border-radius: 5px; margin-top: 25px;">
                                <small><strong>💡 {new_type}</strong><br>
                                Spécification en Ah<br>
                                (standard du marché)</small>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Calcul automatique des kWh équivalents
                        if new_capacite > 0 and new_voltage > 0:
                            kwh_equivalent = (new_capacite * new_voltage) / 1000
                            st.info(f"💡 Équivalent énergétique: ~{kwh_equivalent:.1f} kWh à {new_voltage}V")
                    
                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, help="Prix de vente en FCFA")
                    
                    # Calcul du prix par kWh pour comparaison
                    if new_price > 0:
                        if new_type == "Lithium HV" and new_kwh and new_kwh > 0:
                            prix_par_kwh = new_price / new_kwh
                            st.caption(f"💰 Prix par kWh: {prix_par_kwh:,.0f} FCFA/kWh")
                        elif new_capacite > 0 and new_voltage > 0:
                            kwh_calc = (new_capacite * new_voltage) / 1000
                            if kwh_calc > 0:
                                prix_par_kwh = new_price / kwh_calc
                                st.caption(f"💰 Prix par kWh: {prix_par_kwh:,.0f} FCFA/kWh")
                    
                    new_item = {
                        "capacite": int(new_capacite),
                        "voltage": int(new_voltage),
                        "type": new_type,
                        "cycles": int(new_cycles),
                        "decharge_max": int(new_decharge),
                        "prix": int(new_price)
                    }
                    
                    # Ajouter le champ kWh si c'est une batterie Lithium HV
                    if new_type == "Lithium HV" and new_kwh is not None:
                        new_item["kwh"] = float(new_kwh)
                elif selected_category == "onduleurs":
                    col1, col2 = st.columns(2)
                    with col1:
                        new_puissance = st.number_input("Puissance (W)", min_value=0, step=100, value=1000,
                                                      help="Puissance de sortie continue en Watts")
                        new_voltage = st.number_input("Voltage entrée (V)", min_value=0, step=12, value=12,
                                                    help="Tension d'entrée DC (batterie)")
                        onduleur_types = ["Off-Grid", "Hybride", "Online", "Online Tri"]
                        new_type = st.selectbox("Type", onduleur_types,
                                              help="Off-Grid: autonome, Hybride: réseau+batterie, Online: UPS") 
                    with col2:
                        phase_options = ["monophase", "triphase"]
                        new_phase = st.selectbox("Phase", phase_options, 
                                               help="Monophasé: usage domestique, Triphasé: usage industriel")
                        new_mppt = st.text_input("MPPT (optionnel)", placeholder="ex: 60A",
                                                help="Courant MPPT intégré (pour onduleurs hybrides)")
                        new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, help="Prix de vente en FCFA")
                    
                    # Calculs d'information utiles
                    if new_puissance > 0:
                        if new_price > 0:
                            prix_par_watt = new_price / new_puissance
                            st.caption(f"💰 Prix par Watt: {prix_par_watt:.0f} FCFA/W")
                        
                        if new_voltage > 0:
                            courant_entree = new_puissance / new_voltage
                            st.info(f"⚡ Courant d'entrée estimé: ~{courant_entree:.1f}A à {new_voltage}V")
                            
                            # Recommandations selon le type
                            if new_type == "Hybride":
                                st.success(f"🔋 Idéal pour systèmes avec réseau électrique + batteries")
                            elif new_type == "Off-Grid":
                                st.info(f"🏠 Parfait pour sites isolés sans réseau électrique")
                            elif new_type in ["Online", "Online Tri"]:
                                st.warning(f"🏢 Recommandé pour applications critiques (UPS)")
                    
                    new_item = {
                        "puissance": int(new_puissance),
                        "voltage": int(new_voltage),
                        "type": new_type,
                        "phase": new_phase,
                        "mppt": new_mppt,
                        "prix": int(new_price)
                    }
                elif selected_category == "regulateurs":
                    col1, col2 = st.columns(2)
                    with col1:
                        new_amperage = st.number_input("Ampérage (A)", min_value=0, step=5, value=30,
                                                     help="Courant maximum de charge en Ampères")
                        regulateur_types = ["PWM", "MPPT"]
                        new_type = st.selectbox("Type", regulateur_types,
                                              help="MPPT: +30% efficacité vs PWM, PWM: plus économique")
                    with col2:
                        new_voltage_max = st.number_input("Voltage max (V)", min_value=0, step=12, value=100,
                                                        help="Tension maximale d'entrée des panneaux")
                        new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000, help="Prix de vente en FCFA")
                    
                    # Calculs d'information utiles
                    if new_amperage > 0:
                        # Estimation de la puissance pour différents voltages système
                        st.markdown("**💡 Puissance maximale recommandée :**")
                        col_12v, col_24v, col_48v = st.columns(3)
                        
                        with col_12v:
                            puissance_12v = new_amperage * 12
                            st.metric("12V", f"{puissance_12v}W")
                        with col_24v:
                            puissance_24v = new_amperage * 24
                            st.metric("24V", f"{puissance_24v}W")
                        with col_48v:
                            puissance_48v = new_amperage * 48
                            st.metric("48V", f"{puissance_48v}W")
                        
                        if new_type == "MPPT":
                            puissance_panneau_max = max(puissance_12v, puissance_24v, puissance_48v) * 1.3
                            st.success(f"🔋 Avec MPPT: Panneaux jusqu'à {puissance_panneau_max:.0f}W possibles (+30% efficacité)")
                        else:
                            st.info(f"🔋 Avec PWM: Respecter la tension batterie = tension panneaux")
                    
                    if new_price > 0 and new_amperage > 0:
                        prix_par_ampere = new_price / new_amperage
                        st.caption(f"💰 Prix par Ampère: {prix_par_ampere:.0f} FCFA/A")
                    
                    new_item = {
                        "amperage": int(new_amperage),
                        "type": new_type,
                        "voltage_max": int(new_voltage_max),
                        "prix": int(new_price)
                    }
                else:
                    new_price = st.number_input("Prix (FCFA)", min_value=0, step=1000)
                    new_item = {"prix": int(new_price)}
                
                add_submit = st.form_submit_button("➕ Ajouter l'article", type="primary")
                if add_submit:
                    # Validation du nom
                    if not new_name or len(new_name.strip()) < 2:
                        st.error("❌ **Nom invalide** : Veuillez renseigner un nom d'article d'au moins 2 caractères")
                    # Vérification de l'unicité du nom
                    elif new_name in current_prices.get(selected_category, {}):
                        st.error(f"❌ **Nom déjà existant** : L'article '{new_name}' existe déjà dans '{selected_category}'")
                    # Validations spécifiques par catégorie
                    elif selected_category == "panneaux":
                        if new_item.get("puissance", 0) <= 0:
                            st.error("❌ **Puissance invalide** : La puissance du panneau doit être supérieure à 0W")
                        elif new_item.get("prix", 0) <= 0:
                            st.error("❌ **Prix invalide** : Le prix doit être supérieur à 0 FCFA")
                        else:
                            # Validation réussie pour panneaux
                            updated_prices = current_prices.copy()
                            if selected_category not in updated_prices:
                                updated_prices[selected_category] = {}
                            updated_prices[selected_category][new_name] = new_item
                            if save_equipment_prices(updated_prices):
                                st.success(f"✅ **Panneau ajouté !** '{new_name}' dans '{selected_category}'")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Erreur lors de l'ajout de l'article")
                    elif selected_category == "batteries":
                        if new_item.get("voltage", 0) <= 0:
                            st.error("❌ **Voltage invalide** : Le voltage doit être supérieur à 0V")
                        elif new_item.get("capacite", 0) <= 0 and new_item.get("kwh", 0) <= 0:
                            st.error("❌ **Capacité invalide** : La capacité doit être supérieure à 0")
                        elif new_item.get("prix", 0) <= 0:
                            st.error("❌ **Prix invalide** : Le prix doit être supérieur à 0 FCFA")
                        elif new_item.get("cycles", 0) <= 0:
                            st.error("❌ **Cycles invalides** : Le nombre de cycles doit être supérieur à 0")
                        elif not (0 < new_item.get("decharge_max", 0) <= 100):
                            st.error("❌ **Décharge invalide** : La décharge max doit être entre 1% et 100%")
                        else:
                            # Validation réussie pour batteries
                            updated_prices = current_prices.copy()
                            if selected_category not in updated_prices:
                                updated_prices[selected_category] = {}
                            updated_prices[selected_category][new_name] = new_item
                            if save_equipment_prices(updated_prices):
                                st.success(f"✅ **Batterie ajoutée !** '{new_name}' dans '{selected_category}'")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Erreur lors de l'ajout de l'article")
                    elif selected_category == "onduleurs":
                        if new_item.get("puissance", 0) <= 0:
                            st.error("❌ **Puissance invalide** : La puissance de l'onduleur doit être supérieure à 0W")
                        elif new_item.get("voltage", 0) <= 0:
                            st.error("❌ **Voltage invalide** : Le voltage d'entrée doit être supérieur à 0V")
                        elif new_item.get("prix", 0) <= 0:
                            st.error("❌ **Prix invalide** : Le prix doit être supérieur à 0 FCFA")
                        else:
                            # Validation réussie pour onduleurs
                            updated_prices = current_prices.copy()
                            if selected_category not in updated_prices:
                                updated_prices[selected_category] = {}
                            updated_prices[selected_category][new_name] = new_item
                            if save_equipment_prices(updated_prices):
                                st.success(f"✅ **Onduleur ajouté !** '{new_name}' dans '{selected_category}'")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Erreur lors de l'ajout de l'article")
                    elif selected_category == "regulateurs":
                        if new_item.get("amperage", 0) <= 0:
                            st.error("❌ **Ampérage invalide** : L'ampérage doit être supérieur à 0A")
                        elif new_item.get("voltage_max", 0) <= 0:
                            st.error("❌ **Voltage max invalide** : Le voltage max doit être supérieur à 0V")
                        elif new_item.get("prix", 0) <= 0:
                            st.error("❌ **Prix invalide** : Le prix doit être supérieur à 0 FCFA")
                        else:
                            # Validation réussie pour régulateurs
                            updated_prices = current_prices.copy()
                            if selected_category not in updated_prices:
                                updated_prices[selected_category] = {}
                            updated_prices[selected_category][new_name] = new_item
                            if save_equipment_prices(updated_prices):
                                st.success(f"✅ **Régulateur ajouté !** '{new_name}' dans '{selected_category}'")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Erreur lors de l'ajout de l'article")
                    else:
                        # Catégorie générique
                        if new_item.get("prix", 0) <= 0:
                            st.error("❌ **Prix invalide** : Le prix doit être supérieur à 0 FCFA")
                        else:
                            updated_prices = current_prices.copy()
                            if selected_category not in updated_prices:
                                updated_prices[selected_category] = {}
                            updated_prices[selected_category][new_name] = new_item
                            if save_equipment_prices(updated_prices):
                                st.success(f"✅ **Article ajouté !** '{new_name}' dans '{selected_category}'")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Erreur lors de l'ajout de l'article")
            
            # Réinitialisation seulement
            st.markdown("---")
            st.markdown("### 🔁 Réinitialiser aux valeurs par défaut")
            if st.button("🔄 Réinitialiser aux valeurs par défaut", type="secondary"):
                if save_equipment_prices(PRIX_EQUIPEMENTS):
                    st.success("✅ Tous les prix ont été réinitialisés aux valeurs par défaut!")
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.error("❌ Erreur lors de la réinitialisation")
        
        with admin_tab2:
            st.markdown("### 🔧 Configuration des pourcentages de main d'œuvre par région")
            st.markdown("Configurez les pourcentages de main d'œuvre appliqués selon les régions du Sénégal.")
            
            # Récupérer les pourcentages actuels depuis Firebase
            try:
                pourcentages_actuels = get_labor_percentages()
                if not pourcentages_actuels:
                    # Initialiser avec les valeurs par défaut si aucune donnée n'existe
                    success, message = initialize_labor_percentages_in_firebase(POURCENTAGES_MAIN_OEUVRE_DEFAUT.copy())
                    if success:
                        pourcentages_actuels = POURCENTAGES_MAIN_OEUVRE_DEFAUT.copy()
                        st.info("✅ Pourcentages initialisés avec les valeurs par défaut")
                    else:
                        st.warning(f"⚠️ {message}")
                        pourcentages_actuels = POURCENTAGES_MAIN_OEUVRE_DEFAUT.copy()
            except Exception as e:
                st.error(f"❌ Erreur lors de la récupération des données: {str(e)}")
                pourcentages_actuels = POURCENTAGES_MAIN_OEUVRE_DEFAUT.copy()
            
            # Interface de modification
            st.markdown("#### 📊 Pourcentages actuels")
            
            # Créer un formulaire pour modifier les pourcentages
            with st.form("labor_percentages_form"):
                st.markdown("**Modifiez les pourcentages de main d'œuvre par région :**")
                
                # Créer des colonnes pour une meilleure présentation
                col1, col2 = st.columns(2)
                
                nouveaux_pourcentages = {}
                regions_list = list(REGIONS_SENEGAL)
                
                # Diviser les régions en deux colonnes
                mid_point = len(regions_list) // 2
                
                with col1:
                    for region in regions_list[:mid_point]:
                        valeur_actuelle = pourcentages_actuels.get(region, POURCENTAGES_MAIN_OEUVRE_DEFAUT.get(region, 20.0))
                        nouveaux_pourcentages[region] = st.number_input(
                            f"🏛️ {region}",
                            min_value=0.0,
                            max_value=50.0,
                            value=float(valeur_actuelle),
                            step=0.5,
                            format="%.1f",
                            help=f"Pourcentage de main d'œuvre pour {region}"
                        )
                
                with col2:
                    for region in regions_list[mid_point:]:
                        valeur_actuelle = pourcentages_actuels.get(region, POURCENTAGES_MAIN_OEUVRE_DEFAUT.get(region, 20.0))
                        nouveaux_pourcentages[region] = st.number_input(
                            f"🏛️ {region}",
                            min_value=0.0,
                            max_value=50.0,
                            value=float(valeur_actuelle),
                            step=0.5,
                            format="%.1f",
                            help=f"Pourcentage de main d'œuvre pour {region}"
                        )
                
                # Boutons d'action
                col_btn1, col_btn2, col_btn3 = st.columns(3)
                
                with col_btn1:
                    submit_button = st.form_submit_button("💾 Sauvegarder", type="primary")
                
                with col_btn2:
                    reset_button = st.form_submit_button("🔄 Réinitialiser")
                
                with col_btn3:
                    preview_button = st.form_submit_button("👁️ Aperçu")
                
                # Traitement des actions
                if submit_button:
                    try:
                        if save_labor_percentages(nouveaux_pourcentages):
                            clear_labor_percentages_cache()  # Vider le cache pour afficher les nouvelles données
                            st.success("✅ Pourcentages sauvegardés avec succès!")
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors de la sauvegarde")
                    except Exception as e:
                        st.error(f"❌ Erreur: {str(e)}")
                
                if reset_button:
                    try:
                        if save_labor_percentages(POURCENTAGES_MAIN_OEUVRE_DEFAUT):
                            clear_labor_percentages_cache()  # Vider le cache pour afficher les nouvelles données
                            st.success("✅ Pourcentages réinitialisés aux valeurs par défaut!")
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors de la réinitialisation")
                    except Exception as e:
                        st.error(f"❌ Erreur: {str(e)}")
                
                if preview_button:
                    st.markdown("#### 👁️ Aperçu des modifications")
                    
                    # Afficher les changements
                    col_prev1, col_prev2 = st.columns(2)
                    
                    with col_prev1:
                        st.markdown("**Valeurs actuelles:**")
                        for region in REGIONS_SENEGAL:
                            valeur_actuelle = pourcentages_actuels.get(region, POURCENTAGES_MAIN_OEUVRE_DEFAUT.get(region, 20.0))
                            st.write(f"• {region}: {valeur_actuelle:.1f}%")
                    
                    with col_prev2:
                        st.markdown("**Nouvelles valeurs:**")
                        for region in REGIONS_SENEGAL:
                            nouvelle_valeur = nouveaux_pourcentages.get(region, 20.0)
                            valeur_actuelle = pourcentages_actuels.get(region, POURCENTAGES_MAIN_OEUVRE_DEFAUT.get(region, 20.0))
                            
                            if nouvelle_valeur != valeur_actuelle:
                                st.write(f"• {region}: {nouvelle_valeur:.1f}% ⚠️")
                            else:
                                st.write(f"• {region}: {nouvelle_valeur:.1f}%")
            
            # Statistiques et informations
            st.markdown("---")
            st.markdown("#### 📈 Statistiques")
            
            col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
            
            valeurs_actuelles = list(pourcentages_actuels.values())
            
            with col_stat1:
                st.metric("Minimum", f"{min(valeurs_actuelles):.1f}%")
            with col_stat2:
                st.metric("Maximum", f"{max(valeurs_actuelles):.1f}%")
            with col_stat3:
                st.metric("Moyenne", f"{sum(valeurs_actuelles)/len(valeurs_actuelles):.1f}%")
            with col_stat4:
                st.metric("Régions", len(REGIONS_SENEGAL))
            
            # Aide et informations
            st.markdown("---")
            st.markdown("#### ℹ️ Aide")
            
            st.info("""
            **Comment utiliser cette interface :**

            1. **Modifier les pourcentages** : Ajustez les valeurs selon les coûts de main d'œuvre locaux
            2. **Sauvegarder** : Cliquez sur "Sauvegarder" pour appliquer les changements
            3. **Réinitialiser** : Restaure les valeurs par défaut du système
            4. **Aperçu** : Visualisez les changements avant de les sauvegarder

            **Notes importantes :**
            - Les pourcentages s'appliquent au coût total des équipements
            - Les valeurs sont comprises entre 0% et 50%
            - Les modifications sont appliquées immédiatement aux nouveaux devis
            """)
            
            # Section Taux accessoires
            st.markdown("---")
            st.markdown("### 🔧 Configuration du taux accessoires")
            st.markdown("Configurez le taux accessoires appliqué aux devis (câbles, fusibles, disjoncteurs, etc.).")
            
            # Récupérer le taux actuel depuis Firebase
            try:
                taux_data = get_accessories_rate()
                if isinstance(taux_data, dict):
                    taux_actuel = taux_data.get('rate')
                else:
                    taux_actuel = taux_data
                if taux_actuel is None:
                    # Initialiser avec la valeur par défaut si aucune donnée n'existe
                    success, message = initialize_accessories_rate_in_firebase({'rate': TAUX_ACCESSOIRES_DEFAUT})
                    if success:
                        taux_actuel = TAUX_ACCESSOIRES_DEFAUT
                        st.info("✅ Taux accessoires initialisé avec la valeur par défaut")
                    else:
                        st.warning(f"⚠️ {message}")
                        taux_actuel = TAUX_ACCESSOIRES_DEFAUT
            except Exception as e:
                st.error(f"❌ Erreur lors de la récupération du taux accessoires: {str(e)}")
                taux_actuel = TAUX_ACCESSOIRES_DEFAUT
            
            # Interface de modification
            with st.form("accessories_rate_form"):
                st.markdown("**Configurez le taux accessoires :**")
                
                nouveau_taux = st.number_input(
                    "🔌 Taux accessoires (%)",
                    min_value=0.0,
                    max_value=50.0,
                    value=float(taux_actuel),
                    step=0.5,
                    format="%.1f",
                    help="Pourcentage appliqué au coût total des équipements pour les accessoires (câbles, fusibles, disjoncteurs, etc.)"
                )
                
                # Boutons d'action
                col_btn1, col_btn2 = st.columns(2)
                
                with col_btn1:
                    submit_button = st.form_submit_button("💾 Sauvegarder", type="primary")
                
                with col_btn2:
                    reset_button = st.form_submit_button("🔄 Réinitialiser")
                
                # Traitement des actions
                if submit_button:
                    try:
                        if save_accessories_rate({'rate': nouveau_taux}):
                            clear_accessories_rate_cache()  # Vider le cache pour afficher les nouvelles données
                            st.success("✅ Taux accessoires sauvegardé avec succès!")
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors de la sauvegarde")
                    except Exception as e:
                        st.error(f"❌ Erreur: {str(e)}")
                
                if reset_button:
                    try:
                        if save_accessories_rate({'rate': TAUX_ACCESSOIRES_DEFAUT}):
                            clear_accessories_rate_cache()  # Vider le cache pour afficher les nouvelles données
                            st.success("✅ Taux accessoires réinitialisé à la valeur par défaut!")
                            st.rerun()
                        else:
                            st.error("❌ Erreur lors de la réinitialisation")
                    except Exception as e:
                        st.error(f"❌ Erreur: {str(e)}")
            
            # Informations sur le taux actuel
            st.markdown("#### 📊 Taux actuel")
            col_info1, col_info2 = st.columns(2)
            
            with col_info1:
                st.metric("Taux accessoires", f"{taux_actuel:.1f}%")
            
            with col_info2:
                st.metric("Valeur par défaut", f"{TAUX_ACCESSOIRES_DEFAUT:.1f}%")
            
            # Aide pour le taux accessoires
            st.markdown("---")
            st.markdown("#### ℹ️ À propos du taux accessoires")
            
            st.info("""
            **Le taux accessoires inclut :**
            - Câbles et connecteurs
            - Fusibles et disjoncteurs
            - Supports et fixations
            - Petits accessoires d'installation
            
            **Notes importantes :**
            - Ce taux s'applique au coût total des équipements principaux
            - Il remplace l'ancien choix client pour une gestion centralisée
            - Les modifications sont appliquées immédiatement aux nouveaux devis
            """)
        
        with admin_tab3:
            st.subheader("📋 Devis Partagés par les Clients")
            
            # Charger tous les devis depuis Firebase
            all_quotes = get_all_quotes()
            
            if all_quotes:
                st.success(f"✅ {len(all_quotes)} devis trouvé(s)")
                
                # Statistiques rapides
                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                
                # Calculer les statistiques
                total_value = sum([quote.get('total', 0) for quote in all_quotes])
                avg_value = total_value / len(all_quotes) if all_quotes else 0
                
                # Compter par type de batterie
                battery_types = {}
                for quote in all_quotes:
                    battery_type = quote.get('type_batterie', 'N/A')
                    battery_types[battery_type] = battery_types.get(battery_type, 0) + 1
                
                with col_stat1:
                    st.metric("📊 Total devis", len(all_quotes))
                with col_stat2:
                    st.metric("💰 Valeur totale", f"{total_value:,.0f} FCFA")
                with col_stat3:
                    st.metric("📈 Valeur moyenne", f"{avg_value:,.0f} FCFA")
                with col_stat4:
                    most_popular_battery = max(battery_types.items(), key=lambda x: x[1])[0] if battery_types else "N/A"
                    st.metric("🔋 Batterie populaire", most_popular_battery)
                
                # Filtres
                st.markdown("### 🔍 Filtres")
                col_filter1, col_filter2, col_filter3 = st.columns(3)
                
                with col_filter1:
                    # Filtre par type de batterie
                    battery_options = ["Tous"] + list(set([quote.get('type_batterie', 'N/A') for quote in all_quotes]))
                    filter_battery = st.selectbox("Type de batterie", battery_options)
                
                with col_filter2:
                    # Filtre par plage de prix
                    filter_price = st.selectbox("Plage de prix", [
                        "Tous",
                        "< 500k FCFA",
                        "500k - 1M FCFA", 
                        "1M - 2M FCFA",
                        "2M - 5M FCFA",
                        "> 5M FCFA"
                    ])
                
                with col_filter3:
                    # Filtre par puissance
                    filter_power = st.selectbox("Puissance", [
                        "Tous",
                        "< 1 kWc",
                        "1-3 kWc",
                        "3-5 kWc", 
                        "5-10 kWc",
                        "> 10 kWc"
                    ])
                
                # Appliquer les filtres
                filtered_quotes = all_quotes
                
                if filter_battery != "Tous":
                    filtered_quotes = [q for q in filtered_quotes if q.get('type_batterie') == filter_battery]
                
                if filter_price != "Tous":
                    if filter_price == "< 500k FCFA":
                        filtered_quotes = [q for q in filtered_quotes if q.get('total', 0) < 500000]
                    elif filter_price == "500k - 1M FCFA":
                        filtered_quotes = [q for q in filtered_quotes if 500000 <= q.get('total', 0) < 1000000]
                    elif filter_price == "1M - 2M FCFA":
                        filtered_quotes = [q for q in filtered_quotes if 1000000 <= q.get('total', 0) < 2000000]
                    elif filter_price == "2M - 5M FCFA":
                        filtered_quotes = [q for q in filtered_quotes if 2000000 <= q.get('total', 0) < 5000000]
                    elif filter_price == "> 5M FCFA":
                        filtered_quotes = [q for q in filtered_quotes if q.get('total', 0) >= 5000000]
                
                if filter_power != "Tous":
                    if filter_power == "< 1 kWc":
                        filtered_quotes = [q for q in filtered_quotes if q.get('puissance_totale', 0) < 1]
                    elif filter_power == "1-3 kWc":
                        filtered_quotes = [q for q in filtered_quotes if 1 <= q.get('puissance_totale', 0) < 3]
                    elif filter_power == "3-5 kWc":
                        filtered_quotes = [q for q in filtered_quotes if 3 <= q.get('puissance_totale', 0) < 5]
                    elif filter_power == "5-10 kWc":
                        filtered_quotes = [q for q in filtered_quotes if 5 <= q.get('puissance_totale', 0) < 10]
                    elif filter_power == "> 10 kWc":
                        filtered_quotes = [q for q in filtered_quotes if q.get('puissance_totale', 0) >= 10]
                
                st.info(f"📊 {len(filtered_quotes)} devis après filtrage")
                
                # Trier par date (plus récent en premier)
                filtered_quotes.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
                
                # Afficher les devis
                for i, quote in enumerate(filtered_quotes):
                    # Informations de base
                    timestamp = quote.get('timestamp', '')[:16].replace('T', ' ')
                    client_name = quote.get('nom_client', 'Client anonyme')
                    total_price = quote.get('total', 0)
                    power = quote.get('puissance_totale', 0)
                    
                    # Titre de l'expandeur
                    title = f"💰 {client_name} - {total_price:,} FCFA - {power:.1f} kWc - {timestamp}"
                    
                    with st.expander(title):
                        # Informations système
                        col_sys1, col_sys2, col_sys3 = st.columns(3)
                        
                        with col_sys1:
                            st.markdown("**⚡ Système**")
                            st.write(f"**Puissance:** {quote.get('puissance_totale', 0):.2f} kWc")
                            st.write(f"**Consommation:** {quote.get('consommation', 0):.1f} kWh/jour")
                            st.write(f"**Autonomie:** {quote.get('autonomie_pct', 100)}%")
                            st.write(f"**Voltage:** {quote.get('voltage', 'N/A')}V")
                        
                        with col_sys2:
                            st.markdown("**🔋 Équipements**")
                            st.write(f"**Batterie:** {quote.get('type_batterie', 'N/A')}")
                            st.write(f"**Onduleur:** {quote.get('type_onduleur', 'N/A')}")
                            st.write(f"**Région:** {quote.get('region', 'Non spécifiée')}")
                        
                        with col_sys3:
                            st.markdown("**👤 Client**")
                            st.write(f"**Nom:** {quote.get('nom_client', 'N/A')}")
                            st.write(f"**Email:** {quote.get('email_client', 'N/A')}")
                            st.write(f"**Téléphone:** {quote.get('telephone', 'N/A')}")
                            st.write(f"**Ville:** {quote.get('ville', 'N/A')}")
                        
                        # Détails des équipements
                        if quote.get('details'):
                            st.markdown("---")
                            st.markdown("**📦 Détails des équipements**")
                            
                            # Tableau des équipements
                            col_eq1, col_eq2, col_eq3, col_eq4 = st.columns([3, 1, 2, 2])
                            
                            with col_eq1:
                                st.markdown("**Équipement**")
                            with col_eq2:
                                st.markdown("**Qté**")
                            with col_eq3:
                                st.markdown("**Prix unitaire**")
                            with col_eq4:
                                st.markdown("**Sous-total**")
                            
                            st.markdown("---")
                            
                            for detail in quote.get('details', []):
                                col_eq1, col_eq2, col_eq3, col_eq4 = st.columns([3, 1, 2, 2])
                                
                                with col_eq1:
                                    st.write(detail.get('item', 'N/A'))
                                with col_eq2:
                                    st.write(f"x{detail.get('quantite', 0)}")
                                with col_eq3:
                                    st.write(f"{detail.get('prix_unitaire', 0):,} FCFA")
                                with col_eq4:
                                    st.write(f"**{detail.get('sous_total', 0):,} FCFA**")
                        
                        # Actions admin
                        st.markdown("---")
                        st.markdown("**⚙️ Actions Admin**")
                        
                        col_action1, col_action2, col_action3 = st.columns(3)
                        
                        with col_action1:
                            # Lien pour contacter le client
                            email = quote.get('email_client', '')
                            if email:
                                subject = f"Votre devis solaire - {quote.get('nom_client', '')}"
                                st.markdown(f"📧 [Contacter par email](mailto:{email}?subject={subject})")
                        
                        with col_action2:
                            # Lien pour appeler
                            phone = quote.get('telephone', '').replace(' ', '').replace('+', '')
                            if phone:
                                st.markdown(f"📞 [Appeler le client](tel:{phone})")
                        
                        with col_action3:
                            # Bouton de suppression
                            _confirm_del_q = st.checkbox(
                                "Confirmer suppression",
                                key=f"confirm_del_quote_{quote.get('id','')}"
                            )
                            if st.button("🗑️ Supprimer", key=f"btn_del_quote_{quote.get('id','')}"):
                                if _confirm_del_q:
                                    if delete_quote(quote.get('id')):
                                        st.success("✅ Devis supprimé.")
                                        st.rerun()
                                    else:
                                        st.error("❌ Échec de suppression du devis.")
                                else:
                                    st.warning("Veuillez cocher la confirmation avant suppression.")
                        
                        # Informations système
                        st.markdown("---")
                        st.caption(f"**ID:** {quote.get('id', 'N/A')[:8]}... | **Créé:** {timestamp}")
                
                # Actions en lot
                st.markdown("---")
                st.markdown("### 🔧 Actions en lot")
                col_bulk1, col_bulk2 = st.columns(2)
                
                with col_bulk1:
                    if st.button("📊 Exporter en CSV"):
                        # Préparer les données pour export
                        export_data = []
                        for quote in filtered_quotes:
                            export_data.append({
                                'Date': quote.get('timestamp', '')[:10],
                                'Nom_Client': quote.get('nom_client', ''),
                                'Email': quote.get('email_client', ''),
                                'Telephone': quote.get('telephone', ''),
                                'Ville': quote.get('ville', ''),
                                'Region': quote.get('region', ''),
                                'Puissance_kWc': quote.get('puissance_totale', 0),
                                'Consommation_kWh': quote.get('consommation', 0),
                                'Type_Batterie': quote.get('type_batterie', ''),
                                'Type_Onduleur': quote.get('type_onduleur', ''),
                                'Voltage': quote.get('voltage', ''),
                                'Autonomie_Pct': quote.get('autonomie_pct', 100),
                                'Prix_Total_FCFA': quote.get('total', 0)
                            })
                        
                        if export_data:
                            df_export = pd.DataFrame(export_data)
                            csv = df_export.to_csv(index=False)
                            st.download_button(
                                label="💾 Télécharger CSV",
                                data=csv,
                                file_name=f"devis_clients_{pd.Timestamp.now().strftime('%Y%m%d')}.csv",
                                mime="text/csv"
                            )
                
                with col_bulk2:
                    if st.button("🔄 Actualiser"):
                        st.rerun()
            
            else:
                st.info("📭 Aucun devis partagé pour le moment")
                st.markdown("Les devis apparaîtront ici quand les clients sauvegarderont leurs devis dans l'onglet Devis Estimatif Détaillé.")
        
        # Onglet 4: Gestion des demandes clients
        with admin_tab4:
            st.subheader("📞 Gestion des Demandes Clients")
            
            # Charger toutes les demandes depuis Firebase
            client_requests = get_all_client_requests()
            
            if client_requests:
                st.success(f"✅ {len(client_requests)} demande(s) trouvée(s)")
                
                # Statistiques rapides
                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                
                # Compter par statut
                status_counts = {}
                for req in client_requests:
                    status = req.get('status', 'nouveau')
                    status_counts[status] = status_counts.get(status, 0) + 1
                
                with col_stat1:
                    st.metric("🆕 Nouvelles", status_counts.get('nouveau', 0))
                with col_stat2:
                    st.metric("📞 En cours", status_counts.get('en_cours', 0))
                with col_stat3:
                    st.metric("✅ Traitées", status_counts.get('traite', 0))
                with col_stat4:
                    st.metric("📊 Total", len(client_requests))
                
                # Filtres
                st.markdown("### 🔍 Filtres")
                col_filter1, col_filter2, col_filter3 = st.columns(3)
                
                with col_filter1:
                    filter_status = st.selectbox("Statut", ["Tous", "nouveau", "en_cours", "traite"])
                with col_filter2:
                    filter_urgence = st.selectbox("Urgence", ["Toutes", "Urgent (< 1 mois)", "Court terme (1-3 mois)", "Moyen terme (3-6 mois)", "Pas urgent (> 6 mois)"])
                with col_filter3:
                    filter_ville = st.selectbox("Ville", ["Toutes"] + list(set([req.get('ville', '') for req in client_requests if req.get('ville')])))
                
                # Appliquer les filtres
                filtered_requests = client_requests
                if filter_status != "Tous":
                    filtered_requests = [r for r in filtered_requests if r.get('status', 'nouveau') == filter_status]
                if filter_urgence != "Toutes":
                    filtered_requests = [r for r in filtered_requests if r.get('urgence', '') == filter_urgence]
                if filter_ville != "Toutes":
                    filtered_requests = [r for r in filtered_requests if r.get('ville', '') == filter_ville]
                
                st.info(f"📊 {len(filtered_requests)} demande(s) après filtrage")
                
                # Trier par date (plus récent en premier)
                filtered_requests.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
                
                # Afficher les demandes
                for i, request in enumerate(filtered_requests):
                    _nom_client = request.get('nom_client', '') or 'Client'
                    _ville_client = request.get('ville', '')
                    _urgence = request.get('urgence', 'Normal')
                    _status = request.get('status', 'nouveau')
                    
                    # Icônes selon le statut
                    if _status == "nouveau":
                        status_icon = "🆕"
                    elif _status == "en_cours":
                        status_icon = "📞"
                    else:
                        status_icon = "✅"
                    
                    # Couleur selon l'urgence
                    if "Urgent (" in _urgence:
                        urgence_color = "🔴"
                    elif "Court terme" in _urgence:
                        urgence_color = "🟠"
                    elif "Moyen terme" in _urgence:
                        urgence_color = "🟡"
                    else:
                        urgence_color = "🟢"
                    
                    timestamp = request.get('timestamp', 'Date inconnue')[:16].replace('T', ' ')
                    _titre_demande = f"{status_icon} Demande #{i+1} - {timestamp} - {_nom_client}" + (f" - {_ville_client}" if _ville_client else "")
                    
                    with st.expander(f"{urgence_color} {_titre_demande}"):
                        # Informations principales
                        col_info1, col_info2, col_info3 = st.columns(3)
                        
                        with col_info1:
                            st.markdown("**👤 Contact**")
                            st.write(f"Nom: {request.get('nom_client', 'N/A')}")
                            st.write(f"Téléphone: {request.get('telephone', 'N/A')}")
                            st.write(f"Email: {request.get('email_client', 'N/A')}")
                            st.write(f"Ville: {request.get('ville', 'N/A')}")
                            st.write(f"Quartier: {request.get('quartier', 'N/A')}")
                        
                        with col_info2:
                            st.markdown("**🏠 Projet**")
                            st.write(f"Type: {request.get('type_batiment', 'N/A')}")
                            st.write(f"Urgence: {_urgence}")
                            st.write(f"Budget: {request.get('budget_estime', 'N/A')}")
                            st.write(f"Installation existante: {request.get('installation_existante', 'N/A')}")
                            st.write(f"Visite technique: {'Oui' if request.get('visite_technique', False) else 'Non'}")
                        
                        with col_info3:
                            st.markdown("**⚡ Dimensionnement**")
                            dim = request.get('dimensionnement', {})
                            st.write(f"Consommation: {dim.get('consommation_kwh_jour', 'N/A')} kWh/j")
                            st.write(f"Puissance: {dim.get('puissance_totale_kwc', 'N/A')} kWc")
                            st.write(f"Prix estimé: {dim.get('prix_total_fcfa', 'N/A'):,} FCFA" if dim.get('prix_total_fcfa') else "Prix: N/A")
                            st.write(f"Batterie: {dim.get('type_batterie', 'N/A')}")
                            voltage_display = dim.get('voltage_systeme', 'N/A')
                            if voltage_display == "High Voltage":
                                st.write(f"Voltage: {voltage_display}")
                            else:
                                st.write(f"Voltage: {voltage_display}V")
                        
                        # Commentaires
                        if request.get('commentaires'):
                            st.markdown("**💬 Commentaires:**")
                            st.write(request.get('commentaires', 'Aucun commentaire'))
                        
                        # Gestion du statut
                        st.markdown("---")
                        st.markdown("**📋 Gestion Admin**")
                        
                        col_status1, col_status2 = st.columns(2)
                        
                        with col_status1:
                            new_status = st.selectbox(
                                "Changer le statut",
                                ["nouveau", "en_cours", "traite"],
                                index=["nouveau", "en_cours", "traite"].index(_status),
                                key=f"status_{request.get('id', '')}"
                            )
                            
                            if st.button(f"💾 Mettre à jour", key=f"update_status_{request.get('id', '')}"):
                                if update_client_request_status(request.get('id'), new_status):
                                    st.success("✅ Statut mis à jour!")
                                    st.rerun()
                                else:
                                    st.error("❌ Erreur lors de la mise à jour")
                        
                        with col_status2:
                            admin_notes = st.text_area(
                                "Notes admin",
                                value=request.get('admin_notes', ''),
                                key=f"notes_{request.get('id', '')}",
                                height=100
                            )
                        
                        # Actions de suppression
                        st.markdown("---")
                        _confirm_del_r = st.checkbox(
                            "Confirmer la suppression de cette demande",
                            key=f"confirm_del_request_{request.get('id','')}"
                        )
                        if st.button(
                            "🗑️ Supprimer cette demande",
                            key=f"btn_del_request_{request.get('id','')}"
                        ):
                            if _confirm_del_r:
                                if delete_client_request(request.get('id')):
                                    st.success("✅ Demande supprimée.")
                                    st.rerun()
                                else:
                                    st.error("❌ Échec de suppression de la demande.")
                            else:
                                st.warning("Veuillez cocher la confirmation avant suppression.")
                        
                        # Informations système
                        st.markdown("---")
                        st.caption(f"**ID:** {request.get('id', 'N/A')[:8]}... | **Créé:** {timestamp} | **Source:** {request.get('source', 'N/A')}")
                
                # Actions en lot
                st.markdown("---")
                st.markdown("### 🔧 Actions en lot")
                col_bulk1, col_bulk2 = st.columns(2)
                
                with col_bulk1:
                    if st.button("📊 Exporter en CSV"):
                        # Préparer les données pour export
                        export_data = []
                        for req in filtered_requests:
                            dim = req.get('dimensionnement', {})
                            export_data.append({
                                'Date': req.get('timestamp', '')[:10],
                                'Nom': req.get('nom_client', ''),
                                'Téléphone': req.get('telephone', ''),
                                'Email': req.get('email_client', ''),
                                'Ville': req.get('ville', ''),
                                'Type_Batiment': req.get('type_batiment', ''),
                                'Urgence': req.get('urgence', ''),
                                'Budget': req.get('budget_estime', ''),
                                'Consommation_kWh': dim.get('consommation_kwh_jour', 0),
                                'Puissance_kWc': dim.get('puissance_totale_kwc', 0),
                                'Prix_FCFA': dim.get('prix_total_fcfa', 0),
                                'Statut': req.get('status', 'nouveau'),
                                'Notes_Admin': req.get('admin_notes', '')
                            })
                        
                        if export_data:
                            df_export = pd.DataFrame(export_data)
                            csv = df_export.to_csv(index=False)
                            st.download_button(
                                label="💾 Télécharger CSV",
                                data=csv,
                                file_name=f"demandes_clients_{pd.Timestamp.now().strftime('%Y%m%d')}.csv",
                                mime="text/csv"
                            )
                
                with col_bulk2:
                    if st.button("🔄 Actualiser"):
                        st.rerun()
            
            else:
                st.info("📭 Aucune demande client pour le moment")
                st.markdown("Les demandes apparaîtront ici quand les clients utiliseront le formulaire de contact dans l'onglet Dimensionnement.")
            
        # Onglet Calculateur
        with admin_tab6:
            st.subheader("🧮 Calculateur pour Installateurs Solaires")
            st.caption("Outils de calcul spécialisés pour les professionnels du solaire")
            
            # Sous-onglets pour différents calculateurs
            calc_tab1, calc_tab2, calc_tab3, calc_tab4 = st.tabs([
                "⚡ Section Câbles", 
                "🔌 Disjoncteurs", 
                "🔋 Autonomie", 
                "💰 Rentabilité"
            ])
            
            # Calculateur de section de câbles
            with calc_tab1:
                st.markdown("### ⚡ Calculateur de Section de Câbles")
                st.info("Calculez la section de câble nécessaire en fonction de la puissance, distance et chute de tension admissible.")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Paramètres électriques**")
                    puissance_cable = st.number_input("Puissance (W)", min_value=1, max_value=50000, value=3000, step=100)
                    tension_cable = st.selectbox("Tension système (V)", [12, 24, 48, 220, 380], index=2)
                    type_courant = st.selectbox("Type de courant", ["DC (Continu)", "AC (Alternatif)"], index=0)
                    
                with col2:
                    st.markdown("**Paramètres installation**")
                    distance_cable = st.number_input("Distance (m)", min_value=0.1, max_value=1000.0, value=10.0, step=0.5)
                    chute_tension_max = st.slider("Chute de tension max (%)", min_value=1.0, max_value=5.0, value=3.0, step=0.5)
                    temperature_amb = st.number_input("Température ambiante (°C)", min_value=20, max_value=60, value=30, step=5)
                
                if st.button("🔍 Calculer la section", type="primary"):
                    # Calcul du courant
                    courant = puissance_cable / tension_cable
                    
                    # Calcul de la résistance maximale admissible
                    chute_tension_v = (chute_tension_max / 100) * tension_cable
                    resistance_max = chute_tension_v / courant
                    
                    # Calcul de la section minimale (résistivité cuivre = 0.017 ohm.mm²/m)
                    resistivite_cuivre = 0.017
                    section_min = (resistivite_cuivre * 2 * distance_cable) / resistance_max
                    
                    # Facteur de correction température (simplifié)
                    facteur_temp = 1 + (temperature_amb - 20) * 0.004
                    section_corrigee = section_min * facteur_temp
                    
                    # Sections normalisées disponibles
                    sections_normalisees = [1.5, 2.5, 4, 6, 10, 16, 25, 35, 50, 70, 95, 120, 150, 185, 240, 300]
                    section_recommandee = next((s for s in sections_normalisees if s >= section_corrigee), 300)
                    
                    # Affichage des résultats
                    st.success("✅ Calcul terminé")
                    
                    col_res1, col_res2, col_res3 = st.columns(3)
                    with col_res1:
                        st.metric("Courant", f"{courant:.1f} A")
                        st.metric("Section minimale", f"{section_min:.2f} mm²")
                    with col_res2:
                        st.metric("Chute de tension", f"{chute_tension_v:.2f} V")
                        st.metric("Section corrigée", f"{section_corrigee:.2f} mm²")
                    with col_res3:
                        st.metric("**Section recommandée**", f"**{section_recommandee} mm²**")
                        
                        # Couleur du câble selon la section
                        if section_recommandee <= 2.5:
                            couleur = "🔵 Bleu/Noir"
                        elif section_recommandee <= 6:
                            couleur = "🟡 Jaune/Vert"
                        else:
                            couleur = "🔴 Rouge/Noir"
                        st.write(f"Couleur suggérée: {couleur}")
                    
                    # Informations complémentaires
                    st.markdown("---")
                    st.markdown("**📋 Informations complémentaires:**")
                    st.write(f"• Résistance linéique: {(resistivite_cuivre / section_recommandee):.4f} Ω/m")
                    st.write(f"• Chute de tension réelle: {((resistivite_cuivre * 2 * distance_cable * courant) / section_recommandee):.2f} V ({((resistivite_cuivre * 2 * distance_cable * courant) / section_recommandee / tension_cable * 100):.1f}%)")
                    st.write(f"• Capacité de courant estimée: {section_recommandee * 6:.0f} A (en air libre)")
            
            # Calculateur de disjoncteurs
            with calc_tab2:
                st.markdown("### 🔌 Dimensionnement des Disjoncteurs")
                st.info("Calculez le calibre des disjoncteurs et fusibles pour votre installation.")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Équipement à protéger**")
                    type_protection = st.selectbox("Type d'équipement", [
                        "Panneau solaire", "Batterie", "Onduleur", "Charge DC", "Charge AC"
                    ])
                    puissance_protection = st.number_input("Puissance (W)", min_value=1, max_value=50000, value=1000, step=50)
                    tension_protection = st.selectbox("Tension (V)", [12, 24, 48, 220, 380], index=1)
                    
                with col2:
                    st.markdown("**Paramètres de protection**")
                    facteur_securite = st.slider("Facteur de sécurité", min_value=1.1, max_value=2.0, value=1.25, step=0.05)
                    type_disjoncteur = st.selectbox("Type de protection", ["Disjoncteur DC", "Disjoncteur AC", "Fusible"])
                    courbe_declenchement = st.selectbox("Courbe de déclenchement", ["B (1-3 In)", "C (5-10 In)", "D (10-20 In)"], index=1)
                
                if st.button("🔍 Calculer la protection", type="primary"):
                    # Calcul du courant nominal
                    courant_nominal = puissance_protection / tension_protection
                    
                    # Application du facteur de sécurité
                    courant_protection = courant_nominal * facteur_securite
                    
                    # Calibres normalisés
                    calibres_standard = [1, 2, 4, 6, 10, 16, 20, 25, 32, 40, 50, 63, 80, 100, 125, 160, 200, 250]
                    calibre_recommande = next((c for c in calibres_standard if c >= courant_protection), 250)
                    
                    # Recommandations spécifiques par type
                    if type_protection == "Panneau solaire":
                        # Pour les panneaux, on utilise souvent 1.56 x Isc
                        facteur_pv = 1.56
                        courant_protection_pv = courant_nominal * facteur_pv
                        calibre_recommande = next((c for c in calibres_standard if c >= courant_protection_pv), 250)
                        note_specifique = f"Pour panneaux PV: facteur 1.56 x Isc appliqué"
                    elif type_protection == "Batterie":
                        note_specifique = "Protection batterie: vérifiez le courant de décharge max"
                    elif type_protection == "Onduleur":
                        note_specifique = "Protection onduleur: considérez le courant de démarrage"
                    else:
                        note_specifique = "Protection standard appliquée"
                    
                    # Affichage des résultats
                    st.success("✅ Calcul terminé")
                    
                    col_res1, col_res2, col_res3 = st.columns(3)
                    with col_res1:
                        st.metric("Courant nominal", f"{courant_nominal:.1f} A")
                        st.metric("Courant de protection", f"{courant_protection:.1f} A")
                    with col_res2:
                        st.metric("**Calibre recommandé**", f"**{calibre_recommande} A**")
                        st.metric("Type", type_disjoncteur)
                    with col_res3:
                        st.metric("Courbe", courbe_declenchement)
                        
                        # Couleur selon le calibre
                        if calibre_recommande <= 16:
                            couleur_protection = "🟢 Vert"
                        elif calibre_recommande <= 63:
                            couleur_protection = "🟡 Jaune"
                        else:
                            couleur_protection = "🔴 Rouge"
                        st.write(f"Indication: {couleur_protection}")
                    
                    st.markdown("---")
                    st.markdown("**📋 Notes importantes:**")
                    st.write(f"• {note_specifique}")
                    st.write(f"• Pouvoir de coupure: minimum {tension_protection * 10} A")
                    st.write(f"• Vérifiez la compatibilité DC/AC selon votre installation")
            
            # Calculateur d'autonomie
            with calc_tab3:
                st.markdown("### 🔋 Calculateur d'Autonomie")
                st.info("Calculez l'autonomie de votre système en fonction des batteries et de la consommation.")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Configuration batteries**")
                    capacite_batterie = st.number_input("Capacité unitaire (Ah)", min_value=1, max_value=1000, value=100, step=10)
                    tension_batterie = st.selectbox("Tension batterie (V)", [12, 24, 48], index=0)
                    nb_batteries = st.number_input("Nombre de batteries", min_value=1, max_value=20, value=4, step=1)
                    type_batterie_auto = st.selectbox("Type de batterie", ["Plomb", "AGM", "GEL", "Lithium", "Lithium HV"])
                    
                with col2:
                    st.markdown("**Consommation**")
                    consommation_jour = st.number_input("Consommation journalière (kWh)", min_value=0.1, max_value=100.0, value=5.0, step=0.1)
                    jours_autonomie_souhaite = st.number_input("Jours d'autonomie souhaités", min_value=1, max_value=10, value=3, step=1)
                    rendement_onduleur = st.slider("Rendement onduleur (%)", min_value=80, max_value=98, value=90, step=1)
                
                if st.button("🔍 Calculer l'autonomie", type="primary"):
                    # Paramètres selon le type de batterie
                    if type_batterie_auto in ["Plomb", "AGM"]:
                        decharge_max = 50  # 50% pour plomb/AGM
                        rendement_batterie = 85
                    elif type_batterie_auto == "GEL":
                        decharge_max = 60  # 60% pour GEL
                        rendement_batterie = 88
                    elif type_batterie_auto in ["Lithium", "Lithium HV"]:
                        decharge_max = 90  # 90% pour Lithium
                        rendement_batterie = 95
                    
                    # Calculs
                    capacite_totale_ah = capacite_batterie * nb_batteries
                    capacite_totale_kwh = (capacite_totale_ah * tension_batterie) / 1000
                    capacite_utilisable_kwh = capacite_totale_kwh * (decharge_max / 100)
                    
                    # Consommation réelle avec pertes
                    consommation_reelle = consommation_jour / (rendement_onduleur / 100) / (rendement_batterie / 100)
                    
                    # Autonomie réelle
                    autonomie_jours = capacite_utilisable_kwh / consommation_reelle
                    
                    # Capacité recommandée pour l'autonomie souhaitée
                    capacite_recommandee_kwh = consommation_reelle * jours_autonomie_souhaite
                    capacite_recommandee_ah = (capacite_recommandee_kwh * 1000) / tension_batterie
                    nb_batteries_recommande = math.ceil(capacite_recommandee_ah / capacite_batterie)
                    
                    # Affichage des résultats
                    st.success("✅ Calcul terminé")
                    
                    col_res1, col_res2, col_res3 = st.columns(3)
                    with col_res1:
                        st.metric("Capacité totale", f"{capacite_totale_kwh:.1f} kWh")
                        st.metric("Capacité utilisable", f"{capacite_utilisable_kwh:.1f} kWh")
                    with col_res2:
                        st.metric("**Autonomie réelle**", f"**{autonomie_jours:.1f} jours**")
                        st.metric("Consommation avec pertes", f"{consommation_reelle:.2f} kWh/j")
                    with col_res3:
                        if autonomie_jours < jours_autonomie_souhaite:
                            st.metric("Batteries recommandées", f"{nb_batteries_recommande} unités", delta=f"+{nb_batteries_recommande - nb_batteries}")
                        else:
                            st.metric("Configuration", "✅ Suffisante")
                        st.metric("Décharge max", f"{decharge_max}%")
                    
                    # Graphique d'autonomie
                    st.markdown("---")
                    st.markdown("**📊 Évolution de l'autonomie:**")
                    
                    # Simulation sur plusieurs jours
                    jours = list(range(0, int(autonomie_jours) + 1))
                    capacite_restante = [capacite_utilisable_kwh - (consommation_reelle * j) for j in jours]
                    capacite_restante = [max(0, c) for c in capacite_restante]  # Pas en dessous de 0
                    
                    chart_data = pd.DataFrame({
                        'Jour': jours,
                        'Capacité restante (kWh)': capacite_restante
                    })
                    
                    st.line_chart(chart_data.set_index('Jour'))
                    
                    # Recommandations
                    st.markdown("**💡 Recommandations:**")
                    if autonomie_jours < 2:
                        st.warning("⚠️ Autonomie faible - Augmentez la capacité des batteries")
                    elif autonomie_jours > 5:
                        st.info("ℹ️ Autonomie élevée - Configuration optimale pour sites isolés")
                    else:
                        st.success("✅ Autonomie correcte pour usage standard")
            
            # Calculateur de rentabilité solaire
            with calc_tab4:
                st.markdown("### 💰 Calculateur de Rentabilité Solaire")
                st.info("Analysez la rentabilité financière de votre installation solaire et calculez le retour sur investissement.")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**💡 Installation solaire**")
                    cout_installation = st.number_input("Coût total installation (FCFA)", min_value=500000, max_value=50000000, value=3000000, step=100000)
                    puissance_installee = st.number_input("Puissance installée (kWc)", min_value=0.5, max_value=100.0, value=5.0, step=0.5)
                    production_annuelle = st.number_input("Production annuelle estimée (kWh)", min_value=500, max_value=200000, value=8000, step=100)
                    
                with col2:
                    st.markdown("**⚡ Situation énergétique actuelle**")
                    facture_mensuelle = st.number_input("Facture électricité mensuelle (FCFA)", min_value=5000, max_value=500000, value=45000, step=5000)
                    prix_kwh_senelec = st.number_input("Prix kWh SENELEC (FCFA)", min_value=80, max_value=200, value=118, step=1)
                    augmentation_annuelle = st.slider("Augmentation annuelle prix électricité (%)", min_value=0.0, max_value=15.0, value=5.0, step=0.5)
                
                col3, col4 = st.columns(2)
                
                with col3:
                    st.markdown("**🔧 Paramètres techniques**")
                    degradation_annuelle = st.slider("Dégradation panneaux (%/an)", min_value=0.3, max_value=1.0, value=0.5, step=0.1)
                    cout_maintenance = st.number_input("Coût maintenance annuel (FCFA)", min_value=0, max_value=200000, value=50000, step=10000)
                    duree_vie = st.selectbox("Durée de vie système (années)", [15, 20, 25, 30], index=2)
                    
                with col4:
                    st.markdown("**💼 Paramètres financiers**")
                    taux_actualisation = st.slider("Taux d'actualisation (%)", min_value=3.0, max_value=12.0, value=6.0, step=0.5)
                    subvention = st.number_input("Subventions/Aides (FCFA)", min_value=0, max_value=2000000, value=0, step=50000)
                    revente_surplus = st.checkbox("Revente surplus possible", value=False)
                    if revente_surplus:
                        prix_revente = st.number_input("Prix revente (FCFA/kWh)", min_value=50, max_value=150, value=80, step=5)
                    else:
                        prix_revente = 0
                
                if st.button("📊 Calculer la rentabilité", type="primary"):
                    # Calculs de rentabilité
                    cout_net_installation = cout_installation - subvention
                    facture_annuelle_actuelle = facture_mensuelle * 12
                    
                    # Calcul des économies annuelles
                    economies_annuelles = []
                    productions_annuelles = []
                    couts_maintenance_cumules = []
                    
                    for annee in range(1, duree_vie + 1):
                        # Production dégradée
                        production_degradee = production_annuelle * ((1 - degradation_annuelle/100) ** (annee - 1))
                        productions_annuelles.append(production_degradee)
                        
                        # Prix électricité avec augmentation
                        prix_kwh_annee = prix_kwh_senelec * ((1 + augmentation_annuelle/100) ** (annee - 1))
                        
                        # Économies sur facture
                        economie_facture = min(production_degradee * prix_kwh_annee, facture_annuelle_actuelle * ((1 + augmentation_annuelle/100) ** (annee - 1)))
                        
                        # Revenus de revente (si applicable)
                        if revente_surplus:
                            surplus = max(0, production_degradee - (facture_annuelle_actuelle / prix_kwh_senelec))
                            revenus_revente = surplus * prix_revente
                        else:
                            revenus_revente = 0
                        
                        # Économies totales moins maintenance
                        economie_nette = economie_facture + revenus_revente - cout_maintenance
                        economies_annuelles.append(economie_nette)
                        couts_maintenance_cumules.append(cout_maintenance * annee)
                    
                    # Calcul du retour sur investissement
                    flux_cumules = []
                    flux_actualises = []
                    cumul = -cout_net_installation
                    
                    for i, economie in enumerate(economies_annuelles):
                        cumul += economie
                        flux_cumules.append(cumul)
                        
                        # Flux actualisé
                        flux_actualise = economie / ((1 + taux_actualisation/100) ** (i + 1))
                        flux_actualises.append(flux_actualise)
                    
                    # Temps de retour simple
                    temps_retour = None
                    for i, flux in enumerate(flux_cumules):
                        if flux >= 0:
                            temps_retour = i + 1
                            break
                    
                    # VAN (Valeur Actuelle Nette)
                    van = sum(flux_actualises) - cout_net_installation
                    
                    # TRI approximatif (méthode simplifiée)
                    economie_moyenne = sum(economies_annuelles) / len(economies_annuelles)
                    tri_approx = (economie_moyenne / cout_net_installation) * 100
                    
                    # Affichage des résultats
                    st.success("✅ Analyse de rentabilité terminée")
                    
                    # Métriques principales
                    col_res1, col_res2, col_res3, col_res4 = st.columns(4)
                    
                    with col_res1:
                        if temps_retour:
                            st.metric("⏱️ Temps de retour", f"{temps_retour} ans")
                        else:
                            st.metric("⏱️ Temps de retour", "> 25 ans")
                        
                    with col_res2:
                        st.metric("💰 VAN", f"{van:,.0f} FCFA")
                        
                    with col_res3:
                        st.metric("📈 TRI approximatif", f"{tri_approx:.1f}%")
                        
                    with col_res4:
                        economies_totales = sum(economies_annuelles)
                        st.metric("💵 Économies 25 ans", f"{economies_totales:,.0f} FCFA")
                    
                    # Graphique des flux de trésorerie
                    st.markdown("---")
                    st.markdown("**📈 Évolution des flux de trésorerie cumulés**")
                    
                    import pandas as pd
                    df_flux = pd.DataFrame({
                        'Année': list(range(0, duree_vie + 1)),
                        'Flux cumulés (FCFA)': [-cout_net_installation] + flux_cumules
                    })
                    
                    st.line_chart(df_flux.set_index('Année'))
                    
                    # Analyse détaillée
                    st.markdown("**📋 Analyse détaillée:**")
                    
                    col_analyse1, col_analyse2 = st.columns(2)
                    
                    with col_analyse1:
                        st.write(f"• **Investissement net:** {cout_net_installation:,.0f} FCFA")
                        st.write(f"• **Production 1ère année:** {production_annuelle:,.0f} kWh")
                        st.write(f"• **Économie 1ère année:** {economies_annuelles[0]:,.0f} FCFA")
                        st.write(f"• **Coût maintenance total:** {cout_maintenance * duree_vie:,.0f} FCFA")
                        
                    with col_analyse2:
                        if van > 0:
                            st.success("✅ **Projet rentable** - VAN positive")
                        else:
                            st.error("❌ **Projet non rentable** - VAN négative")
                            
                        if temps_retour and temps_retour <= 10:
                            st.success(f"✅ **Retour rapide** - {temps_retour} ans")
                        elif temps_retour and temps_retour <= 15:
                            st.warning(f"⚠️ **Retour moyen** - {temps_retour} ans")
                        else:
                            st.error("❌ **Retour trop long** - > 15 ans")
                    
                    # Recommandations
                    st.markdown("---")
                    st.markdown("**💡 Recommandations:**")
                    
                    if van > 0 and temps_retour and temps_retour <= 12:
                        st.success("🎯 **Excellent investissement** - Procédez à l'installation")
                    elif van > 0:
                        st.info("👍 **Bon investissement** - Rentable sur le long terme")
                    else:
                        st.warning("⚠️ **Investissement à reconsidérer** - Optimisez les paramètres")
                        st.write("Suggestions d'amélioration:")
                        st.write("- Rechercher des subventions supplémentaires")
                        st.write("- Optimiser la taille de l'installation")
                        st.write("- Négocier le prix d'installation")
                        st.write("- Considérer la revente du surplus")
                    
                    st.caption("💡 Analyse basée sur les données fournies - Consultez un expert pour validation")

            # Historique déplacé dans l'onglet Admin → 🕘 Historique
        
        # Onglet Gestion de Stock
        with admin_tab7:
            st.subheader("📦 Gestion de Stock")
            st.caption("Gérez vos produits, clients, factures et mouvements de stock avec synchronisation Firebase.")
            
            # Sous-onglets pour la gestion de stock
            stock_tab1, stock_tab2, stock_tab3, stock_tab4, stock_tab5 = st.tabs(["📊 Tableau de Bord", "📦 Produits", "👥 Clients", "📄 Factures", "📈 Mouvements"])
            
            # Tableau de Bord Stock
            with stock_tab1:
                # Affichage des alertes dans la sidebar
                show_stock_alerts_sidebar()
                
                # Vue d'ensemble financière moderne
                try:
                    products = get_all_products_from_firebase()
                    if products:
                        products_list = []
                        for product_id, product in products.items():
                            products_list.append({
                                'id': product_id,
                                'nom': product.get('nom', ''),
                                'categorie': product.get('categorie', ''),
                                'prix_achat': product.get('prix_achat', 0),
                                'prix_vente': product.get('prix_vente', 0),
                                'quantite': product.get('stock_actuel', product.get('quantite', 0)),
                                'stock_min': product.get('stock_minimum', product.get('stock_min', 0)),
                                'unite': product.get('unite', 'pièce')
                            })
                        df_products = pd.DataFrame(products_list)
                        create_financial_overview(df_products)
                    else:
                        st.info("Aucun produit trouvé pour l'aperçu financier")
                except Exception as e:
                    st.error(f"Erreur lors du chargement des données financières: {e}")
                
                # Section de synchronisation retirée selon demande utilisateur
                # (Import dimensionnement → stock, cohérence, synchronisation complète, actions rapides)
                # La gestion des graphiques et alertes stock reste disponible ci-dessous.
                st.markdown("---")
                
                # Graphiques avancés de stock
                try:
                    products = get_all_products_from_firebase()
                    if products:
                        products_list = []
                        for product_id, product in products.items():
                            products_list.append({
                                'id': product_id,
                                'nom': product.get('nom', ''),
                                'categorie': product.get('categorie', ''),
                                'prix_achat': product.get('prix_achat', 0),
                                'prix_vente': product.get('prix_vente', 0),
                                'quantite': product.get('stock_actuel', product.get('quantite', 0)),
                                'stock_min': product.get('stock_minimum', product.get('stock_min', 0)),
                                'unite': product.get('unite', 'pièce')
                            })
                        df_products = pd.DataFrame(products_list)
                        create_advanced_stock_chart(df_products)
                    else:
                        st.info("Aucun produit trouvé pour les graphiques")
                except Exception as e:
                    st.error(f"Erreur lors du chargement des graphiques: {e}")
                
                st.markdown("---")
                
                # Alertes stock avec cartes modernes
                try:
                    products = get_all_products_from_firebase()
                    if products:
                        low_stock_products = []
                        for product_id, product in products.items():
                            current_stock = product.get('stock_actuel', 0)
                            min_stock = product.get('stock_minimum', 0)
                            if current_stock <= min_stock:
                                low_stock_products.append({
                                    'product_name': product.get('nom', ''),
                                    'current_stock': current_stock,
                                    'min_stock': min_stock,
                                    'category': product.get('categorie', ''),
                                    'urgency': 'critical' if current_stock == 0 else 'warning'
                                })
                        
                        if low_stock_products:
                            st.subheader("⚠️ Alertes Stock Critiques")
                            for alert in low_stock_products:
                                create_stock_alert_card(alert)
                        else:
                            st.success("✅ Aucune alerte stock critique")
                    
                except Exception as e:
                    st.error(f"❌ Erreur lors du chargement des données: {e}")
            
            # Gestion des Produits
            with stock_tab2:
                # Utiliser le tableau interactif moderne pour la gestion des produits
                try:
                    products = get_all_products_from_firebase()
                    if products:
                        products_list = []
                        for product_id, product in products.items():
                            products_list.append({
                                'id': product_id,
                                'nom': product.get('nom', ''),
                                'categorie': product.get('categorie', ''),
                                'prix_achat': product.get('prix_achat', 0),
                                'prix_vente': product.get('prix_vente', 0),
                                'quantite': product.get('stock_actuel', product.get('quantite', 0)),
                                'stock_min': product.get('stock_minimum', product.get('stock_min', 0)),
                                'unite': product.get('unite', 'pièce')
                            })
                        df_products = pd.DataFrame(products_list)
                        create_interactive_product_table(df_products)
                    else:
                        st.info("Aucun produit trouvé")
                except Exception as e:
                    st.error(f"Erreur lors du chargement des produits: {e}")
                
                st.markdown("---")
                
                # Formulaire d'ajout de produit amélioré
                with st.expander("➕ Ajouter un Nouveau Produit", expanded=False):
                    with st.form("add_product"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            nom = st.text_input("Nom du produit")
                            categorie = st.selectbox("Catégorie", [
                                "Panneaux Solaires", "Batteries", "Onduleurs", 
                                "Régulateurs", "Accessoires", "Câbles", "Autres"
                            ])
                            prix_achat = st.number_input("Prix d'achat (FCFA)", min_value=0, value=0, step=1000)
                            prix_vente = st.number_input("Prix de vente (FCFA)", min_value=0, value=0, step=1000)
                        
                        with col2:
                            stock_initial = st.number_input("Stock initial", min_value=0, value=0, step=1)
                            stock_minimum = st.number_input("Stock minimum", min_value=0, value=5, step=1)
                            unite = st.selectbox("Unité", ["pièce", "mètre", "kg", "litre", "lot"])
                            description = st.text_area("Description (optionnel)", height=100)
                        
                        # Caractéristiques techniques par catégorie
                        specifications = {}
                        if categorie == "Panneaux Solaires":
                            st.subheader("Caractéristiques techniques du panneau")
                            spec_col1, spec_col2, spec_col3 = st.columns(3)
                            with spec_col1:
                                puissance_w = st.number_input("Puissance (W)", min_value=0, value=0, step=10, help="Puissance nominale en watts")
                                voltage_v = st.number_input("Voltage (V)", min_value=0, value=0, step=1)
                            with spec_col2:
                                type_panneau = st.selectbox("Type", ["Monocristallin", "Polycristallin", "PERC", "Thin-film"], index=0)
                                courant_a = st.number_input("Courant (A)", min_value=0.0, value=0.0, step=0.1)
                            with spec_col3:
                                dimensions = st.text_input("Dimensions (mm)", value="")
                            specifications = {
                                "puissance": puissance_w,
                                "voltage": voltage_v,
                                "type": type_panneau,
                                "courant": courant_a,
                                "dimensions": dimensions,
                            }
                        elif categorie == "Batteries":
                            st.subheader("Caractéristiques techniques de la batterie")
                            spec_col1, spec_col2, spec_col3 = st.columns(3)
                            with spec_col1:
                                capacite_ah = st.number_input("Capacité (Ah)", min_value=0, value=0, step=10)
                                voltage_v = st.number_input("Voltage (V)", min_value=0, value=12, step=1)
                            with spec_col2:
                                cycles = st.number_input("Cycles", min_value=0, value=0, step=100)
                                decharge_max = st.number_input("Décharge max (%)", min_value=0, max_value=100, value=50, step=1)
                            with spec_col3:
                                type_batterie = st.selectbox("Type", ["Gel", "AGM", "Lithium", "Plomb"], index=0)
                            specifications = {
                                "capacite": capacite_ah,
                                "voltage": voltage_v,
                                "cycles": cycles,
                                "decharge_max": decharge_max,
                                "type": type_batterie,
                            }
                        elif categorie == "Onduleurs":
                            st.subheader("Caractéristiques techniques de l'onduleur")
                            spec_col1, spec_col2, spec_col3 = st.columns(3)
                            with spec_col1:
                                puissance_w = st.number_input("Puissance (W)", min_value=0, value=0, step=100)
                                voltage_v = st.number_input("Voltage (V)", min_value=0, value=220, step=10)
                            with spec_col2:
                                phase = st.selectbox("Phase", ["Monophasé", "Triphasé"], index=0)
                                mppt = st.checkbox("MPPT")
                            with spec_col3:
                                efficacite = st.number_input("Efficacité (%)", min_value=0, max_value=100, value=90, step=1)
                            specifications = {
                                "puissance": puissance_w,
                                "voltage": voltage_v,
                                "phase": phase,
                                "mppt": mppt,
                                "efficacite": efficacite,
                            }
                        elif categorie == "Régulateurs":
                            st.subheader("Caractéristiques techniques du régulateur")
                            spec_col1, spec_col2, spec_col3 = st.columns(3)
                            with spec_col1:
                                voltage_v = st.number_input("Voltage (V)", min_value=0, value=12, step=1)
                            with spec_col2:
                                amperage_a = st.number_input("Ampérage (A)", min_value=0, value=0, step=1)
                                mppt = st.checkbox("MPPT")
                            with spec_col3:
                                type_reg = st.selectbox("Type", ["PWM", "MPPT"], index=1)
                            specifications = {
                                "voltage": voltage_v,
                                "amperage": amperage_a,
                                "mppt": mppt,
                                "type": type_reg,
                            }
                        
                        if st.form_submit_button("➕ Ajouter le Produit", type="primary"):
                            if nom and categorie:
                                try:
                                    product_data = {
                                        'nom': nom,
                                        'categorie': categorie,
                                        'prix_achat': prix_achat,
                                        'prix_vente': prix_vente,
                                        'stock_actuel': stock_initial,
                                        'stock_minimum': stock_minimum,
                                        'unite': unite,
                                        'description': description,
                                        'specifications': specifications,
                                        'source': 'manual',
                                        'date_creation': pd.Timestamp.now().isoformat()
                                    }
                                    
                                    success = save_product_to_firebase(product_data)
                                    if success:
                                        st.success("✅ Produit ajouté avec succès!")
                                        # Synchroniser vers Gestion des Prix des Équipements (catégories principales)
                                        try:
                                            prix_data = get_equipment_prices() or {}
                                            cat_map = {
                                                "Panneaux Solaires": "panneaux",
                                                "Batteries": "batteries",
                                                "Onduleurs": "onduleurs",
                                                "Régulateurs": "regulateurs",
                                            }
                                            if categorie in cat_map:
                                                cat = cat_map[categorie]
                                                if cat not in prix_data:
                                                    prix_data[cat] = {}
                                                key = re.sub(r"\s+", "_", nom.strip().lower())
                                                entry = prix_data[cat].get(key, {})
                                                entry.update({
                                                    'prix': float(prix_vente) if prix_vente is not None else 0,
                                                    'puissance': specifications.get('puissance', entry.get('puissance')),
                                                    'voltage': specifications.get('voltage', entry.get('voltage')),
                                                    'type': specifications.get('type', entry.get('type')),
                                                    'capacite': specifications.get('capacite', entry.get('capacite')),
                                                    'cycles': specifications.get('cycles', entry.get('cycles')),
                                                    'decharge_max': specifications.get('decharge_max', entry.get('decharge_max')),
                                                    'phase': specifications.get('phase', entry.get('phase')),
                                                    'mppt': specifications.get('mppt', entry.get('mppt')),
                                                    'amperage': specifications.get('amperage', entry.get('amperage')),
                                                    'courant': specifications.get('courant', entry.get('courant')),
                                                    'dimensions': specifications.get('dimensions', entry.get('dimensions')),
                                                    'efficacite': specifications.get('efficacite', entry.get('efficacite')),
                                                })
                                                prix_data[cat][key] = entry
                                                save_equipment_prices(prix_data)
                                                clear_prices_cache()
                                                st.info("🔄 Caractéristiques synchronisées avec Gestion des Prix des Équipements")
                                        except Exception as sync_e:
                                            st.warning(f"⚠️ Synchronisation des prix échouée: {sync_e}")
                                        clear_stock_cache()
                                        st.rerun()
                                    else:
                                        st.error("❌ Erreur lors de l'ajout")
                                except Exception as e:
                                    st.error(f"❌ Erreur: {e}")
                            else:
                                st.error("❌ Veuillez remplir tous les champs obligatoires")
                
                # Interface d'importation Excel avec mappage de colonnes
                with st.expander("📊 Importer des Produits depuis Excel", expanded=False):
                    st.markdown("### 📥 Importation Excel avec Mappage de Colonnes")
                    
                    uploaded_file = st.file_uploader(
                        "Choisir un fichier Excel (.xlsx, .xls)", 
                        type=['xlsx', 'xls'],
                        help="Téléchargez votre fichier Excel contenant la liste des produits"
                    )
                    
                    if uploaded_file is not None:
                        try:
                            # Lire le fichier Excel
                            df_excel = pd.read_excel(uploaded_file)
                            
                            st.success(f"✅ Fichier chargé avec succès! {len(df_excel)} lignes trouvées.")
                            
                            # Aperçu des données
                            st.markdown("**📋 Aperçu des données:**")
                            st.dataframe(df_excel.head(), use_container_width=True)
                            
                            # Mappage des colonnes
                            st.markdown("### 🔗 Mappage des Colonnes")
                            st.info("Associez les colonnes de votre fichier Excel aux champs du système:")
                            
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                nom_col = st.selectbox("Nom du produit", [""] + list(df_excel.columns), key="nom_mapping")
                                categorie_col = st.selectbox("Catégorie", [""] + list(df_excel.columns), key="cat_mapping")
                                prix_achat_col = st.selectbox("Prix d'achat", [""] + list(df_excel.columns), key="achat_mapping")
                                prix_vente_col = st.selectbox("Prix de vente", [""] + list(df_excel.columns), key="vente_mapping")
                            
                            with col2:
                                stock_col = st.selectbox("Stock initial", [""] + list(df_excel.columns), key="stock_mapping")
                                stock_min_col = st.selectbox("Stock minimum", [""] + list(df_excel.columns), key="min_mapping")
                                unite_col = st.selectbox("Unité", [""] + list(df_excel.columns), key="unite_mapping")
                                description_col = st.selectbox("Description (optionnel)", [""] + list(df_excel.columns), key="desc_mapping")
                            
                            # Validation et importation
                            if nom_col:
                                st.markdown("### ✅ Validation et Importation")
                                
                                # Aperçu du mappage
                                preview_data = []
                                for i, row in df_excel.head(3).iterrows():
                                    preview_data.append({
                                        'Nom': row[nom_col] if nom_col else '',
                                        'Catégorie': row[categorie_col] if categorie_col else 'Autres',
                                        'Prix Achat': row[prix_achat_col] if prix_achat_col else 0,
                                        'Prix Vente': row[prix_vente_col] if prix_vente_col else 0,
                                        'Stock': row[stock_col] if stock_col else 0,
                                        'Stock Min': row[stock_min_col] if stock_min_col else 0,
                                        'Unité': row[unite_col] if unite_col else 'pièce',
                                        'Description': row[description_col] if description_col else ''
                                    })
                                
                                st.markdown("**📋 Aperçu du mappage (3 premières lignes):**")
                                st.dataframe(pd.DataFrame(preview_data), use_container_width=True)
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    update_on_duplicate = st.checkbox("Mettre à jour si doublon (par nom)", value=True)
                                    if st.button("📥 Importer tous les produits", type="primary"):
                                        progress_bar = st.progress(0)
                                        success_count = 0
                                        error_count = 0
                                        
                                        # Charger les produits existants pour détecter les doublons par nom
                                        existing_by_name = {}
                                        try:
                                            from firebase_config import get_all_products_from_firebase, update_product_in_firebase
                                            existing_products = get_all_products_from_firebase()
                                            if existing_products:
                                                for pid, pdata in existing_products.items():
                                                    name_key = str(pdata.get('nom', '')).strip().lower()
                                                    if name_key:
                                                        existing_by_name[name_key] = pid
                                        except Exception as e:
                                            st.warning(f"Impossible de charger les produits existants pour la mise à jour: {e}")
                                        
                                        for i, row in df_excel.iterrows():
                                            try:
                                                product_data = {
                                                    'nom': str(row[nom_col]) if nom_col and pd.notna(row[nom_col]) else f"Produit_{i+1}",
                                                    'categorie': str(row[categorie_col]) if categorie_col and pd.notna(row[categorie_col]) else 'Autres',
                                                    'prix_achat': float(row[prix_achat_col]) if prix_achat_col and pd.notna(row[prix_achat_col]) else 0,
                                                    'prix_vente': float(row[prix_vente_col]) if prix_vente_col and pd.notna(row[prix_vente_col]) else 0,
                                                    'stock_actuel': int(row[stock_col]) if stock_col and pd.notna(row[stock_col]) else 0,
                                                    'stock_minimum': int(row[stock_min_col]) if stock_min_col and pd.notna(row[stock_min_col]) else 0,
                                                    'unite': str(row[unite_col]) if unite_col and pd.notna(row[unite_col]) else 'pièce',
                                                    'description': str(row[description_col]) if description_col and pd.notna(row[description_col]) else '',
                                                    'date_creation': pd.Timestamp.now().isoformat(),
                                                    'source': 'import_excel'
                                                }
                                                
                                                name_key = str(product_data.get('nom', '')).strip().lower()
                                                
                                                # Mettre à jour si doublon, sinon créer
                                                if update_on_duplicate and name_key in existing_by_name:
                                                    try:
                                                        if update_product_in_firebase(existing_by_name[name_key], product_data):
                                                            success_count += 1
                                                        else:
                                                            error_count += 1
                                                    except Exception as e:
                                                        error_count += 1
                                                        st.error(f"Erreur mise à jour ligne {i+1}: {e}")
                                                else:
                                                    if save_product_to_firebase(product_data):
                                                        success_count += 1
                                                    else:
                                                        error_count += 1
                                                
                                            except Exception as e:
                                                error_count += 1
                                                st.error(f"Erreur ligne {i+1}: {e}")
                                            
                                            progress_bar.progress((i + 1) / len(df_excel))
                                        
                                        if success_count > 0:
                                            st.success(f"✅ {success_count} produits traités avec succès!")
                                            clear_stock_cache()
                                            if error_count > 0:
                                                st.warning(f"⚠️ {error_count} erreurs lors de l'importation")
                                            st.rerun()
                                        else:
                                            st.error("❌ Aucun produit n'a pu être importé")
                                
                                with col2:
                                    if st.button("🔄 Réinitialiser le mappage"):
                                        st.rerun()
                            
                        except Exception as e:
                            st.error(f"❌ Erreur lors de la lecture du fichier: {e}")
                
                # Synchronisation des produits de dimensionnement
                with st.expander("🔄 Synchroniser les Produits de Dimensionnement", expanded=False):
                    st.markdown("### ⚡ Ajouter les Produits du Dimensionnement au Stock")
                    st.info("Cette fonction permet d'ajouter automatiquement tous les produits utilisés dans le dimensionnement vers le gestionnaire de stock.")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**📦 Produits disponibles dans le dimensionnement:**")
                        
                        # Compter les produits par catégorie
                        panneaux_count = len(PRIX_EQUIPEMENTS["panneaux"])
                        batteries_count = len(PRIX_EQUIPEMENTS["batteries"])
                        onduleurs_count = len(PRIX_EQUIPEMENTS["onduleurs"])
                        regulateurs_count = len(PRIX_EQUIPEMENTS["regulateurs"])
                        
                        st.metric("Panneaux Solaires", panneaux_count)
                        st.metric("Batteries", batteries_count)
                        st.metric("Onduleurs", onduleurs_count)
                        st.metric("Régulateurs", regulateurs_count)
                        
                        total_products = panneaux_count + batteries_count + onduleurs_count + regulateurs_count
                        st.metric("**Total**", total_products)
                    
                    with col2:
                        st.markdown("**⚙️ Options de synchronisation:**")
                        
                        sync_mode = st.radio(
                            "Mode de synchronisation:",
                            ["Ajouter seulement les nouveaux produits", "Mettre à jour les prix existants", "Synchronisation complète"],
                            help="Choisissez comment traiter les produits déjà existants"
                        )
                        
                        default_stock = st.number_input("Stock initial par défaut", min_value=0, value=10, step=1)
                        default_stock_min = st.number_input("Stock minimum par défaut", min_value=0, value=2, step=1)
                        
                        if st.button("🔄 Synchroniser les Produits", type="primary"):
                            progress_bar = st.progress(0)
                            success_count = 0
                            updated_count = 0
                            error_count = 0
                            
                            # Obtenir les produits existants
                            existing_products = get_all_products_from_firebase() or {}
                            existing_names = {prod.get('nom', '').lower(): prod_id for prod_id, prod in existing_products.items()}
                            
                            all_equipment = []
                            
                            # Préparer tous les équipements
                            for category, items in PRIX_EQUIPEMENTS.items():
                                category_name = {
                                    "panneaux": "Panneaux Solaires",
                                    "batteries": "Batteries", 
                                    "onduleurs": "Onduleurs",
                                    "regulateurs": "Régulateurs"
                                }.get(category, "Autres")
                                
                                for name, specs in items.items():
                                    all_equipment.append((category_name, name, specs))
                            
                            for i, (category, name, specs) in enumerate(all_equipment):
                                try:
                                    # Calculer prix d'achat estimé (70% du prix de vente)
                                    prix_vente = specs.get('prix', 0)
                                    prix_achat_estime = int(prix_vente * 0.7)
                                    
                                    product_data = {
                                        'nom': name,
                                        'categorie': category,
                                        'prix_achat': prix_achat_estime,
                                        'prix_vente': prix_vente,
                                        'stock_actuel': default_stock,
                                        'stock_minimum': default_stock_min,
                                        'unite': 'pièce',
                                        'description': f"Synchronisé depuis le dimensionnement - {category}",
                                        'date_creation': pd.Timestamp.now().isoformat(),
                                        'source': 'dimensionnement',
                                        'specifications': specs
                                    }
                                    
                                    # Vérifier si le produit existe déjà
                                    existing_id = existing_names.get(name.lower())
                                    
                                    if existing_id:
                                        if sync_mode in ["Mettre à jour les prix existants", "Synchronisation complète"]:
                                            # Mettre à jour le produit existant
                                            if update_product_in_firebase(existing_id, product_data):
                                                updated_count += 1
                                            else:
                                                error_count += 1
                                        # Sinon, ignorer (produit déjà existant)
                                    else:
                                        # Ajouter nouveau produit
                                        if save_product_to_firebase(product_data):
                                            success_count += 1
                                        else:
                                            error_count += 1
                                            
                                except Exception as e:
                                    error_count += 1
                                    st.error(f"Erreur pour {name}: {e}")
                                
                                progress_bar.progress((i + 1) / len(all_equipment))
                            
                            # Afficher les résultats
                            if success_count > 0 or updated_count > 0:
                                st.success(f"✅ Synchronisation terminée!")
                                if success_count > 0:
                                    st.info(f"📦 {success_count} nouveaux produits ajoutés")
                                if updated_count > 0:
                                    st.info(f"🔄 {updated_count} produits mis à jour")
                                if error_count > 0:
                                    st.warning(f"⚠️ {error_count} erreurs")
                                
                                clear_stock_cache()
                                st.rerun()
                            else:
                                if error_count > 0:
                                    st.error(f"❌ {error_count} erreurs lors de la synchronisation")
                                else:
                                    st.info("ℹ️ Aucune modification nécessaire")
                
                # Interface d'édition rapide des produits existants
                with st.expander("✏️ Édition Rapide des Produits", expanded=False):
                    st.markdown("### 🛠️ Modifier les Produits Existants")
                    
                    try:
                        products = get_all_products_from_firebase()
                        if products:
                            # Sélection du produit à modifier
                            product_names = {f"{prod.get('nom', '')} ({prod.get('categorie', '')})": prod_id 
                                           for prod_id, prod in products.items()}
                            
                            selected_product_display = st.selectbox(
                                "Sélectionner un produit à modifier:",
                                [""] + list(product_names.keys())
                            )
                            
                            if selected_product_display:
                                selected_product_id = product_names[selected_product_display]
                                selected_product = products[selected_product_id]
                                
                                st.markdown(f"**Modification de:** {selected_product.get('nom', '')}")
                                
                                with st.form("edit_product"):
                                    col1, col2 = st.columns(2)
                                    
                                    with col1:
                                        new_nom = st.text_input("Nom", value=selected_product.get('nom', ''))
                                        new_categorie = st.selectbox("Catégorie", [
                                            "Panneaux Solaires", "Batteries", "Onduleurs", 
                                            "Régulateurs", "Accessoires", "Câbles", "Autres"
                                        ], index=["Panneaux Solaires", "Batteries", "Onduleurs", 
                                                "Régulateurs", "Accessoires", "Câbles", "Autres"].index(
                                                selected_product.get('categorie', 'Autres')))
                                        new_prix_achat = st.number_input("Prix d'achat (FCFA)", 
                                                                        value=float(selected_product.get('prix_achat', 0)), 
                                                                        min_value=0.0, step=1000.0)
                                        new_prix_vente = st.number_input("Prix de vente (FCFA)", 
                                                                        value=float(selected_product.get('prix_vente', 0)), 
                                                                        min_value=0.0, step=1000.0)
                                    
                                    with col2:
                                        new_stock = st.number_input("Stock actuel", 
                                                                   value=int(selected_product.get('stock_actuel', 0)), 
                                                                   min_value=0, step=1)
                                        new_stock_min = st.number_input("Stock minimum", 
                                                                       value=int(selected_product.get('stock_minimum', 0)), 
                                                                       min_value=0, step=1)
                                        new_unite = st.selectbox("Unité", ["pièce", "mètre", "kg", "litre", "lot"],
                                                                index=["pièce", "mètre", "kg", "litre", "lot"].index(
                                                                    selected_product.get('unite', 'pièce')))
                                        new_description = st.text_area("Description", 
                                                                      value=selected_product.get('description', ''), 
                                                                      height=100)
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        if st.form_submit_button("💾 Sauvegarder les Modifications", type="primary"):
                                            try:
                                                updated_data = {
                                                    'nom': new_nom,
                                                    'categorie': new_categorie,
                                                    'prix_achat': new_prix_achat,
                                                    'prix_vente': new_prix_vente,
                                                    'stock_actuel': new_stock,
                                                    'stock_minimum': new_stock_min,
                                                    'unite': new_unite,
                                                    'description': new_description,
                                                    'date_modification': pd.Timestamp.now().isoformat()
                                                }
                                                
                                                if update_product_in_firebase(selected_product_id, updated_data):
                                                    st.success("✅ Produit modifié avec succès!")
                                                    clear_stock_cache()
                                                    st.rerun()
                                                else:
                                                    st.error("❌ Erreur lors de la modification")
                                            except Exception as e:
                                                st.error(f"❌ Erreur: {e}")
                                    
                                    with col2:
                                        if st.form_submit_button("🗑️ Supprimer le Produit", type="secondary"):
                                            if delete_product_from_firebase(selected_product_id):
                                                # Suppression croisée: retirer l’article du dimensionnement si présent
                                                delete_dimensionnement_article_if_exists(
                                                    selected_product.get('nom', ''),
                                                    selected_product.get('categorie', '')
                                                )
                                                st.success("✅ Produit supprimé avec succès!")
                                                clear_stock_cache()
                                                st.rerun()
                                            else:
                                                st.error("❌ Erreur lors de la suppression")
                        else:
                            st.info("Aucun produit disponible pour modification")
                    except Exception as e:
                        st.error(f"Erreur lors du chargement des produits: {e}")
            
            # Gestion des Clients
            with stock_tab3:
                st.subheader("👥 Gestion des Clients")
                
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.markdown("**Liste des Clients**")
                    try:
                        clients = get_all_clients_from_firebase()
                        if clients:
                            clients_list = []
                            for client_id, client in clients.items():
                                clients_list.append({
                                    'ID': client_id,
                                    'Nom': client.get('nom', ''),
                                    'Téléphone': client.get('telephone', ''),
                                    'Email': client.get('email', ''),
                                    'Adresse': client.get('adresse', '')
                                })
                            
                            df_clients = pd.DataFrame(clients_list)
                            st.dataframe(df_clients, use_container_width=True)
                        else:
                            st.info("Aucun client trouvé")
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")
                
                with col2:
                    st.markdown("**Ajouter un Client**")
                    with st.form("add_client"):
                        nom = st.text_input("Nom du client")
                        telephone = st.text_input("Téléphone")
                        email = st.text_input("Email")
                        adresse = st.text_area("Adresse")
                        
                        if st.form_submit_button("➕ Ajouter"):
                            if nom:
                                try:
                                    client_data = {
                                        'nom': nom,
                                        'telephone': telephone,
                                        'email': email,
                                        'adresse': adresse,
                                        'date_creation': pd.Timestamp.now().isoformat()
                                    }
                                    
                                    success = save_client_to_firebase(client_data)
                                    if success:
                                        st.success("✅ Client ajouté avec succès!")
                                        clear_stock_cache()
                                        st.rerun()
                                    else:
                                        st.error("❌ Erreur lors de l'ajout")
                                except Exception as e:
                                    st.error(f"❌ Erreur: {e}")
                            else:
                                st.error("❌ Le nom du client est obligatoire")
            
            # Gestion des Factures - Éditeur Complet
            with stock_tab4:
                # Utiliser l'éditeur de factures complet
                show_invoice_editor()
            
            # Mouvements de Stock
            with stock_tab5:
                # Timeline moderne des mouvements
                try:
                    movements = get_stock_movements_from_firebase()
                    if movements:
                        movements_list = []
                        for movement_id, movement in movements.items():
                            movements_list.append({
                                'id': movement_id,
                                'produit_nom': movement.get('produit_nom', ''),
                                'type': movement.get('type', ''),
                                'quantite': movement.get('quantite', 0),
                                'date': movement.get('date', ''),
                                'motif': movement.get('motif', ''),
                                'stock_apres': movement.get('stock_apres', 0)
                            })
                        df_movements = pd.DataFrame(movements_list)
                        create_movement_timeline(df_movements)
                    else:
                        st.info("Aucun mouvement trouvé")
                except Exception as e:
                    st.error(f"Erreur lors du chargement des mouvements: {e}")
                
                st.markdown("---")
                
                # Formulaire d'enregistrement de mouvement amélioré
                with st.expander("📝 Enregistrer un Nouveau Mouvement", expanded=False):
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.markdown("**Historique Récent**")
                        try:
                            movements = get_stock_movements_from_firebase()
                            if movements:
                                # Afficher les 5 derniers mouvements
                                movements_list = []
                                for movement_id, movement in movements.items():
                                    movements_list.append({
                                        'Produit': movement.get('produit_nom', ''),
                                        'Type': movement.get('type', ''),
                                        'Quantité': movement.get('quantite', 0),
                                        'Date': movement.get('date', ''),
                                        'Motif': movement.get('motif', ''),
                                        'Stock Après': movement.get('stock_apres', 0)
                                    })
                                
                                # Trier par date et prendre les 5 derniers
                                df_movements = pd.DataFrame(movements_list)
                                df_movements = df_movements.sort_values('Date', ascending=False).head(5)
                                st.dataframe(df_movements, use_container_width=True, hide_index=True)
                            else:
                                st.info("Aucun mouvement trouvé")
                        except Exception as e:
                            st.error(f"❌ Erreur: {e}")
                    
                    with col2:
                        st.markdown("**Nouveau Mouvement**")
                    try:
                        products = get_all_products_from_firebase()
                        if products:
                            with st.form("add_movement"):
                                # Sélection du produit
                                product_options = {f"{product.get('nom', '')} (Stock: {product.get('stock_actuel', 0)})": product_id 
                                                 for product_id, product in products.items()}
                                selected_product_display = st.selectbox("Produit", list(product_options.keys()))
                                selected_product_id = product_options[selected_product_display] if selected_product_display else None
                                
                                type_mouvement = st.selectbox("Type de mouvement", ["Entrée", "Sortie"])
                                quantite = st.number_input("Quantité", min_value=1, step=1)
                                motif = st.text_input("Motif", placeholder="Ex: Achat, Vente, Ajustement...")
                                
                                if st.form_submit_button("📝 Enregistrer"):
                                    if selected_product_id and quantite > 0 and motif:
                                        try:
                                            # Récupérer les infos du produit
                                            product = products[selected_product_id]
                                            stock_actuel = product.get('stock_actuel', 0)
                                            
                                            # Vérifier le stock pour les sorties
                                            if type_mouvement == "Sortie" and quantite > stock_actuel:
                                                st.error(f"❌ Stock insuffisant! Stock actuel: {stock_actuel}")
                                            else:
                                                # Calculer le nouveau stock
                                                if type_mouvement == "Entrée":
                                                    nouveau_stock = stock_actuel + quantite
                                                else:
                                                    nouveau_stock = stock_actuel - quantite
                                                
                                                # Enregistrer le mouvement
                                                movement_data = {
                                                    'produit_id': selected_product_id,
                                                    'produit_nom': product.get('nom', ''),
                                                    'type': type_mouvement,
                                                    'quantite': quantite,
                                                    'stock_avant': stock_actuel,
                                                    'stock_apres': nouveau_stock,
                                                    'motif': motif,
                                                    'date': pd.Timestamp.now().isoformat(),
                                                    'utilisateur': st.session_state.get('user_email', '')
                                                }
                                                
                                                success_movement = save_stock_movement_to_firebase(movement_data)
                                                
                                                # Mettre à jour le stock du produit
                                                product_update = {**product, 'stock_actuel': nouveau_stock}
                                                success_product = update_product_in_firebase(selected_product_id, product_update)
                                                
                                                if success_movement and success_product:
                                                    st.success("✅ Mouvement enregistré avec succès!")
                                                    clear_stock_cache()
                                                    st.rerun()
                                                else:
                                                    st.error("❌ Erreur lors de l'enregistrement")
                                        except Exception as e:
                                            st.error(f"❌ Erreur: {e}")
                                    else:
                                        st.error("❌ Veuillez remplir tous les champs")
                        else:
                            st.info("Aucun produit disponible. Ajoutez d'abord des produits.")
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")

st.markdown("---")

# Footer réorganisé
# Logo centré en haut
col1, col2, col3 = st.columns([2, 1, 2])
with col2:
    try:
        st.image("logo-solaire.svg", width=150)
    except FileNotFoundError:
        st.markdown("<h3 style='text-align: center;'>☀️</h3>", unsafe_allow_html=True)

# Contenu textuel centré en dessous
st.markdown("""
<div style='text-align: center; color: #666; padding: 10px 20px;'>
    <p><strong>☀️ Application de Dimensionnement Solaire - Sénégal</strong></p>
    <p style='color: #4CAF50; font-weight: bold; margin: 5px 0;'>🥇 Premier outil de dimensionnement solaire en ligne au Sénégal</p>
    <p>🌍 Développé par la Team Mo.TL (773591509).</p>
    <p>📞 Pour acheter vos équipements : <a href='https://energiesolairesenegal.com' target='_blank'>energiesolairesenegal.com</a></p>
    <p style='font-size: 0.9em; margin-top: 10px;'>
        💡 <b>Conseil :</b> Consultez toujours un professionnel certifié pour l'installation<br>
        ⚡ Prix indicatifs - Demandez un devis personnalisé pour votre projet
    </p>
</div>
""", unsafe_allow_html=True)
