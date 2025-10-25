"""
Éditeur de factures complet avec génération PDF
"""

import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.pdfgen import canvas
import io
import base64
from firebase_config import (
    save_invoice_to_firebase,
    get_all_invoices_from_firebase,
    get_all_clients_from_firebase,
    get_all_products_from_firebase,
    update_product_in_firebase,
    save_stock_movement_to_firebase
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class InvoiceEditor:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
    
    def setup_custom_styles(self):
        """Configuration des styles personnalisés"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.darkblue,
            alignment=1  # Center
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceAfter=12,
            textColor=colors.darkblue
        ))
    
    def create_invoice_interface(self):
        """Interface principale de création de facture"""
        st.subheader("📄 Éditeur de Factures")
        
        # Onglets pour différentes actions
        tab1, tab2, tab3 = st.tabs(["Nouvelle Facture", "Factures Existantes", "Modèles"])
        
        with tab1:
            self.create_new_invoice()
        
        with tab2:
            self.display_existing_invoices()
        
        with tab3:
            self.invoice_templates()
    
    def create_new_invoice(self):
        """Interface de création d'une nouvelle facture"""
        st.write("### Créer une nouvelle facture")
        
        # Informations de base
        col1, col2 = st.columns(2)
        
        with col1:
            # Sélection du client
            clients = get_all_clients_from_firebase()
            client_options = ["Nouveau client..."] + [f"{c.get('nom', '')} - {c.get('email', '')}" for c in clients]
            selected_client = st.selectbox("Client", client_options)
            
            if selected_client == "Nouveau client...":
                with st.expander("Informations du nouveau client"):
                    new_client_name = st.text_input("Nom du client")
                    new_client_email = st.text_input("Email")
                    new_client_phone = st.text_input("Téléphone")
                    new_client_address = st.text_area("Adresse")
                    
                    if st.button("Ajouter ce client"):
                        from firebase_config import save_client_to_firebase
                        client_data = {
                            "nom": new_client_name,
                            "email": new_client_email,
                            "telephone": new_client_phone,
                            "adresse": new_client_address
                        }
                        if save_client_to_firebase(client_data):
                            st.success("Client ajouté avec succès!")
                            st.rerun()
            
            # Numéro de facture
            invoice_number = st.text_input("Numéro de facture", value=f"FACT-{datetime.now().strftime('%Y%m%d')}-{len(get_all_invoices_from_firebase()) + 1:03d}")
            
            # Date de facture
            invoice_date = st.date_input("Date de facture", value=datetime.now())
            
            # Date d'échéance
            due_date = st.date_input("Date d'échéance", value=datetime.now() + timedelta(days=30))
        
        with col2:
            # Type de document
            doc_type = st.selectbox("Type de document", ["Facture", "Devis", "Proforma"])
            
            # Statut
            status = st.selectbox("Statut", ["Brouillon", "Envoyé", "Validé", "Payé", "En retard", "Annulé"])
            
            # Conditions de paiement
            payment_terms = st.text_area("Conditions de paiement", 
                                       value="Paiement à 30 jours.")
        
        # Section produits/services
        st.write("### Produits et Services")
        
        # Initialiser les lignes de facture dans la session
        if 'invoice_lines' not in st.session_state:
            st.session_state.invoice_lines = []
        
        # Ajouter une ligne
        with st.expander("Ajouter un produit/service"):
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                # Sélection produit ou service personnalisé
                products = get_all_products_from_firebase()
                # Convertir le dictionnaire en liste pour l'interface
                products_list = list(products.values()) if products else []
                product_options = ["Service personnalisé..."] + [p.get('nom', '') for p in products_list]
                selected_product = st.selectbox("Produit/Service", product_options, key="new_product")
            
            with col2:
                if selected_product == "Service personnalisé...":
                    description = st.text_input("Description", key="new_desc")
                    unit_price = st.number_input("Prix unitaire (FCFA)", min_value=0.0, step=1.0, key="new_price")
                else:
                    product_data = next((p for p in products_list if p.get('nom') == selected_product), None)
                    description = selected_product
                    unit_price = product_data.get('prix_vente', 0) if product_data else 0
                    st.write(f"Prix: {unit_price:,.0f} FCFA")
            
            with col3:
                quantity = st.number_input("Quantité", min_value=0.0, value=1.0, step=1.0, key="new_qty")
            
            with col4:
                st.write("Total")
                total_line = unit_price * quantity
                st.write(f"{total_line:,.0f} FCFA")
            
            if st.button("Ajouter à la facture"):
                line_data = {
                    "description": description,
                    "quantity": quantity,
                    "unit_price": unit_price,
                    "total": total_line,
                    "product_id": product_data.get('id') if selected_product != "Service personnalisé..." else None
                }
                st.session_state.invoice_lines.append(line_data)
                st.success("Ligne ajoutée!")
                st.rerun()
        
        # Afficher les lignes de facture
        if st.session_state.invoice_lines:
            st.write("### Lignes de facture")
            
            for i, line in enumerate(st.session_state.invoice_lines):
                col1, col2, col3, col4, col5 = st.columns([3, 1, 1, 1, 1])
                
                with col1:
                    st.write(line['description'])
                with col2:
                    st.write(f"{line['quantity']}")
                with col3:
                    st.write(f"{line['unit_price']:,.0f}")
                with col4:
                    st.write(f"{line['total']:,.0f}")
                with col5:
                    if st.button("🗑️", key=f"delete_{i}"):
                        st.session_state.invoice_lines.pop(i)
                        st.rerun()
            
            # Calculs totaux
            subtotal = sum(line['total'] for line in st.session_state.invoice_lines)
            
            col1, col2 = st.columns([2, 1])
            with col2:
                st.write("### Totaux")
                
                # TVA
                tva_rate = st.number_input("TVA (%)", min_value=0.0, max_value=100.0, value=18.0, step=0.1)
                tva_amount = subtotal * (tva_rate / 100)
                
                # Remise
                discount_rate = st.number_input("Remise (%)", min_value=0.0, max_value=100.0, value=0.0, step=0.1)
                discount_amount = subtotal * (discount_rate / 100)
                
                # Calculs finaux
                subtotal_after_discount = subtotal - discount_amount
                total_ttc = subtotal_after_discount + tva_amount
                
                st.write(f"**Sous-total:** {subtotal:,.0f} FCFA")
                if discount_amount > 0:
                    st.write(f"**Remise ({discount_rate}%):** -{discount_amount:,.0f} FCFA")
                    st.write(f"**Sous-total après remise:** {subtotal_after_discount:,.0f} FCFA")
                st.write(f"**TVA ({tva_rate}%):** {tva_amount:,.0f} FCFA")
                st.write(f"**TOTAL TTC:** {total_ttc:,.0f} FCFA")
        
        # Boutons d'action
        if st.session_state.invoice_lines:
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("💾 Sauvegarder"):
                    self.save_invoice(
                        selected_client, invoice_number, invoice_date, due_date,
                        doc_type, status, payment_terms, st.session_state.invoice_lines,
                        subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total_ttc
                    )
                # Bouton de validation avec déduction du stock (pour Facture)
                if doc_type == "Facture" and st.button("✅ Valider et déduire du stock"):
                    self.save_invoice(
                        selected_client, invoice_number, invoice_date, due_date,
                        doc_type, "Validé", payment_terms, st.session_state.invoice_lines,
                        subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total_ttc
                    )
            
            with col2:
                if st.button("📄 Générer PDF"):
                    pdf_data = self.generate_pdf(
                        selected_client, invoice_number, invoice_date, due_date,
                        doc_type, st.session_state.invoice_lines,
                        subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total_ttc
                    )
                    if pdf_data:
                        st.download_button(
                            label="📥 Télécharger PDF",
                            data=pdf_data,
                            file_name=f"{doc_type}_{invoice_number}.pdf",
                            mime="application/pdf"
                        )
                # Nouveau bouton: générer le document Word (DOCX)
                if st.button("📄 Générer Word (DOCX)"):
                    docx_data = self.generate_docx(
                        selected_client, invoice_number, invoice_date, due_date,
                        doc_type, st.session_state.invoice_lines,
                        subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total_ttc
                    )
                    if docx_data:
                        st.download_button(
                            label="📥 Télécharger DOCX",
                            data=docx_data,
                            file_name=f"{doc_type}_{invoice_number}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            with col3:
                if st.button("🔄 Nouvelle facture"):
                    st.session_state.invoice_lines = []
                    st.rerun()
    
    def save_invoice(self, client, number, date, due_date, doc_type, status, terms, lines, 
                    subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total):
        """Sauvegarde la facture dans Firebase"""
        try:
            # Garder le numéro de facture en session pour le motif des mouvements
            st.session_state['current_invoice_number'] = number
            
            invoice_data = {
                "numero": number,
                "client": client,
                "date": date.isoformat(),
                "date_echeance": due_date.isoformat(),
                "type": doc_type,
                "statut": status,
                "conditions_paiement": terms,
                "lignes": lines,
                "sous_total": subtotal,
                "taux_tva": tva_rate,
                "montant_tva": tva_amount,
                "taux_remise": discount_rate,
                "montant_remise": discount_amount,
                "total_ttc": total,
                "date_creation": datetime.now().isoformat(),
                "utilisateur": st.session_state.get("user_email", "admin"),
                # Indicateur de mise à jour du stock selon le statut
                "stock_updated": True if (doc_type == "Facture" and status in ["Validé", "Payé"]) else False
            }
            
            if save_invoice_to_firebase(invoice_data):
                st.success(f"{doc_type} sauvegardée avec succès!")
                
                # Mettre à jour le stock pour facture validée ou payée
                if doc_type == "Facture" and status in ["Validé", "Payé"]:
                    self.update_stock_from_invoice(lines)
                
                return True
            else:
                st.error("Erreur lors de la sauvegarde")
                return False
                
        except Exception as e:
            st.error(f"Erreur: {e}")
            return False
    
    def update_stock_from_invoice(self, lines):
        """Met à jour le stock après une facture validée/payée"""
        try:
            products = get_all_products_from_firebase()
            
            for line in lines:
                if line.get('product_id'):
                    # Récupérer le produit directement par son ID depuis le dictionnaire
                    product = products.get(line['product_id']) if products else None
                    if product:
                        qty = int(line['quantity']) if isinstance(line.get('quantity'), (int, float)) else 0
                        # Décrémenter le stock
                        new_stock = max(0, product.get('stock_actuel', 0) - qty)
                        product['stock_actuel'] = new_stock
                        
                        # Mettre à jour dans Firebase
                        update_product_in_firebase(product['id'], product)
                        
                        # Enregistrer le mouvement
                        movement = {
                            "produit_id": product['id'],
                            "produit_nom": product.get('nom', ''),
                            "type": "sortie",
                            "quantite": qty,
                            "motif": f"Vente - Facture {st.session_state.get('current_invoice_number', '')}",
                            "date": datetime.now().isoformat(),
                            "utilisateur": st.session_state.get("user_email", "admin")
                        }
                        save_stock_movement_to_firebase(movement)
            
            st.success("Stock mis à jour automatiquement")
            
        except Exception as e:
            st.error(f"Erreur lors de la mise à jour du stock: {e}")
    
    def generate_pdf(self, client, number, date, due_date, doc_type, lines, 
                    subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total):
        """Génère un PDF de la facture"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, 
                                  topMargin=72, bottomMargin=18)
            
            # Contenu du PDF
            story = []
            
            # En-tête
            story.append(Paragraph(f"{doc_type.upper()}", self.styles['CustomTitle']))
            story.append(Spacer(1, 12))
            
            # Informations de base
            info_data = [
                ['Numéro:', number, 'Date:', date.strftime('%d/%m/%Y')],
                ['Client:', client, 'Échéance:', due_date.strftime('%d/%m/%Y')],
            ]
            
            info_table = Table(info_data, colWidths=[1*inch, 2*inch, 1*inch, 1.5*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ]))
            
            story.append(info_table)
            story.append(Spacer(1, 20))
            
            # Tableau des lignes
            table_data = [['Description', 'Qté', 'Prix Unit.', 'Total']]
            
            for line in lines:
                table_data.append([
                    line['description'],
                    f"{line['quantity']:.0f}",
                    f"{line['unit_price']:,.0f}",
                    f"{line['total']:,.0f}"
                ])
            
            # Ligne de sous-total
            table_data.append(['', '', 'Sous-total:', f"{subtotal:,.0f} FCFA"])
            
            if discount_amount > 0:
                table_data.append(['', '', f'Remise ({discount_rate}%):', f"-{discount_amount:,.0f} FCFA"])
            
            table_data.append(['', '', f'TVA ({tva_rate}%):', f"{tva_amount:,.0f} FCFA"])
            table_data.append(['', '', 'TOTAL TTC:', f"{total:,.0f} FCFA"])
            
            table = Table(table_data, colWidths=[3*inch, 0.8*inch, 1.2*inch, 1.2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -4), colors.beige),
                ('BACKGROUND', (0, -3), (-1, -1), colors.lightgrey),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            
            # Construire le PDF
            doc.build(story)
            
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return pdf_data
            
        except Exception as e:
            st.error(f"Erreur lors de la génération du PDF: {e}")
            return None

    def generate_docx(self, client, number, date, due_date, doc_type, lines, 
                    subtotal, tva_rate, tva_amount, discount_rate, discount_amount, total):
        """Génère un document Word (DOCX) de la facture"""
        try:
            doc = Document()
            
            # Titre
            title = doc.add_heading(f"{doc_type.upper()}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informations de base
            info = doc.add_paragraph()
            info.add_run(f"Numéro: ").bold = True
            info.add_run(str(number))
            info.add_run("\t\tDate: ").bold = True
            info.add_run(date.strftime('%d/%m/%Y'))
            info2 = doc.add_paragraph()
            info2.add_run(f"Client: ").bold = True
            info2.add_run(str(client))
            info2.add_run("\t\tÉchéance: ").bold = True
            info2.add_run(due_date.strftime('%d/%m/%Y'))
            
            doc.add_paragraph("")
            
            # Tableau des lignes
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Qté'
            hdr_cells[2].text = 'Prix Unit.'
            hdr_cells[3].text = 'Total'
            
            for line in lines:
                row_cells = table.add_row().cells
                row_cells[0].text = str(line.get('description', ''))
                row_cells[1].text = f"{line.get('quantity', 0):.0f}"
                row_cells[2].text = f"{line.get('unit_price', 0):,.0f}"
                row_cells[3].text = f"{line.get('total', 0):,.0f}"
            
            doc.add_paragraph("")
            
            # Totaux
            totals = doc.add_paragraph()
            totals.add_run(f"Sous-total: ").bold = True
            totals.add_run(f"{subtotal:,.0f} FCFA\n")
            if discount_amount > 0:
                disc = doc.add_paragraph()
                disc.add_run(f"Remise ({discount_rate}%): ").bold = True
                disc.add_run(f"-{discount_amount:,.0f} FCFA\n")
            tva = doc.add_paragraph()
            tva.add_run(f"TVA ({tva_rate}%): ").bold = True
            tva.add_run(f"{tva_amount:,.0f} FCFA\n")
            total_par = doc.add_paragraph()
            total_par.add_run("TOTAL TTC: ").bold = True
            total_par.add_run(f"{total:,.0f} FCFA")
            
            # Export en mémoire
            buffer = io.BytesIO()
            doc.save(buffer)
            data = buffer.getvalue()
            buffer.close()
            return data
            
        except Exception as e:
            st.error(f"Erreur lors de la génération DOCX: {e}")
            return None
    
    def display_existing_invoices(self):
        """Affiche les factures existantes"""
        st.write("### Factures Existantes")
        
        invoices = get_all_invoices_from_firebase()
        
        if invoices:
            # Filtres
            col1, col2, col3 = st.columns(3)
            
            with col1:
                status_filter = st.selectbox("Filtrer par statut", 
                                           ["Tous"] + list(set(inv.get('statut', '') for inv in invoices)))
            
            with col2:
                type_filter = st.selectbox("Filtrer par type", 
                                         ["Tous"] + list(set(inv.get('type', '') for inv in invoices)))
            
            with col3:
                search_term = st.text_input("Rechercher (numéro, client)")
            
            # Filtrer les factures
            filtered_invoices = invoices
            
            if status_filter != "Tous":
                filtered_invoices = [inv for inv in filtered_invoices if inv.get('statut') == status_filter]
            
            if type_filter != "Tous":
                filtered_invoices = [inv for inv in filtered_invoices if inv.get('type') == type_filter]
            
            if search_term:
                filtered_invoices = [inv for inv in filtered_invoices 
                                   if search_term.lower() in inv.get('numero', '').lower() 
                                   or search_term.lower() in inv.get('client', '').lower()]
            
            # Afficher les factures
            for invoice in filtered_invoices:
                with st.expander(f"{invoice.get('type', '')} {invoice.get('numero', '')} - {invoice.get('client', '')}"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**Date:** {invoice.get('date', '')}")
                        st.write(f"**Statut:** {invoice.get('statut', '')}")
                        st.write(f"**Total:** {invoice.get('total_ttc', 0):,.0f} FCFA")
                    
                    with col2:
                        if st.button(f"Régénérer PDF", key=f"pdf_{invoice.get('id', '')}"):
                            # Régénérer le PDF avec les données sauvegardées
                            st.info("Fonctionnalité de régénération PDF à implémenter")
        else:
            st.info("Aucune facture trouvée")
    
    def invoice_templates(self):
        """Gestion des modèles de factures"""
        st.write("### Modèles de Factures")
        
        # Modèles prédéfinis
        templates = {
            "Installation Solaire Résidentielle": [
                {"description": "Panneaux solaires 320W", "quantity": 4, "unit_price": 180000},
                {"description": "Onduleur hybride 3KVA", "quantity": 1, "unit_price": 400000},
                {"description": "Batteries Lithium 100Ah", "quantity": 2, "unit_price": 450000},
                {"description": "Installation et mise en service", "quantity": 1, "unit_price": 200000},
            ],
            "Kit Solaire Basique": [
                {"description": "Panneau solaire 100W", "quantity": 2, "unit_price": 75000},
                {"description": "Régulateur MPPT 20A", "quantity": 1, "unit_price": 45000},
                {"description": "Batterie AGM 100Ah", "quantity": 1, "unit_price": 110000},
                {"description": "Onduleur 1000W", "quantity": 1, "unit_price": 150000},
            ]
        }
        
        selected_template = st.selectbox("Choisir un modèle", ["Aucun"] + list(templates.keys()))
        
        if selected_template != "Aucun" and st.button("Appliquer le modèle"):
            st.session_state.invoice_lines = []
            for item in templates[selected_template]:
                line_data = {
                    "description": item["description"],
                    "quantity": item["quantity"],
                    "unit_price": item["unit_price"],
                    "total": item["quantity"] * item["unit_price"],
                    "product_id": None
                }
                st.session_state.invoice_lines.append(line_data)
            
            st.success(f"Modèle '{selected_template}' appliqué!")
            st.rerun()

# Fonction principale pour intégrer dans l'application
def show_invoice_editor():
    """Fonction principale pour afficher l'éditeur de factures"""
    editor = InvoiceEditor()
    editor.create_invoice_interface()
